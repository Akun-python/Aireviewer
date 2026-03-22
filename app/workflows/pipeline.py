from __future__ import annotations

from dataclasses import dataclass

import difflib
import hashlib
import json
import os
import re
import shutil
from pathlib import Path
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed

from app.agents.prompts import build_system_prompt
from app.agents.supervisor import build_agent
from app.settings import AppSettings
from app.tools.agent_tools import build_plan_tools, build_tools
from app.tools.doc_map import build_indexed_sections
from app.tools.json_extract import (
    extract_json_list as extract_json_list_from_text,
    extract_json_object as extract_json_object_from_text,
)
from app.tools.path_utils import (
    copy_to_workspace,
    ensure_parent,
    ensure_workspace_dir,
    is_within_root,
    resolve_path,
    to_virtual_path,
)
from app.tools.revision_engine import apply_revisions
from app.tools.revision_policy import normalize_paragraph
from app.tools.comment_cleaner import clean_comment_text
from app.tools.docx_comments import strip_docx_comments
from app.tools.docx_images import extract_docx_images
from app.tools.table_elements import extract_table_elements
from app.tools.win32_utils import is_com_call_rejected, is_com_not_initialized
from app.formatting.profiles import apply_format_profile, resolve_profile


BASE_EXPERT_VIEW = "文档审阅员"
BASE_CONSTRAINTS = [
    "仅审阅正文段落，不审阅表格/图表/图注/表注中的内容。",
    "标题、图标题、表标题需精准识别并跳过。",
    "审阅时必须结合上下文（前后段落/句子），避免孤立修改。",
    "每处修改必须在批注中说明具体原因，涉及多个句子时逐句说明。",
    "保持原意与事实，不做无依据改写或扩写。",
    "仅在确有问题时修改，避免为了凑句数或固定口吻改写。",
    "不要因空格、标点或格式微调给出批注，重点审阅逻辑与表达。",
    "标题符号/序号/段首编号保持原样，不要改动。",
    "批注必须指向具体句子或整段，并给出明确修改建议。",
    "句级批注需包含原句关键片段引用，确保批注与句子严格对应。",
    "禁止大段删减，删减仅限单句或词语。",
]

SUMMARY_TEMPLATE = """{
  "document": "<input_filename>",
  "expert_view": "<expert_view>",
  "intent": "<user_intent>",
  "constraints": ["<constraint>", "..."],
  "sections": [
    {
      "title": "<section_title>",
      "changes": ["<change_1>", "<change_2>"],
      "comments": ["<comment_1>", "<comment_2>"],
      "risks": ["<risk_1>"]
    }
  ],
  "overall_risks": ["<risk>"] ,
  "final_output": "<output_path>"
}"""

def _parallel_config() -> tuple[bool, int, int, int]:
    enabled = os.getenv("REVIEW_PARALLEL", "true").lower() in {"1", "true", "yes"}
    min_paragraphs = int(os.getenv("REVIEW_PARALLEL_MIN_PARAGRAPHS", "80"))
    workers = int(os.getenv("REVIEW_PARALLEL_WORKERS", "4"))
    chunk_size = int(os.getenv("REVIEW_SECTION_CHUNK_SIZE", "40"))
    return enabled, min_paragraphs, workers, chunk_size


def _chunk_context_size() -> int:
    value = int(os.getenv("REVIEW_CHUNK_CONTEXT", "2"))
    return max(0, min(10, value))


def _inline_context_mode() -> str:
    value = os.getenv("REVIEW_INLINE_CONTEXT", "boundary").strip().lower()
    if value in {"none", "boundary", "all"}:
        return value
    return "boundary"


def _memory_scope() -> str:
    """
    Controls how agent thread_id is assigned for section planning:

    - off: each chunk uses an independent random thread_id (no memory)
    - run: stable within this run only
    - session: stable across runs (settings.thread_id)
    - document: stable per document content (settings.thread_id + doc fingerprint)
    """

    value = os.getenv("REVIEW_MEMORY_SCOPE", "document").strip().lower()
    if value in {"off", "run", "session", "document"}:
        return value
    return "document"


def _fingerprint_doc(path: Path) -> str:
    """
    Generate a short, mostly-stable fingerprint for a docx to scope memory.
    Defaults to hashing extracted text (stable across docx re-saves); can be tuned via:
      - REVIEW_MEMORY_FINGERPRINT_MODE= text|bytes (default: text)
      - REVIEW_MEMORY_FINGERPRINT_TEXT_CHARS (default: 200000, 0 = no limit)
      - REVIEW_MEMORY_FINGERPRINT_BYTES (bytes-mode only; 0 = full file)
    """

    hasher = hashlib.sha1()

    # If we have a stable per-document id embedded in the file, prefer it.
    prop_name = (os.getenv("REVIEW_MEMORY_DOC_ID_PROP", "") or "ReviewerDocId").strip() or "ReviewerDocId"
    try:
        from app.tools.docx_custom_props import read_custom_prop

        doc_id = read_custom_prop(path, prop_name)
    except Exception:
        doc_id = None
    if doc_id:
        hasher.update(str(doc_id).encode("utf-8", errors="ignore"))
        return hasher.hexdigest()[:12]

    mode = os.getenv("REVIEW_MEMORY_FINGERPRINT_MODE", "text").strip().lower() or "text"
    if mode not in {"text", "bytes"}:
        mode = "text"

    if mode == "text":
        max_chars_raw = os.getenv("REVIEW_MEMORY_FINGERPRINT_TEXT_CHARS", "200000").strip()
        max_chars = 200000
        try:
            max_chars = int(max_chars_raw)
        except Exception:
            max_chars = 200000
        max_chars = max(0, int(max_chars))
        try:
            from docx import Document  # type: ignore

            doc = Document(str(path))
            total = 0
            chunks: list[str] = []
            def add_text(value: str) -> bool:
                nonlocal total
                text = (value or "").strip()
                if not text:
                    return False
                chunks.append(text)
                total += len(text)
                return bool(max_chars and total >= max_chars)

            for para in doc.paragraphs:
                if add_text(str(getattr(para, "text", "") or "")):
                    break

            # python-docx Document.paragraphs does not include text inside tables.
            # Include table cell text to avoid fingerprint collisions for table-heavy documents.
            if not (max_chars and total >= max_chars):
                def iter_tables(tables):
                    for tbl in list(tables or []):
                        yield tbl
                        try:
                            rows = list(getattr(tbl, "rows", []) or [])
                        except Exception:
                            rows = []
                        for row in rows:
                            try:
                                cells = list(getattr(row, "cells", []) or [])
                            except Exception:
                                cells = []
                            for cell in cells:
                                try:
                                    nested = list(getattr(cell, "tables", []) or [])
                                except Exception:
                                    nested = []
                                yield from iter_tables(nested)

                for tbl in iter_tables(getattr(doc, "tables", []) or []):
                    try:
                        rows = list(getattr(tbl, "rows", []) or [])
                    except Exception:
                        rows = []
                    done = False
                    for row in rows:
                        if done:
                            break
                        try:
                            cells = list(getattr(row, "cells", []) or [])
                        except Exception:
                            cells = []
                        for cell in cells:
                            if add_text(str(getattr(cell, "text", "") or "")):
                                done = True
                                break
                    if done:
                        break
            hasher.update("\n".join(chunks).encode("utf-8", errors="ignore"))
            return hasher.hexdigest()[:12]
        except Exception:
            # Fall back to bytes-mode below.
            mode = "bytes"

    limit_raw = os.getenv("REVIEW_MEMORY_FINGERPRINT_BYTES", "").strip()
    limit = 0
    if limit_raw:
        try:
            limit = int(limit_raw)
        except Exception:
            limit = 0
    limit = max(0, int(limit))
    try:
        remaining = limit if limit > 0 else None
        with open(path, "rb") as handle:
            while True:
                chunk_size = 1024 * 1024
                if remaining is not None:
                    if remaining <= 0:
                        break
                    chunk_size = min(chunk_size, remaining)
                data = handle.read(chunk_size)
                if not data:
                    break
                hasher.update(data)
                if remaining is not None:
                    remaining -= len(data)
        return hasher.hexdigest()[:12]
    except Exception:
        return hashlib.sha1(str(path).encode("utf-8", errors="ignore")).hexdigest()[:12]


def _plan_thread_id_base(settings: AppSettings, input_real: Path) -> tuple[str | None, str]:
    scope = _memory_scope()
    if scope == "off":
        return None, scope
    if scope == "run":
        return uuid.uuid4().hex, scope
    if scope == "session":
        return str(settings.thread_id), scope
    fingerprint = _fingerprint_doc(input_real)
    return f"{settings.thread_id}:doc:{fingerprint}", scope


def _build_resume_key(
    settings: AppSettings,
    input_real: Path,
    intent: str,
    expert_view: str,
    constraints: list[str],
    *,
    allow_expansion: bool,
    expansion_level: str,
    allow_web_search: bool,
) -> str:
    payload = {
        "doc_fingerprint": _fingerprint_doc(input_real),
        "intent": (intent or "").strip(),
        "expert_view": (expert_view or "").strip(),
        "constraints": [str(item).strip() for item in constraints if str(item).strip()],
        "model": (settings.model or "").strip(),
        "revision_engine": (settings.revision_engine or "").strip(),
        "allow_expansion": bool(allow_expansion),
        "expansion_level": (expansion_level or "").strip(),
        "allow_web_search": bool(allow_web_search),
        "chunk_context": _chunk_context_size(),
        "inline_context": _inline_context_mode(),
    }
    hasher = hashlib.sha1()
    hasher.update(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8", errors="ignore"))
    return hasher.hexdigest()[:24]


def _maybe_embed_memory_doc_id(path: Path) -> None:
    if path.suffix.lower() not in {".docx", ".docm", ".dotx", ".dotm"}:
        return
    enabled = os.getenv("REVIEW_MEMORY_EMBED_DOC_ID", "true").strip().lower() in {"1", "true", "yes", "y", "on"}
    if not enabled:
        return
    prop_name = (os.getenv("REVIEW_MEMORY_DOC_ID_PROP", "") or "ReviewerDocId").strip() or "ReviewerDocId"
    try:
        from app.tools.docx_custom_props import ensure_custom_prop

        ensure_custom_prop(path, prop_name)
    except Exception:
        return


def _append_log(message: str) -> None:
    log_path = os.getenv("REVISION_LOG_PATH", "").strip()
    if not log_path:
        return
    try:
        with open(log_path, "a", encoding="utf-8") as handle:
            handle.write(message.rstrip() + "\n")
    except Exception:
        pass


def _normalize_for_compare(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


_SENTENCE_PATTERN = re.compile(r"[^。！？!?；;]+[。！？!?；;]?")
_COMMENT_TARGET_PATTERN = re.compile(
    r"(第\d+句\s*[:：])"
    r"|(整段\s*[:：])"
    r"|([Pp]\s*\d+\s*-\s*[Ss]\s*\d+\s*[:：]?)"
    r"|([Pp]\s*\d+\s*-\s*整段\s*[:：]?)"
)
_SENTENCE_ID_PATTERN = re.compile(r"[Pp]\s*(\d+)\s*-\s*[Ss]\s*(\d+)")
_PARAGRAPH_ID_PATTERN = re.compile(r"[Pp]\s*(\d+)\s*-\s*整段")
_CHAPTER_HINT_PATTERN = re.compile(r"第\s*([一二三四五六七八九十零〇两0-9]+)\s*(章|节|部分|篇|条)")
_PARAGRAPH_HINT_PATTERN = re.compile(r"第\s*(\d+)\s*段(?!\s*第\s*\d+\s*句)")
_PARAGRAPH_SENTENCE_HINT_PATTERN = re.compile(r"第\s*(\d+)\s*段\s*第\s*(\d+)\s*句")
_PARAGRAPH_ONLY_HINT_PATTERN = re.compile(r"[Pp]\s*(\d+)\s*-\s*整段")
_PARAGRAPH_SENTENCE_ID_HINT_PATTERN = re.compile(r"[Pp]\s*(\d+)\s*-\s*[Ss]\s*(\d+)")


def _sentence_count(text: str) -> int:
    return len([m.group(0) for m in _SENTENCE_PATTERN.finditer(text or "") if m.group(0).strip()])


def _split_sentences(text: str) -> list[str]:
    return [m.group(0).strip() for m in _SENTENCE_PATTERN.finditer(text or "") if m.group(0).strip()]


def _cn_number_to_int(value: str) -> int | None:
    if not value:
        return None
    if value.isdigit():
        return int(value)
    digits = {
        "零": 0,
        "〇": 0,
        "一": 1,
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
    }
    total = 0
    current = 0
    seen = False
    for char in value:
        if char in digits:
            current = digits[char]
            seen = True
            continue
        if char == "十":
            seen = True
            if current == 0:
                current = 1
            total += current * 10
            current = 0
            continue
        return None
    if not seen:
        return None
    return total + current


def _extract_focus_from_intent(intent: str, sections: list) -> tuple[set[int], set[str], list[str]]:
    paragraph_indices: set[int] = set()
    section_titles: set[str] = set()
    hints: list[str] = []
    if not intent:
        return paragraph_indices, section_titles, hints

    for match in _PARAGRAPH_SENTENCE_ID_HINT_PATTERN.finditer(intent):
        paragraph_indices.add(int(match.group(1)))
        hints.append(f"P{match.group(1)}-S{match.group(2)}")

    for match in _PARAGRAPH_ONLY_HINT_PATTERN.finditer(intent):
        paragraph_indices.add(int(match.group(1)))
        hints.append(f"P{match.group(1)}-整段")

    for match in _PARAGRAPH_SENTENCE_HINT_PATTERN.finditer(intent):
        paragraph_idx = int(match.group(1)) - 1
        paragraph_indices.add(paragraph_idx)
        hints.append(f"第{match.group(1)}段第{match.group(2)}句")

    for match in _PARAGRAPH_HINT_PATTERN.finditer(intent):
        paragraph_idx = int(match.group(1)) - 1
        paragraph_indices.add(paragraph_idx)
        hints.append(f"第{match.group(1)}段")

    for match in _CHAPTER_HINT_PATTERN.finditer(intent):
        number = _cn_number_to_int(match.group(1))
        label = match.group(0)
        if number is not None:
            label = f"第{number}{match.group(2)}"
        for section in sections:
            if label in section.title:
                section_titles.add(section.title)
        hints.append(label)

    for quote in _extract_quoted_text(intent):
        for section in sections:
            if quote and quote in section.title:
                section_titles.add(section.title)
                hints.append(f"标题包含“{quote}”")

    return paragraph_indices, section_titles, hints


def _should_enable_focus_filter(intent: str) -> bool:
    env_value = os.getenv("REVIEW_ENABLE_FOCUS_FILTER", "").strip().lower()
    if env_value in {"1", "true", "yes", "y", "on"}:
        return True
    if env_value in {"0", "false", "no", "n", "off"}:
        return False
    keywords = ("只审阅", "仅审阅", "仅检查", "重点审阅", "聚焦", "focus only", "only review")
    return any(keyword in (intent or "") for keyword in keywords)


def _filter_sections_by_titles(sections: list, titles: set[str]) -> list:
    if not titles:
        return sections
    filtered: list = []
    for section in sections:
        if section.title in titles:
            filtered.append(section)
    return filtered


def _filter_sections_by_paragraph_indices(sections: list, indices: set[int]) -> list:
    if not indices:
        return sections
    filtered: list = []
    for section in sections:
        paragraphs = [item for item in section.paragraphs if item.index in indices]
        if not paragraphs:
            continue
        filtered.append(type(section)(title=section.title, level=section.level, paragraphs=paragraphs))
    return filtered


def _should_allow_expansion(intent: str) -> bool:
    keywords = ("扩充", "补充", "完善", "增补", "填充", "扩写", "补写", "丰富")
    return any(keyword in (intent or "") for keyword in keywords)


def _should_allow_web_search(intent: str) -> bool:
    keywords = ("网络搜索", "联网", "检索", "查证", "引用", "来源", "数据", "最新")
    return any(keyword in (intent or "") for keyword in keywords)


def _normalize_expansion_level(value: str | None) -> str:
    raw = (value or "").strip().lower()
    if raw in {"none", "no", "false", "0"}:
        return "none"
    if raw in {"light", "lite", "small", "minor", "轻量", "轻度"}:
        return "light"
    if raw in {"heavy", "large", "full", "major", "大量", "重度"}:
        return "heavy"
    return "none"


def _is_large_deletion(original: str, revised: str) -> bool:
    original = (original or "").strip()
    revised = (revised or "").strip()
    if not original:
        return False
    if not revised:
        return len(original) >= 30 or _sentence_count(original) >= 2
    original_len = len(original)
    revised_len = len(revised)
    original_sentences = _sentence_count(original)
    revised_sentences = _sentence_count(revised)
    if original_sentences >= 2 and (original_sentences - revised_sentences) >= 2:
        return True
    if original_len >= 80 and revised_len / max(original_len, 1) < 0.5:
        return True
    if original_sentences <= 1 and original_len >= 60 and revised_len / max(original_len, 1) < 0.5:
        return True
    return False


def _is_large_rewrite(original: str, revised: str) -> bool:
    original = (original or "").strip()
    revised = (revised or "").strip()
    if not original or not revised:
        return False
    original_sentences = _sentence_count(original)
    revised_sentences = _sentence_count(revised)
    if abs(original_sentences - revised_sentences) >= 2:
        return True
    if len(original) >= 120 and len(revised) / max(len(original), 1) < 0.6:
        return True
    ratio = difflib.SequenceMatcher(a=original, b=revised).ratio()
    if len(original) >= 80 and ratio < 0.65:
        return True
    return False


def _comment_is_specific(comment: str) -> bool:
    if not comment:
        return False
    return bool(_COMMENT_TARGET_PATTERN.search(comment))


def _extract_sentence_refs(comment: str) -> list[int]:
    refs: list[int] = []
    for match in re.finditer(r"第(\d+)句", comment or ""):
        try:
            refs.append(int(match.group(1)))
        except Exception:
            continue
    return refs


def _extract_sentence_refs_with_paragraph(comment: str) -> list[tuple[int, int]]:
    refs: list[tuple[int, int]] = []
    for match in _SENTENCE_ID_PATTERN.finditer(comment or ""):
        try:
            refs.append((int(match.group(1)), int(match.group(2))))
        except Exception:
            continue
    return refs


def _extract_paragraph_refs(comment: str) -> list[int]:
    refs: list[int] = []
    for match in _PARAGRAPH_ID_PATTERN.finditer(comment or ""):
        try:
            refs.append(int(match.group(1)))
        except Exception:
            continue
    return refs


def _extract_quoted_text(comment: str) -> list[str]:
    if not comment:
        return []
    pairs = [
        ("“", "”"),
        ("\"", "\""),
        ("「", "」"),
        ("『", "』"),
        ("《", "》"),
        ("‘", "’"),
    ]
    results: list[str] = []
    for left, right in pairs:
        pattern = re.escape(left) + r"(.+?)" + re.escape(right)
        for match in re.findall(pattern, comment):
            text = match.strip()
            if text:
                results.append(text)
    return results


def _sentence_indices_for_quotes(sentences: list[str], quotes: list[str]) -> list[int]:
    indices: list[int] = []
    for quote in quotes:
        for idx, sentence in enumerate(sentences, start=1):
            if quote and quote in sentence:
                indices.append(idx)
                break
    return indices


def _normalize_comment_for_paragraph(comment: str, paragraph_text: str, paragraph_index: int) -> str:
    comment = (comment or "").strip()
    if not comment:
        return comment
    if _SENTENCE_ID_PATTERN.search(comment):
        return _SENTENCE_ID_PATTERN.sub(
            lambda match: f"P{paragraph_index}-S{match.group(2)}",
            comment,
        )
    if _PARAGRAPH_ID_PATTERN.search(comment):
        return _PARAGRAPH_ID_PATTERN.sub(
            lambda _match: f"P{paragraph_index}-整段",
            comment,
        )
    if re.search(r"第\s*\d+\s*句", comment):
        return re.sub(
            r"第\s*(\d+)\s*句",
            lambda match: f"P{paragraph_index}-S{match.group(1)}",
            comment,
        )
    if re.search(r"整段\s*[:：]?", comment):
        cleaned = re.sub(r"^整段\s*[:：\s]*", "", comment).strip()
        return f"P{paragraph_index}-整段：{cleaned or comment}"
    quotes = _extract_quoted_text(comment)
    if quotes:
        sentences = _split_sentences(paragraph_text)
        if sentences:
            indices = _sentence_indices_for_quotes(sentences, quotes)
            unique = sorted(set(indices))
            if len(unique) == 1:
                return f"P{paragraph_index}-S{unique[0]}：{comment}"
            if len(unique) > 1:
                return f"P{paragraph_index}-整段：{comment}"
    return f"P{paragraph_index}-整段：{comment}"


def _normalize_comment_targets(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    normalized: list[dict] = []
    for item in plan:
        comment = (item.get("comment") or "").strip()
        if not comment:
            normalized.append(item)
            continue
        idx = item.get("paragraph_index")
        paragraph_text = index_map.get(idx)
        if paragraph_text is None:
            normalized.append(item)
            continue
        updated = _normalize_comment_for_paragraph(comment, paragraph_text, idx)
        if updated == comment:
            normalized.append(item)
            continue
        cloned = dict(item)
        cloned["comment"] = updated
        normalized.append(cloned)
    return normalized


def _filter_misaligned_comments(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    filtered: list[dict] = []
    for item in plan:
        action = item.get("action")
        if action != "comment":
            filtered.append(item)
            continue
        idx = item.get("paragraph_index")
        comment = (item.get("comment") or "").strip()
        if not comment or idx not in index_map:
            filtered.append(item)
            continue
        paragraph_text = index_map.get(idx, "")
        para_sentence_refs = _extract_sentence_refs_with_paragraph(comment)
        para_refs = _extract_paragraph_refs(comment)
        refs = _extract_sentence_refs(comment)
        is_paragraph = bool(re.search(r"整段\s*[:：]", comment))
        if para_sentence_refs:
            if any(ref_idx != idx for ref_idx, _ in para_sentence_refs):
                continue
            sentence_count = len(_split_sentences(paragraph_text))
            sent_refs = [sent_idx for _p, sent_idx in para_sentence_refs]
            if sentence_count == 0 or any(ref < 1 or ref > sentence_count for ref in sent_refs):
                continue
            quotes = _extract_quoted_text(comment)
            if quotes:
                sentences = _split_sentences(paragraph_text)
                ref_sentences = " ".join(
                    sentences[ref - 1] for ref in sent_refs if 1 <= ref <= len(sentences)
                )
                if any(quote not in ref_sentences for quote in quotes):
                    continue
        elif para_refs:
            if any(ref_idx != idx for ref_idx in para_refs):
                continue
            quotes = _extract_quoted_text(comment)
            if quotes and not any(q in paragraph_text for q in quotes):
                continue
        elif refs:
            sentence_count = len(_split_sentences(paragraph_text))
            if sentence_count == 0 or any(ref < 1 or ref > sentence_count for ref in refs):
                continue
            quotes = _extract_quoted_text(comment)
            if quotes:
                sentences = _split_sentences(paragraph_text)
                ref_sentences = " ".join(sentences[ref - 1] for ref in refs if 1 <= ref <= len(sentences))
                if any(quote not in ref_sentences for quote in quotes):
                    continue
        else:
            quotes = _extract_quoted_text(comment)
            if is_paragraph:
                if quotes and not any(q in paragraph_text for q in quotes):
                    continue
            else:
                continue
        filtered.append(item)
    return filtered


def _index_map(sections: list) -> dict[int, str]:
    mapping: dict[int, str] = {}
    for section in sections:
        for item in section.paragraphs:
            mapping[item.index] = item.text
    return mapping


def _filter_trivial_plan(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    filtered: list[dict] = []
    for item in plan:
        action = item.get("action")
        idx = item.get("paragraph_index")
        comment = (item.get("comment") or "").strip()
        if action == "comment":
            if comment:
                trivial_markers = (
                    "删除句首/句尾空格",
                    "合并多余空格",
                    "合并重复标点",
                    "格式规范化",
                )
                cleaned = comment
                for marker in trivial_markers:
                    cleaned = cleaned.replace(marker, "")
                cleaned = re.sub(r"[；;，,。.\s]+", "", cleaned)
                if not cleaned and any(marker in comment for marker in trivial_markers):
                    continue
            filtered.append(item)
            continue
        if action == "replace" and idx in index_map:
            original = index_map[idx]
            revised = item.get("content") or ""
            if _normalize_for_compare(original) == _normalize_for_compare(revised):
                continue
        filtered.append(item)
    return filtered


def _filter_unspecific_plan(plan: list[dict]) -> list[dict]:
    filtered: list[dict] = []
    for item in plan:
        action = item.get("action")
        comment = (item.get("comment") or "").strip()
        if action in {"comment", "replace", "delete", "insert_after"}:
            if not _comment_is_specific(comment):
                continue
        filtered.append(item)
    return filtered


def _filter_noop_comments(plan: list[dict]) -> list[dict]:
    filtered: list[dict] = []
    noop_markers = (
        "未发现问题",
        "无需修改",
        "无需更改",
        "无正文内容",
        "无须修改",
        "无须更改",
        "无需批注",
        "不需要修改",
        "不需修改",
        "不需要更改",
        "不需更改",
        "无需调整",
    )
    for item in plan:
        action = item.get("action")
        comment = (item.get("comment") or "").strip()
        if action == "comment" and comment:
            if any(marker in comment for marker in noop_markers):
                continue
        filtered.append(item)
    return filtered


def _comment_mentions_change(comment: str) -> bool:
    markers = ("改写", "改为", "改成", "修改为", "新增", "补充", "扩充", "调整为", "替换为")
    return any(marker in (comment or "") for marker in markers)


def _truncate_text(text: str, limit: int = 60) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1] + "…"


def _ensure_change_comments(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    adjusted: list[dict] = []
    for item in plan:
        action = item.get("action")
        idx = item.get("paragraph_index")
        if action not in {"replace", "insert_after", "delete"} or idx not in index_map:
            adjusted.append(item)
            continue
        comment = (item.get("comment") or "").strip()
        if comment and _comment_mentions_change(comment):
            adjusted.append(item)
            continue

        original = index_map.get(idx, "")
        revised = item.get("content") or ""
        if action == "replace":
            orig_snip = _truncate_text(original, 50)
            rev_snip = _truncate_text(revised, 80)
            change_note = f"改写说明：原文“{orig_snip}”→改为“{rev_snip}”"
        elif action == "insert_after":
            rev_snip = _truncate_text(revised, 80)
            change_note = f"新增补充：{rev_snip}"
        else:
            orig_snip = _truncate_text(original, 60)
            change_note = f"删除内容：{orig_snip}"

        prefix = f"P{idx}-整段："
        if comment:
            comment = f"{comment}；{change_note}"
        else:
            comment = f"{prefix}{change_note}"
        cloned = dict(item)
        cloned["comment"] = comment
        adjusted.append(cloned)
    return adjusted


def _limit_large_deletions(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    adjusted: list[dict] = []
    for item in plan:
        action = item.get("action")
        idx = item.get("paragraph_index")
        if action in {"replace", "delete"} and idx in index_map:
            original = index_map.get(idx, "")
            revised = "" if action == "delete" else (item.get("content") or "")
            if _is_large_deletion(original, revised):
                comment = (item.get("comment") or "").strip()
                if comment:
                    comment = f"{comment}（删减过大，需仅删减单句或词语）"
                else:
                    comment = "删减过大，需仅删减单句或词语。"
                adjusted.append(
                    {
                        "action": "comment",
                        "paragraph_index": idx,
                        "comment": comment,
                    }
                )
                continue
        adjusted.append(item)
    return adjusted


def _limit_large_rewrites(plan: list[dict], index_map: dict[int, str]) -> list[dict]:
    adjusted: list[dict] = []
    for item in plan:
        action = item.get("action")
        idx = item.get("paragraph_index")
        if action == "replace" and idx in index_map:
            original = index_map.get(idx, "")
            revised = item.get("content") or ""
            if _is_large_rewrite(original, revised):
                comment = (item.get("comment") or "").strip()
                if not comment:
                    comment = "整段：疑似大段改写，请仅在单句或词语层面修改并给出建议。"
                adjusted.append(
                    {
                        "action": "comment",
                        "paragraph_index": idx,
                        "comment": comment,
                    }
                )
                continue
        adjusted.append(item)
    return adjusted


def _merge_constraints(extra: list[str]) -> list[str]:
    seen = set()
    ordered: list[str] = []
    for item in BASE_CONSTRAINTS + extra:
        key = item.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        ordered.append(key)
    return ordered


def _format_user_request(
    input_path: str,
    output_path: str,
    intent: str,
    expert_view: str,
    constraints: list[str],
    revision_engine: str,
    summary_path: str,
) -> str:
    constraints_block = "\n".join(f"- {item}" for item in constraints if item)
    if constraints_block:
        constraints_block = "Constraints:\n" + constraints_block

    return (
        "Please revise the Word document using the available tools.\n"
        f"Input: {input_path}\n"
        f"Output: {output_path}\n"
        f"Summary: {summary_path}\n"
        f"Revision engine: {revision_engine}\n"
        f"Expert view: {expert_view}\n"
        f"User intent: {intent}\n"
        f"{constraints_block}\n"
        "Steps:\n"
        "1) Build a document map.\n"
        "2) Revise section by section.\n"
        "3) Save a revision summary using the JSON template.\n"
    )


def _auto_approve_interrupts(agent, result, config):
    try:
        from langgraph.types import Command
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("缺少 langgraph，无法处理中断恢复；请先运行：pip install -r requirements.txt") from exc
    while result.get("__interrupt__"):
        interrupts = result["__interrupt__"][0].value
        action_requests = interrupts.get("action_requests", [])
        decisions = [{"type": "approve"} for _ in action_requests]
        result = agent.invoke(Command(resume={"decisions": decisions}), config=config)
    return result


def _extract_json_list(text: str) -> list[dict] | None:
    data = extract_json_list_from_text(text)
    return data if isinstance(data, list) else None


def _extract_json_object(text: str) -> dict | None:
    data = extract_json_object_from_text(text)
    return data if isinstance(data, dict) else None


def _is_unreadable_plan(plan: list[dict]) -> bool:
    if len(plan) != 1:
        return False
    item = plan[0]
    if item.get("action") != "comment":
        return False
    comment = (item.get("comment") or "").strip()
    if not comment:
        return False
    hints = (
        "无法提取",
        "无法进行详细审阅",
        "文档为空",
        "缺少可读",
        "无法读取文档",
    )
    return any(hint in comment for hint in hints)


def _invoke_plan(agent, message: str, *, thread_id: str | None = None) -> list[dict] | None:
    payload = {"messages": [{"role": "user", "content": message}]}
    resolved_thread_id = (thread_id or "").strip() or str(uuid.uuid4())
    plan_config = {"configurable": {"thread_id": resolved_thread_id}}
    result = agent.invoke(payload, config=plan_config)
    content = result["messages"][-1].content if result.get("messages") else ""
    parsed = _extract_json_list(content)
    if parsed is None:
        return None
    if parsed and any(not isinstance(item, dict) for item in parsed):
        filtered = [item for item in parsed if isinstance(item, dict)]
        if not filtered:
            return None
        parsed = filtered
    if _is_unreadable_plan(parsed):
        print("[plan] unreadable response; treating as invalid")
        return None
    return parsed


def _is_context_overflow_error(exc: Exception) -> bool:
    text = str(exc) or ""
    hints = (
        "maximum context length",
        "context length",
        "Please reduce the length of the messages",
        "requested",
        "tokens",
    )
    return any(hint in text for hint in hints)


def _request_revision_plan(
    agent, input_virtual: str, *, thread_id: str | None = None
) -> list[dict] | None:
    message = (
        "Return ONLY a JSON list of revision instructions for apply_revisions. "
        "No markdown, no explanations. Use paragraph_index based on the document. "
        "Only target body paragraphs (ignore tables/figures/captions). "
        "Do not change leading title symbols or numbering in paragraphs. "
        "Comments must include explicit reasons and suggestions; use 'P{paragraph_index}-S{sentence_index}：问题+建议' per sentence. "
        "For sentence-level comments, include a short exact quote from that sentence in Chinese quotes to anchor the comment. "
        "For paragraph-level issues, use 'P{paragraph_index}-整段：问题+建议'. "
        "Use surrounding context from nearby paragraphs; do not edit a sentence in isolation. "
        "Do not delete large spans; deletions must be limited to a single sentence or words. "
        "If no changes are needed, return an empty list []. "
        f"Input: {input_virtual}"
    )
    return _invoke_plan(agent, message, thread_id=thread_id)


def _build_section_message(
    title: str,
    paragraphs: list,
    focus_hint: str | None = None,
    *,
    heading_path: str | None = None,
    context_before: list | None = None,
    context_after: list | None = None,
    allow_web_search: bool = False,
    allow_expansion: bool = False,
    expansion_level: str = "none",
) -> str:
    context_before = context_before or []
    context_after = context_after or []

    inline_ctx_mode = _inline_context_mode()
    max_ctx_chars_raw = os.getenv("REVIEW_CONTEXT_MAX_CHARS", "1200").strip()
    max_ctx_chars = 0
    if max_ctx_chars_raw:
        try:
            max_ctx_chars = int(max_ctx_chars_raw)
        except Exception:
            max_ctx_chars = 0
    max_ctx_chars = max(0, int(max_ctx_chars))

    def _ctx_text(value: str) -> str:
        text = (value or "").strip()
        if not max_ctx_chars or len(text) <= max_ctx_chars:
            return text
        if max_ctx_chars <= 3:
            return "…" * max_ctx_chars
        return text[: max_ctx_chars - 1].rstrip() + "…"

    blocks: list[str] = []
    total_paragraphs = len(paragraphs)
    for idx, item in enumerate(paragraphs):
        sentences = _split_sentences(item.text)
        if sentences:
            block_lines = [f"P{item.index}:"]
            for sent_idx, sent in enumerate(sentences, start=1):
                block_lines.append(f"SENT_{sent_idx} (P{item.index}-S{sent_idx}): {sent}")
        else:
            block_lines = [f"{item.index}: {item.text}"]

        prev_item = None
        next_item = None
        if inline_ctx_mode == "all":
            prev_item = paragraphs[idx - 1] if idx > 0 else (context_before[-1] if context_before else None)
            next_item = (
                paragraphs[idx + 1] if idx + 1 < total_paragraphs else (context_after[0] if context_after else None)
            )
        elif inline_ctx_mode == "boundary":
            prev_item = context_before[-1] if idx == 0 and context_before else None
            next_item = context_after[0] if idx + 1 == total_paragraphs and context_after else None
        if prev_item is not None:
            block_lines.append(f"CTX_PREV {prev_item.index}: {_ctx_text(prev_item.text)}")
        if next_item is not None:
            block_lines.append(f"CTX_NEXT {next_item.index}: {_ctx_text(next_item.text)}")

        blocks.append("\n".join(block_lines))
    section_text = "\n\n".join(blocks)
    hint = ""
    if focus_hint:
        hint = f"{focus_hint}\n"
    tool_rule = "Do not call tools; use only the provided section text. "
    if allow_web_search:
        tool_rule = (
            "You may use the internet_search tool for expansion or verification when needed. "
            "Do not call apply_revisions, build_doc_map, extract_section_text, or save_revision_summary. "
        )
    expansion_rule = ""
    if allow_expansion:
        level_hint = ""
        if expansion_level == "heavy":
            level_hint = (
                "Expand substantially: add new sentences or paragraphs to enrich background, methodology, or "
                "logical transitions; keep coherence with surrounding context. "
                "Ensure the added content is at least 300 Chinese characters in total. "
            )
        elif expansion_level == "light":
            level_hint = "Expand lightly: add 1-2 sentences to improve clarity or continuity. "
        expansion_rule = (
            "Expansion is allowed when the intent requires it; you may add sentences if needed. "
            "When expansion is requested and the internet_search tool is available, use it to gather missing facts. "
            "Ensure new content bridges previous/next paragraphs and keeps the logic consistent. "
            f"{level_hint}"
        )

    ctx_before = context_before
    ctx_after = context_after
    if inline_ctx_mode == "all":
        ctx_before = []
        ctx_after = []
    elif inline_ctx_mode == "boundary":
        ctx_before = context_before[:-1] if len(context_before) > 1 else []
        ctx_after = context_after[1:] if len(context_after) > 1 else []

    ctx_blocks: list[str] = []
    if ctx_before:
        ctx_lines = ["CONTEXT_BEFORE (reference only; do not target these indices):"]
        for item in ctx_before:
            ctx_lines.append(f"CTX_ONLY {item.index}: {_ctx_text(item.text)}")
        ctx_blocks.append("\n".join(ctx_lines))
    if ctx_after:
        ctx_lines = ["CONTEXT_AFTER (reference only; do not target these indices):"]
        for item in ctx_after:
            ctx_lines.append(f"CTX_ONLY {item.index}: {_ctx_text(item.text)}")
        ctx_blocks.append("\n".join(ctx_lines))
    ctx_block_text = ("\n\n".join(ctx_blocks) + "\n\n") if ctx_blocks else ""

    heading_line = f"Heading path: {heading_path}\n" if heading_path else ""
    return (
        "Return ONLY a JSON list of revision instructions for apply_revisions. "
        "No markdown, no explanations. Use ONLY the paragraph_index values provided below. "
        "Only target body paragraphs (ignore tables/figures/captions). "
        "Do not change leading title symbols or numbering in paragraphs. "
        "Context lines start with 'CTX_PREV'/'CTX_NEXT'/'CTX_ONLY' and are for reference only; do not use their indices as targets. "
        "Comments must include explicit reasons and suggestions; use 'P{paragraph_index}-S{sentence_index}：问题+建议' per sentence "
        "and match SENT_X. "
        "For sentence-level comments, include a short exact quote from that sentence in Chinese quotes to anchor the comment. "
        "For paragraph-level issues, use 'P{paragraph_index}-整段：问题+建议'. "
        "For replace/insert_after/delete, include a brief change summary (原文→改写/新增内容) in the comment. "
        "Do not delete large spans; deletions must be limited to a single sentence or words. "
        f"{tool_rule}"
        f"{expansion_rule}"
        "If no changes are needed, return an empty list []. "
        f"{hint}"
        f"Section: {title}\n"
        f"{heading_line}"
        f"{ctx_block_text}"
        "Paragraphs:\n"
        f"{section_text}"
    )


def _fallback_plan_for_paragraphs(paragraphs: list) -> list[dict]:
    plan: list[dict] = []
    for item in paragraphs:
        revised, reasons = normalize_paragraph(item.text)
        if revised != item.text and reasons:
            plan.append(
                {
                    "action": "replace",
                    "paragraph_index": item.index,
                    "content": revised,
                    "comment": "；".join(reasons),
                }
            )
    return plan


def _request_section_plan(
    settings: AppSettings,
    system_prompt: str,
    title: str,
    paragraphs: list,
    focus_hint: str | None = None,
    *,
    heading_path: str | None = None,
    context_before: list | None = None,
    context_after: list | None = None,
    thread_id: str | None = None,
    allow_web_search: bool = False,
    allow_expansion: bool = False,
    expansion_level: str = "none",
) -> list[dict] | None:
    tools = build_plan_tools(allow_web_search=allow_web_search)
    agent = build_agent(settings, tools, system_prompt)
    message = _build_section_message(
        title,
        paragraphs,
        focus_hint,
        heading_path=heading_path,
        context_before=context_before,
        context_after=context_after,
        allow_web_search=allow_web_search,
        allow_expansion=allow_expansion,
        expansion_level=expansion_level,
    )
    return _invoke_plan(agent, message, thread_id=thread_id)


@dataclass(frozen=True)
class SectionChunk:
    title: str
    heading_path: str | None
    paragraphs: list
    context_before: list
    context_after: list


def _heading_paths(sections: list) -> dict[int, str]:
    paths: dict[int, str] = {}
    stack: list[tuple[int, str]] = []
    for idx, section in enumerate(sections):
        level = int(getattr(section, "level", 0) or 0)
        title = (getattr(section, "title", "") or "").strip() or "Untitled"
        if level <= 0:
            stack = []
            paths[idx] = title
            continue
        while stack and stack[-1][0] >= level:
            stack.pop()
        stack.append((level, title))
        paths[idx] = " > ".join(item[1] for item in stack)
    return paths


def _chunk_sections(sections: list, chunk_size: int, context_size: int) -> list[SectionChunk]:
    chunks: list[SectionChunk] = []
    section_paths = _heading_paths(sections)
    for section_idx, section in enumerate(sections):
        paragraphs = section.paragraphs
        if not paragraphs:
            continue
        heading_path = section_paths.get(section_idx) or getattr(section, "title", None)
        if len(paragraphs) <= chunk_size:
            chunks.append(
                SectionChunk(
                    title=section.title,
                    heading_path=heading_path,
                    paragraphs=paragraphs,
                    context_before=[],
                    context_after=[],
                )
            )
            continue
        for offset in range(0, len(paragraphs), chunk_size):
            part = paragraphs[offset : offset + chunk_size]
            before_start = max(0, offset - context_size)
            before = paragraphs[before_start:offset] if context_size > 0 else []
            after_end = min(len(paragraphs), offset + chunk_size + context_size)
            after = paragraphs[offset + chunk_size : after_end] if context_size > 0 else []
            part_title = f"{section.title} (part {offset // chunk_size + 1})"
            chunks.append(
                SectionChunk(
                    title=part_title,
                    heading_path=heading_path,
                    paragraphs=part,
                    context_before=before,
                    context_after=after,
                )
            )
    return chunks


def _filter_sections_by_index(sections: list, start_index: int) -> list:
    filtered: list = []
    for section in sections:
        paragraphs = [item for item in section.paragraphs if item.index >= start_index]
        if not paragraphs:
            continue
        filtered.append(type(section)(title=section.title, level=section.level, paragraphs=paragraphs))
    return filtered


def _merge_plans(plans: list[list[dict]]) -> list[dict]:
    def _hashable(value):
        if value is None or isinstance(value, (str, int, float, bool)):
            return value
        try:
            return json.dumps(value, ensure_ascii=False, sort_keys=True)
        except Exception:
            return str(value)

    def _coerce_index(value, default: int = 0) -> int:
        try:
            return int(value)
        except Exception:
            return default

    seen = set()
    merged: list[dict] = []
    for plan in plans:
        for item in plan:
            if not isinstance(item, dict):
                continue
            paragraph_index = _coerce_index(item.get("paragraph_index", 0), default=0)
            key = (
                str(item.get("action") or ""),
                paragraph_index,
                _hashable(item.get("content")),
                _hashable(item.get("comment")),
            )
            if key in seen:
                continue
            seen.add(key)
            if item.get("paragraph_index") != paragraph_index:
                item = dict(item)
                item["paragraph_index"] = paragraph_index
            merged.append(item)
    merged.sort(key=lambda item: _coerce_index(item.get("paragraph_index", 0), default=0))
    return merged


def _build_parallel_plan(
    settings: AppSettings,
    system_prompt: str,
    sections: list,
    workers: int,
    chunk_size: int,
    focus_hint: str | None = None,
    *,
    thread_id_base: str | None = None,
    allow_web_search: bool = False,
    allow_expansion: bool = False,
    expansion_level: str = "none",
) -> list[dict]:
    chunks = _chunk_sections(sections, chunk_size, _chunk_context_size())
    if not chunks:
        return []
    workers = max(1, min(workers, len(chunks)))
    print(f"[parallel] chunks={len(chunks)} workers={workers}")
    _append_log(f"[parallel] chunks={len(chunks)} workers={workers}")

    plans: list[list[dict]] = []
    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(
                _request_section_plan,
                settings,
                system_prompt,
                chunk.title,
                chunk.paragraphs,
                focus_hint,
                heading_path=chunk.heading_path,
                context_before=chunk.context_before,
                context_after=chunk.context_after,
                thread_id=(
                    f"{thread_id_base}:chunk:P{chunk.paragraphs[0].index}" if (thread_id_base and chunk.paragraphs) else None
                ),
                allow_web_search=allow_web_search,
                allow_expansion=allow_expansion,
                expansion_level=expansion_level,
            ): chunk
            for chunk in chunks
        }
        for future in as_completed(future_map):
            chunk = future_map[future]
            try:
                plan = future.result()
            except Exception as exc:  # noqa: BLE001
                print(f"[parallel] section failed: {chunk.title} ({exc})")
                plan = None
            if plan is None:
                plan = _fallback_plan_for_paragraphs(chunk.paragraphs)
            plans.append(plan)
            print(f"[parallel] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")
            _append_log(f"[parallel] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")

    return _merge_plans(plans)


def _build_sequential_plan(
    settings: AppSettings,
    system_prompt: str,
    sections: list,
    chunk_size: int,
    focus_hint: str | None = None,
    *,
    thread_id_base: str | None = None,
    allow_web_search: bool = False,
    allow_expansion: bool = False,
    expansion_level: str = "none",
) -> list[dict]:
    chunks = _chunk_sections(sections, chunk_size, _chunk_context_size())
    if not chunks:
        return []
    plans: list[list[dict]] = []
    if thread_id_base:
        thread_mode = os.getenv("REVIEW_PLAN_THREAD_MODE", "per_chunk").strip().lower()
        if thread_mode not in {"per_chunk", "shared"}:
            thread_mode = "per_chunk"
        tools = build_plan_tools(allow_web_search=allow_web_search)
        agent = build_agent(settings, tools, system_prompt)
        for chunk in chunks:
            message = _build_section_message(
                chunk.title,
                chunk.paragraphs,
                focus_hint,
                heading_path=chunk.heading_path,
                context_before=chunk.context_before,
                context_after=chunk.context_after,
                allow_web_search=allow_web_search,
                allow_expansion=allow_expansion,
                expansion_level=expansion_level,
            )
            try:
                resolved_thread_id = thread_id_base
                if thread_mode != "shared":
                    # Avoid accumulating the whole document's section planning messages into a
                    # single thread (can exceed model context on long documents).
                    resolved_thread_id = (
                        f"{thread_id_base}:chunk:P{chunk.paragraphs[0].index}"
                        if chunk.paragraphs
                        else f"{thread_id_base}:chunk:{uuid.uuid4().hex}"
                    )
                plan = _invoke_plan(agent, message, thread_id=resolved_thread_id)
            except Exception as exc:  # noqa: BLE001
                original_exc = exc
                if thread_mode == "shared" and _is_context_overflow_error(original_exc):
                    # Auto-fallback: if shared memory overflows, retry once per-chunk.
                    try:
                        resolved_thread_id = (
                            f"{thread_id_base}:chunk:P{chunk.paragraphs[0].index}"
                            if chunk.paragraphs
                            else f"{thread_id_base}:chunk:{uuid.uuid4().hex}"
                        )
                        plan = _invoke_plan(agent, message, thread_id=resolved_thread_id)
                    except Exception:
                        plan = None
                else:
                    plan = None
                if plan is None:
                    print(f"[sequential] section failed: {chunk.title} ({original_exc})")
            if plan is None:
                plan = _fallback_plan_for_paragraphs(chunk.paragraphs)
            plans.append(plan)
            print(f"[sequential] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")
            _append_log(f"[sequential] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")
        return _merge_plans(plans)

    for chunk in chunks:
        try:
            plan = _request_section_plan(
                settings,
                system_prompt,
                chunk.title,
                chunk.paragraphs,
                focus_hint,
                heading_path=chunk.heading_path,
                context_before=chunk.context_before,
                context_after=chunk.context_after,
                allow_web_search=allow_web_search,
                allow_expansion=allow_expansion,
                expansion_level=expansion_level,
            )
        except Exception as exc:  # noqa: BLE001
            print(f"[sequential] section failed: {chunk.title} ({exc})")
            plan = None
        if plan is None:
            plan = _fallback_plan_for_paragraphs(chunk.paragraphs)
        plans.append(plan)
        print(f"[sequential] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")
        _append_log(f"[sequential] section done: {chunk.title} paragraphs={len(chunk.paragraphs)} items={len(plan)}")
    return _merge_plans(plans)


def _summary_from_plan(
    plan: list[dict],
    input_name: str,
    expert_view: str,
    intent: str,
    constraints: list[str],
    output_path: str,
) -> dict:
    changes = []
    comments = []
    for item in plan[:50]:
        action = item.get("action")
        idx = item.get("paragraph_index")
        content = item.get("content")
        comment = item.get("comment")
        changes.append(f"{action} @ {idx}: {content or ''}".strip())
        if comment:
            cleaned = clean_comment_text(comment)
            if cleaned:
                comments.append(cleaned)
    return {
        "document": input_name,
        "expert_view": expert_view,
        "intent": intent,
        "constraints": constraints,
        "sections": [
            {
                "title": "Document",
                "changes": changes,
                "comments": comments,
                "risks": [],
            }
        ],
        "overall_risks": [],
        "final_output": output_path,
    }


def _prepare_paths(input_path: str, output_path: str) -> tuple[Path, Path, Path | None]:
    input_real = copy_to_workspace(input_path)
    output_real = resolve_path(output_path)
    output_target: Path | None = None

    if not is_within_root(output_real):
        workspace = ensure_workspace_dir() / "exports"
        workspace.mkdir(parents=True, exist_ok=True)
        output_target = output_real
        output_real = workspace / output_target.name

    ensure_parent(output_real)
    return input_real, output_real, output_target


def _strip_existing_comments_enabled() -> bool:
    value = os.getenv("STRIP_EXISTING_COMMENTS", "").strip().lower()
    return value in {"1", "true", "yes", "y", "on"}


def _extract_table_elements_enabled() -> bool:
    value = os.getenv("EXTRACT_TABLE_ELEMENTS", "").strip().lower()
    return value in {"1", "true", "yes", "y", "on"}


def _extract_docx_images_enabled() -> bool:
    value = os.getenv("EXTRACT_DOCX_IMAGES", "").strip().lower()
    return value in {"1", "true", "yes", "y", "on"}


def _preprocess_input_docx(input_real: Path) -> Path:
    if input_real.suffix.lower() not in {".docx", ".docm", ".dotx", ".dotm"}:
        return input_real
    workspace = ensure_workspace_dir() / "preprocessed" / uuid.uuid4().hex
    workspace.mkdir(parents=True, exist_ok=True)
    preprocessed = workspace / input_real.name
    shutil.copy2(input_real, preprocessed)
    removed = False
    try:
        removed = strip_docx_comments(preprocessed, prefer_win32=False)
    except Exception:
        removed = False
    if removed:
        print("[preprocess] stripped existing comments")
        _append_log("[preprocess] stripped existing comments")
    else:
        print("[preprocess] no comments stripped")
        _append_log("[preprocess] no comments stripped")
    return preprocessed


def run_revision(
    settings: AppSettings,
    input_path: str,
    output_path: str,
    intent: str,
    expert_view: str,
    constraints: list[str],
    *,
    allow_expansion: bool = False,
    expansion_level: str | None = None,
    allow_web_search: bool = False,
):
    expansion_level_norm = _normalize_expansion_level(expansion_level)
    if expansion_level_norm != "none":
        allow_expansion = True
    if not allow_expansion and _should_allow_expansion(intent):
        allow_expansion = True
        if expansion_level_norm == "none":
            expansion_level_norm = "light"
    if allow_expansion and not allow_web_search:
        allow_web_search = os.getenv("ENABLE_WEB_SEARCH", "").lower() in {"1", "true", "yes"}
    if not allow_web_search and _should_allow_web_search(intent):
        allow_web_search = True
    if allow_expansion and expansion_level_norm == "none":
        expansion_level_norm = "light"
    resolved_expert = expert_view or BASE_EXPERT_VIEW
    merged_constraints = _merge_constraints(constraints)
    system_prompt = build_system_prompt(
        resolved_expert,
        intent,
        merged_constraints,
        SUMMARY_TEMPLATE,
        allow_expansion=allow_expansion,
        expansion_level=expansion_level_norm,
        allow_web_search=allow_web_search,
    )
    os.environ["REVISION_ENGINE"] = settings.revision_engine
    if settings.revision_engine == "python-docx":
        os.environ["DISABLE_WIN32"] = "true"
    else:
        os.environ.pop("DISABLE_WIN32", None)

    input_real, output_real, output_target = _prepare_paths(input_path, output_path)
    memory_source_real = input_real
    if _memory_scope() == "document":
        _maybe_embed_memory_doc_id(memory_source_real)
    if _strip_existing_comments_enabled():
        input_real = _preprocess_input_docx(input_real)
    input_virtual = to_virtual_path(input_real)
    output_virtual = to_virtual_path(output_real)
    summary_real = output_real.with_suffix(".summary.json")
    summary_virtual = to_virtual_path(summary_real)
    tables_real = output_real.with_suffix(".tables.json")
    tables_virtual = to_virtual_path(tables_real)
    table_images_dir = output_real.parent / f"{output_real.stem}_table_images"
    images_real = output_real.with_suffix(".images.json")
    images_virtual = to_virtual_path(images_real)
    doc_images_dir = output_real.parent / f"{output_real.stem}_images"
    log_real = output_real.with_suffix(".log.txt")
    os.environ["REVISION_LOG_PATH"] = str(log_real)
    if log_real.exists():
        try:
            log_real.unlink()
        except Exception:
            pass

    print("[backend] model=", settings.model)
    print("[backend] root_dir=", settings.root_dir)
    print("[backend] use_store=", settings.use_store)
    print("[backend] revision_engine=", settings.revision_engine)
    print("[paths] input_real=", input_real)
    print("[paths] input_virtual=", input_virtual)
    print("[paths] output_real=", output_real)
    print("[paths] output_virtual=", output_virtual)
    print("[paths] summary_virtual=", summary_virtual)
    if _extract_table_elements_enabled():
        print("[paths] tables_virtual=", tables_virtual)
    if _extract_docx_images_enabled():
        print("[paths] images_virtual=", images_virtual)
    if output_target:
        print("[paths] output_target=", output_target)

    resume_dir = ensure_workspace_dir() / "agent_state" / "revision_resume"
    resume_dir.mkdir(parents=True, exist_ok=True)
    resume_key = _build_resume_key(
        settings,
        memory_source_real,
        intent,
        resolved_expert,
        merged_constraints,
        allow_expansion=allow_expansion,
        expansion_level=expansion_level_norm,
        allow_web_search=allow_web_search,
    )
    os.environ.setdefault("REVIEW_RESUME_ENABLED", "true")
    os.environ["REVIEW_RESUME_DIR"] = str(resume_dir)
    os.environ["REVIEW_RESUME_KEY"] = resume_key
    print("[resume] key=", resume_key)
    print("[resume] dir=", resume_dir)
    _append_log(f"[resume] key={resume_key}")

    thread_id_base, memory_scope = _plan_thread_id_base(settings, memory_source_real)
    if thread_id_base:
        print("[memory] scope=", memory_scope)
        _append_log(f"[memory] scope={memory_scope}")
    else:
        memory_scope = "off"
        print("[memory] scope=off")
        _append_log("[memory] scope=off")

    table_extract_enabled = _extract_table_elements_enabled()
    if table_extract_enabled:
        try:
            extract_table_elements(
                str(input_real),
                output_path=str(tables_real),
                images_dir=str(table_images_dir),
            )
            print("[tables] extracted")
            _append_log(f"[tables] extracted: {tables_real}")
        except Exception as exc:  # noqa: BLE001
            message = f"[tables] failed: {exc}"
            print(message)
            _append_log(message)

    images_extract_enabled = _extract_docx_images_enabled()
    if images_extract_enabled:
        try:
            extract_docx_images(
                str(input_real),
                output_path=str(images_real),
                images_dir=str(doc_images_dir),
            )
            print("[images] extracted")
            _append_log(f"[images] extracted: {images_real}")
        except Exception as exc:  # noqa: BLE001
            message = f"[images] failed: {exc}"
            print(message)
            _append_log(message)

    indexed_sections = build_indexed_sections(str(input_real))
    full_body_paragraphs = sum(len(section.paragraphs) for section in indexed_sections)

    focus_paragraphs, focus_sections, focus_hints = _extract_focus_from_intent(intent, indexed_sections)
    focus_hint = None
    enable_focus_filter = _should_enable_focus_filter(intent)
    focus_filter_applied = False
    if focus_hints:
        hint_text = "；".join(dict.fromkeys(focus_hints))
        if enable_focus_filter:
            focus_hint = f"Only review specified targets: {hint_text}"
        else:
            focus_hint = None

    if (focus_sections or focus_paragraphs) and enable_focus_filter:
        focus_filter_applied = True
        if focus_sections:
            indexed_sections = _filter_sections_by_titles(indexed_sections, focus_sections)
        if focus_paragraphs:
            indexed_sections = _filter_sections_by_paragraph_indices(indexed_sections, focus_paragraphs)
        if not indexed_sections:
            indexed_sections = build_indexed_sections(str(input_real))
            print("[focus] no matching targets; fallback to full document")
            _append_log("[focus] no matching targets; fallback to full document")
            focus_hint = None
            focus_filter_applied = False
        filtered_body_paragraphs = sum(len(section.paragraphs) for section in indexed_sections)
        print(
            "[focus] enabled; paragraphs=",
            len(focus_paragraphs),
            "sections=",
            len(focus_sections),
        )
        _append_log(
            f"[focus] enabled; paragraphs={len(focus_paragraphs)} sections={len(focus_sections)}"
        )
        print(f"[focus] filtered body_paragraphs={filtered_body_paragraphs} (full={full_body_paragraphs})")
        _append_log(f"[focus] filtered body_paragraphs={filtered_body_paragraphs} (full={full_body_paragraphs})")
    elif focus_hints and (focus_sections or focus_paragraphs) and not enable_focus_filter:
        print("[focus] detected targets but focus filtering is disabled; reviewing full document")
        _append_log("[focus] detected targets but focus filtering is disabled; reviewing full document")
    paragraph_map = _index_map(indexed_sections)
    body_paragraphs = sum(len(section.paragraphs) for section in indexed_sections)
    print(f"[doc] sections={len(indexed_sections)} body_paragraphs={body_paragraphs}")
    _append_log(f"[doc] sections={len(indexed_sections)} body_paragraphs={body_paragraphs}")
    if body_paragraphs == 0:
        print("[doc] no readable body paragraphs; check tables/text boxes or file format")
        _append_log("[doc] no readable body paragraphs; check tables/text boxes or file format")
    parallel_enabled, min_paragraphs, workers, chunk_size = _parallel_config()
    use_parallel = parallel_enabled and body_paragraphs >= max(1, int(min_paragraphs))
    allow_parallel_with_memory = os.getenv("REVIEW_MEMORY_ALLOW_PARALLEL", "").strip().lower() in {"1", "true", "yes"}
    if thread_id_base and not allow_parallel_with_memory:
        use_parallel = False
        if parallel_enabled:
            print("[memory] force sequential (disable parallel planning)")
            _append_log("[memory] force sequential (disable parallel planning)")
    if use_parallel:
        print(f"[parallel] enabled; paragraphs={body_paragraphs}")
        _append_log(f"[parallel] enabled; paragraphs={body_paragraphs}")
        plan = _build_parallel_plan(
            settings,
            system_prompt,
            indexed_sections,
            workers,
            chunk_size,
            focus_hint,
            thread_id_base=thread_id_base,
            allow_web_search=allow_web_search,
            allow_expansion=allow_expansion,
            expansion_level=expansion_level_norm,
        )
    else:
        print(f"[sequential] enabled; paragraphs={body_paragraphs}")
        _append_log(f"[sequential] enabled; paragraphs={body_paragraphs}")
        plan = _build_sequential_plan(
            settings,
            system_prompt,
            indexed_sections,
            chunk_size,
            focus_hint,
            thread_id_base=thread_id_base,
            allow_web_search=allow_web_search,
            allow_expansion=allow_expansion,
            expansion_level=expansion_level_norm,
        )

    def _safe_plan_index(item: dict, default: int = -1) -> int:
        try:
            return int(item.get("paragraph_index", default))
        except Exception:
            return default

    max_paragraph_index = max(paragraph_map.keys(), default=-1)
    max_plan_index = max((_safe_plan_index(item, -1) for item in plan), default=-1)
    tail_ratio = float(os.getenv("REVIEW_TAIL_COVERAGE_RATIO", "0.7"))
    tail_ratio = max(0.1, min(0.95, tail_ratio))
    if max_paragraph_index > 0 and max_plan_index < int(max_paragraph_index * tail_ratio):
        tail_start = int(max_paragraph_index * tail_ratio)
        tail_sections = _filter_sections_by_index(indexed_sections, tail_start)
        if tail_sections:
            print(f"[tail] recheck from index={tail_start}")
            _append_log(f"[tail] recheck from index={tail_start}")
            focus_hint = (
                "Second pass on later sections. Ensure coverage and do not skip paragraphs. "
                "Only return comments when there are real issues; otherwise return []."
            )
            if use_parallel:
                tail_plan = _build_parallel_plan(
                    settings,
                    system_prompt,
                    tail_sections,
                    workers,
                    chunk_size,
                    focus_hint,
                    thread_id_base=thread_id_base,
                    allow_web_search=allow_web_search,
                    allow_expansion=allow_expansion,
                    expansion_level=expansion_level_norm,
                )
            else:
                tail_plan = _build_sequential_plan(
                    settings,
                    system_prompt,
                    tail_sections,
                    chunk_size,
                    focus_hint,
                    thread_id_base=thread_id_base,
                    allow_web_search=allow_web_search,
                    allow_expansion=allow_expansion,
                    expansion_level=expansion_level_norm,
                )
            if tail_plan:
                plan = _merge_plans([plan, tail_plan])

    plan = _filter_trivial_plan(plan, paragraph_map)
    plan = _limit_large_deletions(plan, paragraph_map)
    plan = _limit_large_rewrites(plan, paragraph_map)
    plan = _normalize_comment_targets(plan, paragraph_map)
    plan = _ensure_change_comments(plan, paragraph_map)
    plan = _filter_unspecific_plan(plan)
    plan = _filter_noop_comments(plan)
    plan = _filter_misaligned_comments(plan, paragraph_map)
    print(
        "[plan] items=",
        len(plan),
        "max_index=",
        max((_safe_plan_index(item, 0) for item in plan), default=-1),
    )
    _append_log(
        f"[plan] items={len(plan)} max_index="
        f"{max((_safe_plan_index(item, 0) for item in plan), default=-1)}"
    )
    revisions_json = json.dumps(plan, ensure_ascii=False)
    apply_revisions(str(input_real), str(output_real), revisions_json, settings.revision_engine)

    format_profile = resolve_profile(getattr(settings, "format_profile", "none"))
    if format_profile != "none":
        print("[format] profile=", format_profile)
        _append_log(f"[format] profile={format_profile}")
        try:
            apply_format_profile(docx_path=str(output_real), root_dir=settings.root_dir, profile=format_profile)
        except Exception as exc:  # noqa: BLE001
            message = f"[format] failed: {exc}"
            if is_com_call_rejected(exc):
                message += "；Word 正忙或存在弹窗，建议关闭所有 Word 窗口/对话框后重试"
            elif is_com_not_initialized(exc):
                message += "；COM 未初始化（多线程调用需在该线程先 CoInitialize）"
            print(message)
            _append_log(message)

    if memory_scope == "document":
        _maybe_embed_memory_doc_id(output_real)

    summary_payload = _summary_from_plan(
        plan,
        input_real.name,
        resolved_expert,
        intent,
        merged_constraints,
        str(output_real),
    )
    summary_payload["doc_stats"] = {
        "sections": len(indexed_sections),
        "body_paragraphs": body_paragraphs,
        "full_body_paragraphs": full_body_paragraphs,
        "focus_filter_enabled": focus_filter_applied,
        "focus_targets_detected": bool(focus_sections or focus_paragraphs),
    }
    summary_payload["memory"] = {
        "scope": memory_scope,
        "thread_id": thread_id_base or "",
        "inline_context": _inline_context_mode(),
    }
    summary_payload["resume"] = {
        "enabled": os.getenv("REVIEW_RESUME_ENABLED", "").strip().lower() in {"1", "true", "yes", "y", "on"},
        "key": resume_key,
        "dir": str(resume_dir),
    }
    if table_extract_enabled:
        summary_payload["table_elements"] = {
            "enabled": True,
            "path": str(tables_real) if tables_real.exists() else "",
            "images_dir": str(table_images_dir) if table_images_dir.exists() else "",
        }
    if images_extract_enabled:
        summary_payload["docx_images"] = {
            "enabled": True,
            "path": str(images_real) if images_real.exists() else "",
            "images_dir": str(doc_images_dir) if doc_images_dir.exists() else "",
        }
    ensure_parent(summary_real)
    summary_real.write_text(
        json.dumps(summary_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    if output_target and output_real.exists():
        ensure_parent(output_target)
        shutil.copy2(output_real, output_target)
        tables_target = output_target.with_suffix(".tables.json")
        images_target = output_target.with_suffix(".images.json")
        table_images_target_dir = output_target.parent / f"{output_target.stem}_table_images"
        doc_images_target_dir = output_target.parent / f"{output_target.stem}_images"

        summary_target = output_target.with_suffix(".summary.json")
        if table_extract_enabled and tables_real.exists():
            try:
                ensure_parent(tables_target)
                shutil.copy2(tables_real, tables_target)
            except Exception:
                pass
        if images_extract_enabled and images_real.exists():
            try:
                ensure_parent(images_target)
                shutil.copy2(images_real, images_target)
            except Exception:
                pass
        if table_extract_enabled and table_images_dir.exists():
            try:
                shutil.copytree(table_images_dir, table_images_target_dir, dirs_exist_ok=True)
            except Exception:
                pass
        if images_extract_enabled and doc_images_dir.exists():
            try:
                shutil.copytree(doc_images_dir, doc_images_target_dir, dirs_exist_ok=True)
            except Exception:
                pass
        if summary_real.exists():
            summary_export = dict(summary_payload)
            if table_extract_enabled:
                summary_export["table_elements"] = {
                    "enabled": True,
                    "path": str(tables_target) if tables_target.exists() else "",
                    "images_dir": str(table_images_target_dir) if table_images_target_dir.exists() else "",
                }
            if images_extract_enabled:
                summary_export["docx_images"] = {
                    "enabled": True,
                    "path": str(images_target) if images_target.exists() else "",
                    "images_dir": str(doc_images_target_dir) if doc_images_target_dir.exists() else "",
                }
            try:
                ensure_parent(summary_target)
                summary_target.write_text(
                    json.dumps(summary_export, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
            except Exception:
                try:
                    ensure_parent(summary_target)
                    shutil.copy2(summary_real, summary_target)
                except Exception:
                    pass
    return {"messages": []}
