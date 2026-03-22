from __future__ import annotations

from dataclasses import dataclass, field, replace
import datetime as dt
import json
import os
from pathlib import Path
import queue
import re
import shutil
import threading
import traceback
from typing import Iterable
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait

from app.agents.supervisor import build_agent
from app.settings import AppSettings
from app.tools.internet_search import run_internet_search
from app.tools.path_utils import ensure_parent, resolve_path, to_virtual_path
from app.tools.win32_utils import (
    com_retry,
    dispatch_word_application,
    get_win32_constants,
    install_ole_message_filter,
    is_com_call_rejected,
    is_com_not_initialized,
    restore_ole_message_filter,
)
from app.formatting.profiles import apply_format_profile, resolve_profile


DEFAULT_FRAMEWORK = """1. [选题说明] 选题所研究的具体问题、研究视角和核心概念（300字以内）。
2. [选题依据] 国内外相关研究的学术史梳理及研究进展（略写）；相对于已有研究特别是国家社科基金同类项目的独到学术价值和应用价值。
3. [研究内容] 课题研究的主要目标、重点难点、整体框架、研究计划及其可行性等。（框架思路要列出提纲或目录）
4. [创新之处] 在学术观点、研究方法等方面的特色和创新。
5. [预期成果] 成果形式、宣传转化及预期学术价值和社会效益等。（略写）
6. [研究基础] 申请人前期相关代表性研究成果、核心观点等。（略写）
7. [参考文献] 开展本课题研究的主要中外参考文献。（略写）"""

OUTLINE_JSON_EXAMPLE = """[
  {"level": 1, "title": "选题说明"},
  {"level": 2, "title": "研究问题与视角"},
  {"level": 3, "title": "核心概念界定"},
  {"level": 1, "title": "研究内容"},
  {"level": 2, "title": "研究框架与总体思路"},
  {"level": 3, "title": "建模方法与技术路线"}
]"""


@dataclass
class ReportSection:
    title: str
    description: str = ""
    weight: float = 1.0
    min_chars: int | None = None
    max_chars: int | None = None
    target_chars: int | None = None
    brief: bool = False
    is_reference: bool = False


@dataclass
class OutlineItem:
    level: int
    title: str


@dataclass
class OutlineNode:
    title: str
    level: int
    children: list["OutlineNode"] = field(default_factory=list)
    target_chars: int | None = None
    weight: float = 1.0
    brief: bool = False
    is_reference: bool = False


@dataclass
class SourceItem:
    title: str
    url: str
    snippet: str
    score: float | None = None
    query: str | None = None


def _infer_weight(title: str, desc: str) -> float:
    if "研究内容" in title:
        return 3.0
    if "选题依据" in title:
        return 1.8
    if "参考文献" in title:
        return 1.2
    if "选题说明" in title:
        return 0.6
    if "创新" in title:
        return 1.0
    if "预期" in title:
        return 0.9
    if "研究基础" in title:
        return 0.9
    if "略写" in desc:
        return 0.7
    return 1.0


def _extract_json_list(text: str) -> list[dict] | None:
    if not text:
        return None
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1 or end <= start:
        return None
    candidate = text[start : end + 1]
    try:
        data = json.loads(candidate)
    except Exception:
        return None
    if isinstance(data, list):
        return data
    return None


def _strip_prefix_markers(text: str) -> str:
    value = (text or "").strip()
    value = re.sub(r"^\s*(?:\d+|[一二三四五六七八九十]+)[.\u3001)]\s*", "", value)
    value = re.sub(r"^\s*[\-\*•]\s*", "", value)
    return value.strip()


def _extract_title_from_line(line: str) -> str:
    cleaned = _strip_prefix_markers(line)
    cleaned = re.sub(r"^[\[\【]\s*(.+?)\s*[\]\】]\s*", r"\1 ", cleaned).strip()
    if not cleaned:
        return ""
    for sep in ("：", ":", "。"):
        if sep in cleaned:
            left = cleaned.split(sep, 1)[0].strip()
            if left:
                return left
    parts = re.split(r"\s+", cleaned, maxsplit=1)
    return parts[0].strip() if parts else cleaned


def _outline_from_framework(text: str | None) -> list[OutlineItem]:
    raw = (text or "").strip()
    if not raw:
        raw = DEFAULT_FRAMEWORK
    lines = [line.strip() for line in re.split(r"[\n\r;；]+", raw) if line.strip()]
    items: list[OutlineItem] = []
    for line in lines:
        title = _extract_title_from_line(line)
        if not title:
            continue
        items.append(OutlineItem(level=1, title=title))
    if not items:
        items = [OutlineItem(level=1, title="选题说明"), OutlineItem(level=1, title="研究内容")]
    return items


_OUTLINE_LINE_PATTERN = re.compile(r"^(?:\\[outline\\]\\s*)?L?(?P<level>[1-3])\\s+(?P<title>.+)$", re.I)


def _outline_from_explicit_text(text: str | None) -> list[OutlineItem] | None:
    raw = (text or "").strip()
    if not raw:
        return None
    lines = [line.strip() for line in raw.splitlines()]
    if not lines:
        return None

    def _parse_lines(source: list[str]) -> list[OutlineItem]:
        items: list[OutlineItem] = []
        for line in source:
            if not line:
                continue
            if "层级" in line and "标题" in line:
                continue
            if line.startswith(("-", "—", "|")):
                continue
            match = _OUTLINE_LINE_PATTERN.match(line)
            level = None
            title = ""
            if match:
                level = int(match.group("level"))
                title = match.group("title").strip()
            else:
                parts = re.split(r"\\s+", line, maxsplit=1)
                if len(parts) == 2 and parts[0].isdigit():
                    level = int(parts[0])
                    title = parts[1].strip()
            if level is None or not title:
                continue
            if level < 1 or level > 3:
                continue
            normalized = _normalize_outline_title(title)
            items.append(OutlineItem(level=level, title=normalized or title))
        return items

    header_idx = None
    for idx, line in enumerate(lines):
        if "层级" in line and "标题" in line:
            header_idx = idx + 1
            break
    if header_idx is not None:
        items = _parse_lines(lines[header_idx:])
        if items:
            return items

    items = _parse_lines(lines)
    if len(items) >= 3:
        return items
    return None


def _build_outline_system_prompt() -> str:
    return (
        "你是课题报告的大纲规划助手。任务是根据选题与用户框架生成三级标题大纲。"
        "必须输出 JSON 数组，每个元素包含 level(1-3) 与 title。"
        "输出只允许 JSON，不要解释，不要 Markdown。"
        "要求："
        "1) 结构必须到三级标题；每个一级标题至少包含 1-2 个二级与三级标题。"
        "2) 保留并映射用户给出的框架意图；框架可能是自由文本或带编号。"
        "3) title 只写短标题（不超过10字），不要把说明性文字写进标题。"
        "3) 必须体现建模方法（标题中含“建模”“模型”“方法”等关键词），建议放在“研究内容”之下。"
        "4) 最后一级标题应为“参考文献”。"
        "示例："
        f"{OUTLINE_JSON_EXAMPLE}"
    )


def _build_outline_prompt(topic: str, framework_text: str) -> str:
    return (
        f"选题：{topic}\n"
        "用户框架（可能是自由文本）：\n"
        f"{framework_text}\n"
        "请输出符合要求的三级标题 JSON 数组。"
    )


def _parse_outline_items(payload: list[dict]) -> list[OutlineItem]:
    items: list[OutlineItem] = []
    for item in payload:
        try:
            level = int(item.get("level", 1))
        except Exception:
            level = 1
        title = _normalize_outline_title((item.get("title") or "").strip())
        if not title:
            continue
        level = max(1, min(3, level))
        items.append(OutlineItem(level=level, title=title))
    return items


def _canonical_l1_title(title: str) -> str:
    if not title:
        return title
    mappings = {
        "选题说明": ("选题说明",),
        "选题依据": ("选题依据", "相对于已有研究", "学术价值", "应用价值", "研究进展", "国内外研究"),
        "研究内容": ("研究内容", "研究目标", "重点难点", "整体框架", "研究计划", "可行性"),
        "创新之处": ("创新", "研究方法"),
        "预期成果": ("预期成果", "成果形式", "宣传转化", "社会效益", "学术价值"),
        "研究基础": ("研究基础", "前期研究", "代表性研究成果", "核心观点"),
        "参考文献": ("参考文献", "参考资料", "主要中外参考文献"),
    }
    for canonical, keywords in mappings.items():
        if any(keyword in title for keyword in keywords):
            return canonical
    return title


def _normalize_outline_title(title: str) -> str:
    value = (title or "").strip()
    if not value:
        return ""
    value = re.sub(r"^[\d一二三四五六七八九十]+[.\u3001)]\s*", "", value)
    value = re.sub(r"[\[\【]\s*", "", value)
    value = re.sub(r"\s*[\]\】]", "", value)
    value = re.sub(r"（.*?）", "", value).strip()
    for sep in ("：", ":", "，", ",", "。"):
        if sep in value:
            left = value.split(sep, 1)[0].strip()
            if left:
                value = left
                break
    parts = value.split()
    if parts:
        value = parts[0].strip()
    value = _canonical_l1_title(value)
    return value.strip()


def _ensure_modeling_methods(items: list[OutlineItem]) -> list[OutlineItem]:
    if any(re.search(r"建模|模型|方法", item.title) for item in items):
        return items
    updated: list[OutlineItem] = []
    inserted = False
    for item in items:
        updated.append(item)
        if item.level == 1 and ("研究内容" in item.title or "研究方法" in item.title):
            updated.append(OutlineItem(level=2, title="建模方法与技术路线"))
            updated.append(OutlineItem(level=3, title="模型构建与验证"))
            inserted = True
            break
    if not inserted:
        updated.append(OutlineItem(level=1, title="研究内容"))
        updated.append(OutlineItem(level=2, title="建模方法与技术路线"))
        updated.append(OutlineItem(level=3, title="模型构建与验证"))
    return updated


def _build_outline_with_llm(
    settings: AppSettings,
    topic: str,
    framework_text: str,
) -> list[OutlineItem]:
    explicit_json = _extract_json_list(framework_text)
    if explicit_json:
        items = _parse_outline_items(explicit_json)
        if items:
            return items
    explicit_outline = _outline_from_explicit_text(framework_text)
    if explicit_outline:
        return explicit_outline
    system_prompt = _build_outline_system_prompt()
    agent = build_agent(settings, tools=[], system_prompt=system_prompt)
    payload = {"messages": [{"role": "user", "content": _build_outline_prompt(topic, framework_text)}]}
    result = agent.invoke(payload, config={"configurable": {"thread_id": f"outline_{dt.datetime.now().timestamp()}"}})
    content = result["messages"][-1].content if result.get("messages") else ""
    parsed = _extract_json_list(content)
    if not parsed:
        return _ensure_modeling_methods(_outline_from_framework(framework_text))
    items = _parse_outline_items(parsed)
    if not items:
        return _ensure_modeling_methods(_outline_from_framework(framework_text))
    return _ensure_modeling_methods(items)


def _items_to_tree(items: list[OutlineItem]) -> OutlineNode:
    root = OutlineNode(title="ROOT", level=0)
    stack = [root]
    l1_nodes: dict[str, OutlineNode] = {}
    for raw in items:
        level = max(1, min(3, raw.level))
        title = raw.title
        if level == 1:
            title = _canonical_l1_title(title)
            if title in l1_nodes:
                node = l1_nodes[title]
            else:
                node = OutlineNode(title=title, level=1)
                l1_nodes[title] = node
                root.children.append(node)
            stack = [root, node]
            continue
        node = OutlineNode(title=title, level=level)
        while stack and stack[-1].level >= level:
            stack.pop()
        parent = stack[-1] if stack else root
        parent.children.append(node)
        stack.append(node)
    return root


def _outline_items_from_tree(root: OutlineNode) -> list[OutlineItem]:
    items: list[OutlineItem] = []

    def _walk(node: OutlineNode) -> None:
        for child in node.children:
            items.append(OutlineItem(level=child.level, title=child.title))
            _walk(child)

    _walk(root)
    return items


def _ensure_three_levels(root: OutlineNode) -> None:
    for level1 in root.children:
        if "参考文献" in level1.title or "参考资料" in level1.title:
            continue
        if not level1.children:
            level1.children.append(OutlineNode(title="研究框架与总体思路", level=2))
        for level2 in level1.children:
            if not level2.children:
                level2.children.append(OutlineNode(title=f"{level2.title}的关键要点", level=3))


def _mark_reference_sections(root: OutlineNode) -> None:
    for node in root.children:
        if "参考文献" in node.title or "参考资料" in node.title:
            node.is_reference = True


def _apply_section_hints(root: OutlineNode) -> None:
    for node in root.children:
        node.weight = _infer_weight(node.title, "")
        if "选题说明" in node.title:
            node.target_chars = 300
            node.brief = False
        if "选题依据" in node.title:
            node.weight = 1.8
        if "预期成果" in node.title or "研究基础" in node.title:
            node.brief = True
        if "参考文献" in node.title:
            node.is_reference = True


def _allocate_outline_targets(root: OutlineNode, total_chars: int) -> None:
    _apply_section_hints(root)
    fixed_total = 0
    for node in root.children:
        if node.target_chars is not None:
            fixed_total += node.target_chars

    flexible = [n for n in root.children if n.target_chars is None and not n.is_reference]
    remaining = max(total_chars - fixed_total, int(total_chars * 0.6))
    weight_sum = sum(n.weight for n in flexible) or 1.0
    for node in flexible:
        target = int(remaining * (node.weight / weight_sum))
        if node.brief:
            target = min(target, 1200)
        node.target_chars = max(400, target)

def _parse_framework(text: str | None) -> list[ReportSection]:
    raw = (text or "").strip()
    if not raw:
        raw = DEFAULT_FRAMEWORK
    lines = [line.strip() for line in raw.splitlines() if line.strip()]
    sections: list[ReportSection] = []
    pattern = re.compile(
        r"^\s*(?:\d+|[一二三四五六七八九十]+)[.\u3001)]*\s*"
        r"(?:\[(?P<title>[^]]+)\]|\u3010(?P<title2>[^】]+)\u3011)?\s*(?P<desc>.*)$"
    )
    for line in lines:
        match = pattern.match(line)
        title = ""
        desc = line
        if match:
            title = match.group("title") or match.group("title2") or ""
            desc = match.group("desc") or ""
        if not title:
            parts = re.split(r"\s+", line, maxsplit=1)
            title = parts[0].strip()
            desc = parts[1].strip() if len(parts) > 1 else ""
        section = ReportSection(title=title, description=desc)
        if "参考文献" in title:
            section.is_reference = True
        if "300字" in desc:
            section.max_chars = 300
            section.min_chars = 220
        if "略写" in desc:
            section.brief = True
        section.weight = _infer_weight(section.title, section.description)
        sections.append(section)
    if not sections:
        return _parse_framework(DEFAULT_FRAMEWORK)
    return sections


def _allocate_targets(sections: list[ReportSection], total_chars: int) -> list[ReportSection]:
    fixed_total = 0
    for section in sections:
        if section.max_chars is not None:
            target = min(section.max_chars, max(section.min_chars or 0, section.max_chars - 20))
            section.target_chars = target
            fixed_total += target

    flexible_sections = [s for s in sections if s.target_chars is None and not s.is_reference]
    remaining = max(total_chars - fixed_total, int(total_chars * 0.6))
    weight_sum = sum(s.weight for s in flexible_sections) or 1.0
    for section in flexible_sections:
        target = int(remaining * (section.weight / weight_sum))
        if section.brief:
            target = min(target, 1200)
        if section.min_chars is not None:
            target = max(target, section.min_chars)
        section.target_chars = max(200, target)
    return sections


def _build_search_queries(topic: str, sections: Iterable) -> list[str]:
    queries = [
        f"{topic} 研究现状",
        f"{topic} 文献综述",
        f"{topic} 理论框架",
        f"{topic} 研究方法",
        f"{topic} 政策 文件",
        f"{topic} 应用 价值",
    ]
    if any("选题依据" in s.title for s in sections):
        queries.append(f"{topic} 国内外研究进展")
    if any("创新" in s.title for s in sections):
        queries.append(f"{topic} 最新研究 争议 与趋势")
    return list(dict.fromkeys([q for q in queries if q]))


def _run_search(
    queries: list[str],
    *,
    max_results: int = 5,
    logger=None,
) -> tuple[list[SourceItem], list[str]]:
    sources: list[SourceItem] = []
    errors: list[str] = []
    seen: set[str] = set()
    for query in queries:
        if logger:
            logger(f"[search] query: {query}")
        try:
            resp = run_internet_search(
                query=query,
                max_results=max_results,
                topic="general",
                include_raw_content=False,
            )
        except Exception as exc:  # noqa: BLE001
            errors.append(str(exc))
            if logger:
                logger(f"[search] error: {exc}")
            continue
        results = resp.get("results", []) if isinstance(resp, dict) else []
        if logger:
            logger(f"[search] results: {len(results)}")
        for item in results:
            title = (item.get("title") or "").strip()
            url = (item.get("url") or item.get("link") or "").strip()
            snippet = (item.get("content") or item.get("snippet") or "").strip()
            if not title and not url:
                continue
            key = url or title
            if key in seen:
                continue
            seen.add(key)
            sources.append(
                SourceItem(
                    title=title or "无标题",
                    url=url,
                    snippet=snippet,
                    score=item.get("score"),
                    query=query,
                )
            )
    return sources, errors


def _format_sources(sources: list[SourceItem], *, max_items: int = 20, max_chars: int = 6000) -> str:
    lines: list[str] = []
    total = 0
    for idx, source in enumerate(sources[:max_items], start=1):
        snippet = source.snippet.replace("\n", " ").strip()
        if len(snippet) > 220:
            snippet = snippet[:219] + "…"
        block = f"[{idx}] {source.title}\nURL: {source.url}\n摘要: {snippet}"
        if total + len(block) > max_chars:
            break
        lines.append(block)
        total += len(block)
    return "\n\n".join(lines)


def _select_sources_for_node(sources: list[SourceItem], node: OutlineNode, *, max_items: int = 8) -> list[SourceItem]:
    if not sources:
        return []
    keywords: list[str] = [node.title]
    for level2 in node.children:
        keywords.append(level2.title)
        for level3 in level2.children:
            keywords.append(level3.title)
    keywords = [k.strip() for k in keywords if isinstance(k, str) and k.strip()]
    keywords = list(dict.fromkeys(keywords))[:12]
    if not keywords:
        return sources[:max_items]

    scored: list[tuple[int, int, SourceItem]] = []
    for idx, source in enumerate(sources):
        haystack = f"{source.title}\n{source.snippet}\n{source.query or ''}"
        score = 0
        for kw in keywords:
            if kw and kw in haystack:
                score += 2
        if score <= 0:
            continue
        scored.append((score, -idx, source))
    scored.sort(reverse=True)
    selected = [item for _score, _idx, item in scored[:max_items]]
    return selected or sources[:max_items]


def _build_report_system_prompt() -> str:
    return (
        "你是课题报告写作助手，负责根据给定选题、框架与检索资料撰写学术型课题报告。"
        "写作必须遵守框架顺序，保持严谨、客观、逻辑清晰的学术表达。"
        "不得编造具体数据、政策条文或不存在的研究结论。"
        "如资料不足，用概括性表述或提出假设性研究路径，不要虚构来源。"
    )


def _section_outline_block(node: OutlineNode) -> list[str]:
    lines: list[str] = []
    for level2 in node.children:
        lines.append(f"[H2] {level2.title}")
        for level3 in level2.children:
            lines.append(f"[H3] {level3.title}")
    return lines


def _count_leaf_nodes(node: OutlineNode) -> int:
    count = 0
    for level2 in node.children:
        if level2.children:
            count += len(level2.children)
        else:
            count += 1
    return max(1, count)


def _leaf_paths(node: OutlineNode) -> list[tuple[str, str | None]]:
    paths: list[tuple[str, str | None]] = []
    for level2 in node.children:
        if level2.children:
            for level3 in level2.children:
                paths.append((level2.title, level3.title))
        else:
            paths.append((level2.title, None))
    return paths


def _build_section_prompt(
    topic: str,
    node: OutlineNode,
    outline_titles: list[str],
    sources_text: str,
) -> str:
    target = node.target_chars or 1200
    leaf_target = int(target / _count_leaf_nodes(node))
    hint_lines = [
        f"选题：{topic}",
        f"本节标题：{node.title}",
        f"本节字数目标：约{target}字",
        "全文一级结构：" + "；".join(outline_titles),
        "请严格按照以下二级/三级标题输出，并在每个标题下撰写内容。",
        "输出格式要求：必须保留 [H2]/[H3] 标记；不要输出其他标题或编号。",
        f"每个三级标题建议字数：约{leaf_target}字（可上下浮动），不得只列标题。",
    ]
    if "研究内容" in node.title:
        hint_lines.append("本节需体现建模方法，说明模型构建、验证与技术路线。")
    if node.brief:
        hint_lines.append("本节为略写，保持简明但完整。")
    if sources_text:
        hint_lines.append("可参考资料摘要如下（仅限于这些信息，不要虚构来源）：")
        hint_lines.append(sources_text)
    hint_lines.append("二级/三级标题如下：")
    hint_lines.extend(_section_outline_block(node))
    return "\n".join(hint_lines)


def _clean_section_text(text: str, title: str) -> str:
    content = (text or "").strip()
    if not content:
        return ""
    # Remove leading headings if the model echoed them.
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return content
    first = lines[0]
    if title in first and not first.strip().startswith(("[H2]", "[H3]")):
        lines = lines[1:]
    content = "\n".join(lines).strip()
    return content


def _parse_section_blocks(text: str) -> list[tuple[int, str, str]]:
    lines = [line.rstrip() for line in (text or "").splitlines()]
    blocks: list[tuple[int, str, str]] = []
    current_level: int | None = None
    current_title: str | None = None
    buffer: list[str] = []

    def _match_heading(line: str) -> tuple[int, str] | None:
        stripped = line.strip()
        if stripped.startswith("[H2]"):
            return 2, stripped.replace("[H2]", "", 1).strip()
        if stripped.startswith("[H3]"):
            return 3, stripped.replace("[H3]", "", 1).strip()
        if stripped.startswith("【H2】"):
            return 2, stripped.replace("【H2】", "", 1).strip()
        if stripped.startswith("【H3】"):
            return 3, stripped.replace("【H3】", "", 1).strip()
        if re.match(r"^H2[:：\\s]", stripped, flags=re.I):
            return 2, re.sub(r"^H2[:：\\s]+", "", stripped, flags=re.I).strip()
        if re.match(r"^H3[:：\\s]", stripped, flags=re.I):
            return 3, re.sub(r"^H3[:：\\s]+", "", stripped, flags=re.I).strip()
        return None

    def _flush():
        if current_level and current_title is not None:
            content = "\n".join([line for line in buffer if line.strip()]).strip()
            blocks.append((current_level, current_title, content))

    for line in lines:
        match = _match_heading(line)
        if match:
            _flush()
            current_level, current_title = match
            buffer = []
            continue
        buffer.append(line)
    _flush()
    return [block for block in blocks if block[1]]


ContentKey = tuple[str, str, str | None]


def _h2_key(title: str) -> ContentKey:
    return ("h2", (title or "").strip(), None)


def _h3_key(h2_title: str, h3_title: str) -> ContentKey:
    return ("h3", (h2_title or "").strip(), (h3_title or "").strip())


def _format_leaf_key(key: ContentKey) -> str:
    kind, h2_title, h3_title = key
    if kind == "h3" and h3_title:
        return f"{h2_title}/{h3_title}"
    return h2_title or (h3_title or "（未命名）")


def _append_content(content_map: dict[ContentKey, str], key: ContentKey, content: str) -> None:
    content = (content or "").strip()
    if not content:
        return
    existing = content_map.get(key)
    if existing:
        content_map[key] = f"{existing}\n{content}".strip()
    else:
        content_map[key] = content


def _blocks_to_content_map(blocks: list[tuple[int, str, str]]) -> dict[ContentKey, str]:
    """
    Convert parsed [H2]/[H3] blocks into a structured map keyed by (kind, h2, h3).

    - ("h2", H2, None): content directly under H2 (as intro or leaf body).
    - ("h3", H2, H3): content under H3 (requires an active H2 context).
    """
    content_map: dict[ContentKey, str] = {}
    current_h2: str | None = None
    for level, title, content in blocks:
        if level == 2:
            current_h2 = title.strip()
            _append_content(content_map, _h2_key(current_h2), content)
            continue
        if level == 3:
            if not current_h2:
                # Orphan H3 without a preceding H2. Keep it for debugging, but it
                # won't match expected outline and will be refilled later.
                current_h2 = ""
            _append_content(content_map, _h3_key(current_h2, title), content)
            continue
    return content_map


def _split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n{2,}", text or "") if p.strip()]


def _invoke_agent(agent, payload: dict, thread_id: str, timeout: int | None = None):
    # Do not run agent.invoke in a different thread to enforce timeouts.
    # DeepAgents/LangGraph backends/checkpointers can be thread-sensitive and may
    # fail when invoked across threads (often as empty-message AssertionError).
    #
    # Keep the `timeout` argument for compatibility, but invoke in the current
    # thread for stability.
    return agent.invoke(payload, config={"configurable": {"thread_id": thread_id}})


def _expected_leaf_titles(node: OutlineNode) -> list[str]:
    titles: list[str] = []
    for h2, h3 in _leaf_paths(node):
        titles.append(_format_leaf_key(_h3_key(h2, h3) if h3 else _h2_key(h2)))
    return titles


def _expected_leaf_keys(node: OutlineNode) -> list[ContentKey]:
    keys: list[ContentKey] = []
    for h2, h3 in _leaf_paths(node):
        keys.append(_h3_key(h2, h3) if h3 else _h2_key(h2))
    return keys


def _validate_section_output(node: OutlineNode, content_map: dict[ContentKey, str]) -> list[ContentKey]:
    expected = _expected_leaf_keys(node)
    if not expected:
        return []
    leaf_target = int((node.target_chars or 1200) / max(1, len(expected)))
    min_chars = max(120, int(leaf_target * 0.6))
    missing = [key for key in expected if len((content_map.get(key) or "").strip()) < min_chars]
    return missing


def _build_missing_leaf_prompt(
    topic: str,
    node: OutlineNode,
    h2_title: str,
    h3_title: str | None,
    outline_titles: list[str],
    sources_text: str,
    target_chars: int,
) -> str:
    lines = [
        f"选题：{topic}",
        f"所在一级标题：{node.title}",
        f"二级标题：{h2_title}",
    ]
    if h3_title:
        lines.append(f"三级标题：{h3_title}")
    lines.append("全文一级结构：" + "；".join(outline_titles))
    lines.append(f"本段目标字数：约{target_chars}字")
    lines.append("要求：仅输出本段正文内容，不要输出任何标题或编号，不要虚构数据。")
    if sources_text:
        lines.append("可参考资料摘要如下（仅限于这些信息，不要虚构来源）：")
        lines.append(sources_text)
    return "\n".join(lines)


def _fallback_leaf_text(topic: str, node: OutlineNode, h2_title: str, h3_title: str | None) -> str:
    title = h3_title or h2_title
    return (
        f"围绕“{topic}”的{node.title}部分，本节以“{title}”为核心展开。"
        "主要说明该主题的基本内涵与分析范围，梳理需要回答的关键问题，"
        "并提出可操作的研究思路与分析路径，为后续研究内容与方法设计提供支撑。"
        "整体表述保持概括性与可执行性，不引入具体数据或未经验证的结论。"
    )


def _pad_text(text: str, min_chars: int, topic: str, section_title: str, leaf_title: str) -> str:
    if len(text) >= min_chars:
        return text
    padding_sentences = [
        f"在“{leaf_title}”这一部分，将从理论阐释与现实场景两条线索展开，强调问题界定与研究边界的明确性。",
        "研究将通过规范性分析、比较分析与典型情境归纳等方式，形成可验证的分析框架。",
        "同时强调逻辑链条的完整性，确保问题提出、分析路径与结论输出之间具有可追溯性。",
        f"上述内容与“{section_title}”整体目标保持一致，为后续章节提供必要的概念支撑与方法铺垫。",
    ]
    buffer = text
    idx = 0
    while len(buffer) < min_chars:
        buffer = f"{buffer} {padding_sentences[idx % len(padding_sentences)]}".strip()
        idx += 1
    return buffer


def _leaf_length_targets(node: OutlineNode) -> tuple[int, int]:
    expected = _expected_leaf_keys(node)
    leaf_target = int((node.target_chars or 1200) / max(1, len(expected)))
    min_chars = max(120, int(leaf_target * 0.6))
    return leaf_target, min_chars


def _ensure_h2_intros(topic: str, node: OutlineNode, content_map: dict[ContentKey, str]) -> None:
    """
    Ensure every H2 that has H3 children has a short intro paragraph.

    Prevents "H2 immediately followed by H3" with no bridging text.
    """
    _leaf_target, min_leaf_chars = _leaf_length_targets(node)
    min_intro = min(220, max(80, int(min_leaf_chars * 0.6)))
    for level2 in node.children:
        if not level2.children:
            continue
        key = _h2_key(level2.title)
        existing = (content_map.get(key) or "").strip()
        if len(existing) >= min_intro:
            continue
        text = _fallback_leaf_text(topic, node, level2.title, None)
        content_map[key] = _pad_text(text, min_intro, topic, node.title, level2.title)


def _fallback_content_map_for_node(topic: str, node: OutlineNode) -> dict[ContentKey, str]:
    _leaf_target, min_chars = _leaf_length_targets(node)
    fallback: dict[ContentKey, str] = {}
    for key in _expected_leaf_keys(node):
        kind, h2_title, h3_title = key
        if kind == "h3" and h3_title:
            text = _fallback_leaf_text(topic, node, h2_title, h3_title)
        else:
            text = _fallback_leaf_text(topic, node, h2_title, None)
        fallback[key] = _pad_text(text, min_chars, topic, node.title, _format_leaf_key(key))
    _ensure_h2_intros(topic, node, fallback)
    return fallback


def _compute_report_stats(outline_root: OutlineNode, contents: dict[str, dict]) -> dict:
    leaf_total = 0
    leaf_missing = 0
    leaf_short = 0
    char_total = 0
    by_section: list[dict] = []

    for node in outline_root.children:
        if node.is_reference:
            continue
        section_map = contents.get(node.title) if isinstance(contents.get(node.title), dict) else {}
        expected = _expected_leaf_keys(node)
        leaf_total += len(expected)
        leaf_target, min_chars = _leaf_length_targets(node)
        section_missing = 0
        section_short = 0
        for key in expected:
            text = (section_map.get(key) or "").strip() if isinstance(section_map, dict) else ""
            if not text:
                leaf_missing += 1
                section_missing += 1
                continue
            char_total += len(text)
            if len(text) < min_chars:
                leaf_short += 1
                section_short += 1
        by_section.append(
            {
                "title": node.title,
                "leaf_target_chars": leaf_target,
                "leaf_min_chars": min_chars,
                "leaf_count": len(expected),
                "missing": section_missing,
                "short": section_short,
            }
        )

    return {
        "leaf_total": leaf_total,
        "leaf_missing": leaf_missing,
        "leaf_short": leaf_short,
        "leaf_coverage": 0.0 if leaf_total == 0 else round((leaf_total - leaf_missing) / leaf_total, 4),
        "leaf_char_total": char_total,
        "sections": by_section,
    }


def _fill_missing_sections(
    agent,
    topic: str,
    node: OutlineNode,
    outline_titles: list[str],
    sources_text: str,
    content_map: dict[ContentKey, str],
    missing_titles: list[ContentKey],
    section_timeout: int | None,
    logger=None,
) -> dict[ContentKey, str]:
    leaf_paths = _leaf_paths(node)
    leaf_target = int((node.target_chars or 1200) / max(1, len(leaf_paths)))
    min_chars = max(120, int(leaf_target * 0.6))
    for h2_title, h3_title in leaf_paths:
        key = _h3_key(h2_title, h3_title) if h3_title else _h2_key(h2_title)
        if key not in missing_titles:
            continue
        prompt = _build_missing_leaf_prompt(
            topic,
            node,
            h2_title,
            h3_title,
            outline_titles,
            sources_text,
            leaf_target,
        )
        payload = {"messages": [{"role": "user", "content": prompt}]}
        thread_id = f"report_fill_{node.title}_{_format_leaf_key(key)}_{dt.datetime.now().timestamp()}"
        result = _invoke_agent(agent, payload, thread_id, section_timeout)
        if result is None:
            if logger:
                logger(f"[fill] 子标题超时：{_format_leaf_key(key)}")
            text = _fallback_leaf_text(topic, node, h2_title, h3_title)
            content_map[key] = _pad_text(text, min_chars, topic, node.title, _format_leaf_key(key))
            if logger:
                logger(f"[fill] 子标题使用兜底文本：{_format_leaf_key(key)}")
            continue
        content = result["messages"][-1].content if result.get("messages") else ""
        text = (content or "").strip()
        if not text:
            text = _fallback_leaf_text(topic, node, h2_title, h3_title)
            content_map[key] = _pad_text(text, min_chars, topic, node.title, _format_leaf_key(key))
            if logger:
                logger(f"[fill] 子标题使用兜底文本：{_format_leaf_key(key)}")
            continue
        content_map[key] = _pad_text(text, min_chars, topic, node.title, _format_leaf_key(key))
        if logger:
            logger(f"[fill] 子标题补全：{_format_leaf_key(key)}（约{len(text)}字）")
    return content_map


def _build_retry_prompt(
    topic: str,
    node: OutlineNode,
    outline_titles: list[str],
    sources_text: str,
    missing_titles: list[ContentKey],
) -> str:
    base = _build_section_prompt(topic, node, outline_titles, sources_text)
    missing = "；".join([_format_leaf_key(key) for key in missing_titles])
    retry_hint = (
        "\n注意：上一次输出存在内容过短/缺失，请重新完整输出本节内容。"
        f"\n必须覆盖并充实以下标题（每个不少于120字）：{missing}"
    )
    return base + retry_hint


def _ensure_length(agent, node: OutlineNode, text: str) -> str:
    if not text:
        return text
    target = node.target_chars or 1200
    length = len(text)
    hard_limit = 300 if "选题说明" in node.title else None
    if hard_limit and length > hard_limit:
        prompt = (
            f"请将以下内容压缩到不超过{hard_limit}字，保留 [H2]/[H3] 标记与标题，"
            "保持核心信息与逻辑，不新增事实：\n"
            f"{text}"
        )
    elif node.target_chars and length > int(target * 1.4):
        prompt = (
            f"请将以下内容压缩到约{target}字，保留 [H2]/[H3] 标记与标题，"
            "保持核心信息与逻辑，不新增事实：\n"
            f"{text}"
        )
    elif length < int(target * 0.6):
        prompt = (
            f"请在保持结构不变、保留 [H2]/[H3] 标记的前提下扩写至约{target}字，"
            "补充必要背景、过渡或方法说明，但不要新增虚构数据：\n"
            f"{text}"
        )
    else:
        return text
    payload = {"messages": [{"role": "user", "content": prompt}]}
    result = agent.invoke(payload, config={"configurable": {"thread_id": f"report_refine_{dt.datetime.now().timestamp()}"}})
    content = result["messages"][-1].content if result.get("messages") else ""
    return (content or text).strip()


def _write_outline_table_docx(doc, outline_items: list[OutlineItem]) -> None:
    doc.add_heading("报告大纲表", level=1)
    table = doc.add_table(rows=1, cols=2)
    header = table.rows[0].cells
    header[0].text = "层级"
    header[1].text = "标题"
    for item in outline_items:
        row = table.add_row().cells
        row[0].text = str(item.level)
        row[1].text = item.title


def _resolve_style_name(doc, candidates: list[str]) -> str | None:
    for name in candidates:
        try:
            doc.Styles(name)
        except Exception:
            continue
        return name
    return None


def _win32_add_paragraph(doc, text: str, style_name: str | None) -> None:
    rng = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
    para = doc.Paragraphs.Add(rng)
    para.Range.Text = text
    if style_name:
        try:
            para.Range.Style = style_name
        except Exception:
            pass


def _write_report_docx_win32(
    path: Path,
    topic: str,
    outline_root: OutlineNode,
    outline_items: list[OutlineItem],
    contents: dict[str, dict],
) -> None:
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"win32com 初始化失败：{exc}") from exc

    pythoncom.CoInitialize()
    old_filter, message_filter = install_ole_message_filter()
    word = None
    doc = None
    try:
        word = com_retry(dispatch_word_application)
        word.Visible = False
        word.DisplayAlerts = 0
        doc = com_retry(lambda: word.Documents.Add())
        try:
            constants = get_win32_constants(win32)
        except Exception:
            constants = None

        title_style = _resolve_style_name(doc, ["Title", "标题"])
        h1_style = _resolve_style_name(doc, ["Heading 1", "标题 1"])
        h2_style = _resolve_style_name(doc, ["Heading 2", "标题 2"])
        h3_style = _resolve_style_name(doc, ["Heading 3", "标题 3"])
        normal_style = _resolve_style_name(doc, ["Normal", "正文"])

        first_para = doc.Paragraphs(1)
        first_para.Range.Text = f"{topic} 课题报告"
        if title_style:
            try:
                first_para.Range.Style = title_style
            except Exception:
                pass

        _win32_add_paragraph(doc, "报告大纲表", h1_style)
        rng = doc.Range(doc.Content.End - 1, doc.Content.End - 1)
        table = doc.Tables.Add(rng, len(outline_items) + 1, 2)
        try:
            table.Style = _resolve_style_name(doc, ["Table Grid", "表格网格"]) or table.Style
        except Exception:
            pass
        try:
            if constants:
                table.AutoFitBehavior(constants.wdAutoFitWindow)
        except Exception:
            pass
        table.Cell(1, 1).Range.Text = "层级"
        table.Cell(1, 2).Range.Text = "标题"
        try:
            table.Rows(1).Range.Bold = True
        except Exception:
            pass
        for idx, item in enumerate(outline_items, start=2):
            table.Cell(idx, 1).Range.Text = str(item.level)
            table.Cell(idx, 2).Range.Text = item.title

        section_index = 1
        for node in outline_root.children:
            _win32_add_paragraph(doc, f"{section_index}. {node.title}", h1_style)
            section_index += 1
            if node.is_reference:
                ref_text = contents.get(node.title, {}).get("__raw__", "").strip()
                if not ref_text:
                    _win32_add_paragraph(doc, "（本节未生成内容）", normal_style)
                    continue
                for line in ref_text.splitlines():
                    if line.strip():
                        _win32_add_paragraph(doc, line.strip(), normal_style)
                continue

            content_map = contents.get(node.title, {})
            for level2 in node.children:
                _win32_add_paragraph(doc, level2.title, h2_style)
                intro_text = (content_map.get(_h2_key(level2.title)) or "").strip() if isinstance(content_map, dict) else ""
                if level2.children:
                    if intro_text:
                        for paragraph in _split_paragraphs(intro_text):
                            _win32_add_paragraph(doc, paragraph, normal_style)
                    for level3 in level2.children:
                        _win32_add_paragraph(doc, level3.title, h3_style)
                        text = (
                            (content_map.get(_h3_key(level2.title, level3.title)) or "").strip()
                            if isinstance(content_map, dict)
                            else ""
                        )
                        if not text:
                            _win32_add_paragraph(doc, "（本节未生成内容）", normal_style)
                            continue
                        for paragraph in _split_paragraphs(text):
                            _win32_add_paragraph(doc, paragraph, normal_style)
                else:
                    text = intro_text
                    if not text:
                        _win32_add_paragraph(doc, "（本节未生成内容）", normal_style)
                        continue
                    for paragraph in _split_paragraphs(text):
                        _win32_add_paragraph(doc, paragraph, normal_style)

        try:
            doc.Fields.Update()
        except Exception:
            pass
        try:
            if doc.TablesOfContents.Count > 0:
                for idx in range(1, doc.TablesOfContents.Count + 1):
                    doc.TablesOfContents(idx).Update()
        except Exception:
            pass
        ensure_parent(path)
        com_retry(lambda: doc.SaveAs(str(path)))
    finally:
        if doc is not None:
            try:
                com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
            except Exception:
                pass
        if word is not None:
            try:
                com_retry(lambda: word.Quit(), timeout_s=5.0)
            except Exception:
                pass
        if message_filter is not None:
            restore_ole_message_filter(old_filter)
        pythoncom.CoUninitialize()


def _write_report_docx(
    path: Path,
    topic: str,
    outline_root: OutlineNode,
    outline_items: list[OutlineItem],
    contents: dict[str, dict],
) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading(f"{topic} 课题报告", level=0)
    _write_outline_table_docx(doc, outline_items)
    section_index = 1
    for node in outline_root.children:
        doc.add_heading(f"{section_index}. {node.title}", level=1)
        section_index += 1
        if node.is_reference:
            ref_text = contents.get(node.title, {}).get("__raw__", "").strip()
            if not ref_text:
                doc.add_paragraph("（本节未生成内容）")
                continue
            for line in ref_text.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
            continue
        content_map = contents.get(node.title, {})
        for level2 in node.children:
            doc.add_heading(level2.title, level=2)
            intro_text = (content_map.get(_h2_key(level2.title)) or "").strip() if isinstance(content_map, dict) else ""
            if level2.children:
                if intro_text:
                    for paragraph in _split_paragraphs(intro_text):
                        doc.add_paragraph(paragraph)
                for level3 in level2.children:
                    doc.add_heading(level3.title, level=3)
                    text = (
                        (content_map.get(_h3_key(level2.title, level3.title)) or "").strip()
                        if isinstance(content_map, dict)
                        else ""
                    )
                    if not text:
                        doc.add_paragraph("（本节未生成内容）")
                        continue
                    for paragraph in _split_paragraphs(text):
                        doc.add_paragraph(paragraph)
            else:
                text = intro_text
                if not text:
                    doc.add_paragraph("（本节未生成内容）")
                    continue
                for paragraph in _split_paragraphs(text):
                    doc.add_paragraph(paragraph)
    ensure_parent(path)
    doc.save(path)


def _write_report_text(
    path: Path,
    topic: str,
    outline_root: OutlineNode,
    outline_items: list[OutlineItem],
    contents: dict[str, dict],
) -> None:
    lines: list[str] = [f"{topic} 课题报告", ""]
    lines.append("报告大纲表")
    for item in outline_items:
        lines.append(f"  L{item.level} {item.title}")
    lines.append("")
    section_index = 1
    for node in outline_root.children:
        lines.append(f"{section_index}. {node.title}")
        section_index += 1
        if node.is_reference:
            lines.append(contents.get(node.title, {}).get("__raw__", "").strip() or "（本节未生成内容）")
            lines.append("")
            continue
        content_map = contents.get(node.title, {})
        for level2 in node.children:
            lines.append(f"  {level2.title}")
            intro_text = (content_map.get(_h2_key(level2.title)) or "").strip() if isinstance(content_map, dict) else ""
            if level2.children:
                if intro_text:
                    lines.append(intro_text)
                for level3 in level2.children:
                    lines.append(f"    {level3.title}")
                    text = (
                        (content_map.get(_h3_key(level2.title, level3.title)) or "").strip()
                        if isinstance(content_map, dict)
                        else ""
                    )
                    lines.append(text or "（本节未生成内容）")
            else:
                lines.append(intro_text or "（本节未生成内容）")
        lines.append("")
    ensure_parent(path)
    path.write_text("\n".join(lines), encoding="utf-8")


def _has_win32() -> bool:
    if os.name != "nt":
        return False
    try:
        import win32com.client  # type: ignore  # noqa: F401
    except Exception:
        return False
    return True


def _heading_level_from_style(style_name: str) -> int | None:
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name in {"title", "subtitle"}:
        return 1
    if name.startswith("heading"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return int(parts[1])
        match = re.search(r"heading\s*(\d+)", name)
        if match:
            return int(match.group(1))
        return 1
    if style_name.strip().startswith("标题"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return int(parts[1])
        match = re.search(r"标题\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
        return 1
    if style_name.strip().startswith("副标题"):
        return 2
    return None


def _win32_heading_level(para, constants) -> int | None:
    outline_level = None
    try:
        outline_level = int(getattr(para, "OutlineLevel", 0))
    except Exception:
        outline_level = None
    if outline_level:
        try:
            body_level = int(constants.wdOutlineLevelBodyText) if constants is not None else 10
        except Exception:
            body_level = 10
        if outline_level != body_level:
            return outline_level
    style_name = ""
    try:
        style = para.Range.Style
        style_name = str(getattr(style, "NameLocal", "") or getattr(style, "Name", ""))
    except Exception:
        style_name = ""
    return _heading_level_from_style(style_name)


def _extract_topic_from_docx(docx_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return docx_path.stem
    try:
        doc = Document(str(docx_path))
    except Exception:
        return docx_path.stem
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if "课题报告" in text:
            return text.replace("课题报告", "").strip() or text
        style_name = para.style.name if para.style else ""
        if style_name in {"Title", "标题"}:
            return text
        return text
    return docx_path.stem


def _replace_paragraph_with_paragraphs(para, paragraphs: list[str], normal_style) -> None:
    if not paragraphs:
        paragraphs = [""]
    rng = para.Range
    rng.End -= 1
    rng.Text = paragraphs[0]
    if normal_style is not None:
        try:
            rng.Style = normal_style
        except Exception:
            pass
    current = para
    for extra in paragraphs[1:]:
        try:
            current.Range.InsertParagraphAfter()
            current = current.Next()
        except Exception:
            break
        if current is None:
            break
        rng = current.Range
        rng.End -= 1
        rng.Text = extra
        if normal_style is not None:
            try:
                rng.Style = normal_style
            except Exception:
                pass


def _insert_paragraphs_after(para, paragraphs: list[str], normal_style) -> None:
    if not paragraphs:
        paragraphs = [""]
    current = para
    for text in paragraphs:
        try:
            current.Range.InsertParagraphAfter()
            current = current.Next()
        except Exception:
            break
        if current is None:
            break
        rng = current.Range
        rng.End -= 1
        rng.Text = text
        if normal_style is not None:
            try:
                rng.Style = normal_style
            except Exception:
                pass


def _build_reference_text(sources: list[SourceItem], limit: int = 25) -> str:
    if not sources:
        return "（未检索到可用参考文献）"
    date_str = dt.date.today().isoformat()
    lines: list[str] = []
    for idx, source in enumerate(sources[:limit], start=1):
        title = source.title or "无标题"
        url = source.url or ""
        if url:
            lines.append(f"[{idx}] {title}. {url}（访问日期：{date_str}）")
        else:
            lines.append(f"[{idx}] {title}。（访问日期：{date_str}）")
    return "\n".join(lines)


def generate_report(
    settings: AppSettings,
    topic: str,
    output_path: str,
    *,
    framework_text: str | None = None,
    total_chars: int = 10000,
    allow_web_search: bool = True,
    max_results_per_query: int = 5,
    section_timeout: int | None = 300,
    max_section_retries: int = 2,
    section_workers: int = 3,
    format_profile: str | None = None,
    logger=None,
) -> dict:
    topic = (topic or "").strip()
    if not topic:
        raise ValueError("topic is required")

    logs: list[str] = []
    main_thread_id = threading.get_ident()
    log_queue: "queue.SimpleQueue[str]" = queue.SimpleQueue()

    def _flush_logger() -> None:
        if not logger:
            return
        if threading.get_ident() != main_thread_id:
            return
        while True:
            try:
                entry = log_queue.get_nowait()
            except Exception:
                break
            try:
                logger(entry)
            except Exception:
                continue

    def _log(message: str) -> None:
        timestamp = dt.datetime.now().strftime("%H:%M:%S")
        entry = f"[{timestamp}] {message}"
        logs.append(entry)
        if not logger:
            return
        if threading.get_ident() != main_thread_id:
            log_queue.put(entry)
            return
        _flush_logger()
        try:
            logger(entry)
        except Exception:
            return

    framework_text = (framework_text or "").strip() or DEFAULT_FRAMEWORK
    _log("开始生成三级标题大纲")
    outline_items = _build_outline_with_llm(settings, topic, framework_text)
    outline_root = _items_to_tree(outline_items)
    _ensure_three_levels(outline_root)
    _mark_reference_sections(outline_root)
    _allocate_outline_targets(outline_root, total_chars)
    outline_items = _outline_items_from_tree(outline_root)
    outline_titles = [node.title for node in outline_root.children]
    _log("三级标题大纲生成完成")
    for item in outline_items:
        _log(f"[outline] L{item.level} {item.title}")

    sources: list[SourceItem] = []
    search_errors: list[str] = []
    if allow_web_search:
        queries = _build_search_queries(topic, outline_root.children)
        _log(f"开始联网检索：{len(queries)} 条查询")
        sources, search_errors = _run_search(queries, max_results=max_results_per_query, logger=_log)
        if not sources:
            _log("[search] 未检索到有效资料，将在无外部资料情况下继续生成（内容会更概括）")
        else:
            _log(f"检索完成：共收集 {len(sources)} 条结果")
        if search_errors:
            for err in search_errors:
                _log(f"[search] error: {err}")

    sources_text = _format_sources(sources, max_items=12, max_chars=3000) if sources else ""
    system_prompt = _build_report_system_prompt()
    agent_local = threading.local()

    def _get_agent():
        agent = getattr(agent_local, "agent", None)
        if agent is not None:
            return agent
        try:
            report_settings = replace(settings, persist_sessions=False, checkpoint_path=None)
        except Exception:
            report_settings = settings
        agent = build_agent(report_settings, tools=[], system_prompt=system_prompt)
        agent_local.agent = agent
        return agent

    contents: dict[str, dict] = {}

    def _generate_section(node: OutlineNode) -> tuple[str, dict]:
        agent = _get_agent()
        if node.is_reference:
            return node.title, {"__raw__": _build_reference_text(sources)}
        started_at = dt.datetime.now()
        _log(f"开始生成章节：{node.title}")
        section_sources = _select_sources_for_node(sources, node, max_items=8) if sources else []
        section_sources_text = (
            _format_sources(section_sources, max_items=8, max_chars=2200) if section_sources else sources_text
        )
        attempt = 0
        cleaned = ""
        content_map: dict[ContentKey, str] = {}
        missing: list[ContentKey] = []
        leaf_target, min_chars = _leaf_length_targets(node)
        while attempt <= max_section_retries:
            prompt = (
                _build_section_prompt(topic, node, outline_titles, section_sources_text)
                if attempt == 0
                else _build_retry_prompt(topic, node, outline_titles, section_sources_text, missing)
            )
            payload = {"messages": [{"role": "user", "content": prompt}]}
            thread_id = f"report_{node.title}_{dt.datetime.now().timestamp()}_{attempt}"
            result = _invoke_agent(agent, payload, thread_id, section_timeout)
            if result is None:
                _log(f"章节超时：{node.title}（第{attempt + 1}次）")
                attempt += 1
                continue
            content = result["messages"][-1].content if result.get("messages") else ""
            cleaned = _clean_section_text(content, node.title)
            cleaned = _ensure_length(agent, node, cleaned)
            blocks = _parse_section_blocks(cleaned)
            content_map = _blocks_to_content_map(blocks)
            missing = _validate_section_output(node, content_map)
            if not missing:
                break
            _log(f"章节内容不足，准备重试：{node.title} 缺失={len(missing)}")
            attempt += 1
        if not content_map:
            if not missing:
                missing = _expected_leaf_keys(node)
            _log(f"[fill] 章节无有效输出，准备补全：{node.title}")
        if missing:
            _log(f"章节仍有缺失，开始补全子标题：{node.title}")
            content_map = _fill_missing_sections(
                agent,
                topic,
                node,
                outline_titles,
                section_sources_text,
                content_map,
                missing,
                section_timeout,
                logger=_log,
            )
            missing = _validate_section_output(node, content_map)
            if missing:
                _log(f"子标题补全后仍缺失：{node.title} 缺失={len(missing)}")
                # Final local fallback to guarantee completeness.
                for key in missing:
                    kind, h2_title, h3_title = key
                    if kind == "h3" and h3_title:
                        text = _fallback_leaf_text(topic, node, h2_title, h3_title)
                    else:
                        text = _fallback_leaf_text(topic, node, h2_title, None)
                    content_map[key] = _pad_text(text, min_chars, topic, node.title, _format_leaf_key(key))
                missing = _validate_section_output(node, content_map)
        if not content_map:
            content_map = {}
        _ensure_h2_intros(topic, node, content_map)
        content_map["__raw__"] = cleaned
        elapsed = (dt.datetime.now() - started_at).total_seconds()
        if section_timeout and elapsed > float(section_timeout):
            _log(f"[timeout] 章节耗时超出阈值：{node.title} {elapsed:.1f}s > {section_timeout}s")
        _log(f"章节完成：{node.title}（约{len(cleaned)}字，耗时{elapsed:.1f}s）")
        return node.title, content_map

    worker_count = max(1, min(section_workers, len(outline_root.children)))
    _log(f"章节并行生成：workers={worker_count}")
    with ThreadPoolExecutor(max_workers=worker_count) as executor:
        future_map = {
            executor.submit(_generate_section, node): node
            for node in outline_root.children
        }
        pending = set(future_map.keys())
        while pending:
            done, pending = wait(pending, timeout=0.25, return_when=FIRST_COMPLETED)
            _flush_logger()
            for future in done:
                node = future_map[future]
                try:
                    section_title, section_content = future.result()
                except Exception as exc:  # noqa: BLE001
                    _log(f"章节失败：{node.title} ({type(exc).__name__}: {exc!r})")
                    _log(traceback.format_exc(limit=6).rstrip())
                    contents[node.title] = _fallback_content_map_for_node(topic, node)
                    continue
                contents[section_title] = section_content
        _flush_logger()

    output_real = resolve_path(output_path)
    ensure_parent(output_real)
    docx_path = output_real if output_real.suffix.lower() == ".docx" else output_real.with_suffix(".docx")
    text_path = docx_path.with_suffix(".txt")
    sources_path = docx_path.with_suffix(".sources.json")
    outline_path = docx_path.with_suffix(".outline.json")
    stats_path = docx_path.with_suffix(".stats.json")

    report_engine = os.getenv("REPORT_DOCX_ENGINE", "auto").strip().lower()
    use_win32 = False
    if report_engine in {"win32", "win32com", "word", "com"}:
        use_win32 = _has_win32()
    elif report_engine in {"python-docx", "docx"}:
        use_win32 = False
    else:
        use_win32 = _has_win32()

    if use_win32:
        _write_report_docx_win32(docx_path, topic, outline_root, outline_items, contents)
    else:
        _write_report_docx(docx_path, topic, outline_root, outline_items, contents)

    profile = (format_profile or os.getenv("REPORT_FORMAT_PROFILE", "")).strip()
    if not profile:
        profile = "thesis_standard"
    profile = resolve_profile(profile)
    if profile != "none":
        try:
            apply_format_profile(docx_path=str(docx_path), root_dir=settings.root_dir, profile=profile)
            _log(f"[format] profile={profile}")
        except Exception as exc:  # noqa: BLE001
            message = f"[format] failed: {exc}"
            if is_com_call_rejected(exc):
                message += "；Word 正忙或存在弹窗，建议关闭所有 Word 窗口/对话框后重试"
            elif is_com_not_initialized(exc):
                message += "；COM 未初始化（多线程调用需在该线程先 CoInitialize）"
            _log(message)

    _write_report_text(text_path, topic, outline_root, outline_items, contents)
    sources_payload = {
        "topic": topic,
        "generated_at": dt.datetime.now().isoformat(sep=" "),
        "queries": _build_search_queries(topic, outline_root.children),
        "sources": [source.__dict__ for source in sources],
        "errors": search_errors,
    }
    sources_path.write_text(json.dumps(sources_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    outline_payload = {
        "topic": topic,
        "generated_at": dt.datetime.now().isoformat(sep=" "),
        "framework": framework_text,
        "outline": [item.__dict__ for item in outline_items],
    }
    outline_path.write_text(json.dumps(outline_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    stats_payload = {
        "topic": topic,
        "generated_at": dt.datetime.now().isoformat(sep=" "),
        "stats": _compute_report_stats(outline_root, contents),
    }
    stats_path.write_text(json.dumps(stats_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    _log(
        "[stats] leaf_coverage="
        f"{stats_payload['stats'].get('leaf_coverage')} "
        f"missing={stats_payload['stats'].get('leaf_missing')} "
        f"short={stats_payload['stats'].get('leaf_short')}"
    )

    return {
        "output_path": to_virtual_path(docx_path),
        "text_path": to_virtual_path(text_path),
        "sources_path": to_virtual_path(sources_path),
        "outline_path": to_virtual_path(outline_path),
        "stats_path": to_virtual_path(stats_path),
        "stats": stats_payload["stats"],
        "sections": outline_titles,
        "logs": logs,
    }


def complete_report_docx(
    settings: AppSettings,
    input_path: str,
    output_path: str,
    *,
    topic: str | None = None,
    allow_web_search: bool = True,
    max_results_per_query: int = 5,
    section_timeout: int | None = 300,
    fill_empty_headings: bool = True,
    format_profile: str | None = None,
    logger=None,
) -> dict:
    input_real = resolve_path(input_path)
    output_real = resolve_path(output_path)
    ensure_parent(output_real)
    if input_real != output_real:
        shutil.copy2(input_real, output_real)

    logs: list[str] = []

    def _log(message: str) -> None:
        timestamp = dt.datetime.now().strftime("%H:%M:%S")
        entry = f"[{timestamp}] {message}"
        logs.append(entry)
        if logger:
            logger(entry)

    resolved_topic = (topic or "").strip() or _extract_topic_from_docx(output_real) or output_real.stem
    _log(f"补全报告：{resolved_topic}")

    sources: list[SourceItem] = []
    search_errors: list[str] = []
    outline_titles: list[str] = []
    sources_text = ""
    system_prompt = _build_report_system_prompt()
    agent = build_agent(settings, tools=[], system_prompt=system_prompt)

    if not _has_win32():
        raise RuntimeError("未检测到 win32com，无法使用 Word COM 补全报告。")

    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"win32com 初始化失败：{exc}") from exc

    placeholder_texts = {"（本节未生成内容）", "(本节未生成内容)", "（本节未生成）"}
    ignored_titles = {"报告大纲表", "目录"}
    targets: list[dict[str, str | int | None]] = []

    pythoncom.CoInitialize()
    old_filter, message_filter = install_ole_message_filter()
    word = None
    doc = None
    try:
        word = com_retry(dispatch_word_application)
        word.Visible = False
        word.DisplayAlerts = 0
        doc = com_retry(lambda: word.Documents.Open(str(output_real)))
        try:
            constants = get_win32_constants(win32)
        except Exception:
            constants = None

        normal_style = None
        total = doc.Paragraphs.Count
        current_h1 = None
        current_h2 = None
        current_h3 = None
        records: list[dict] = []

        for idx in range(1, total + 1):
            para = doc.Paragraphs(idx)
            try:
                in_table = bool(para.Range.Information(constants.wdWithInTable)) if constants else False
            except Exception:
                in_table = False
            if in_table:
                continue
            text = (para.Range.Text or "").strip()
            if not text:
                continue
            level = _win32_heading_level(para, constants)
            if level is not None:
                cleaned = _strip_prefix_markers(text)
                if level <= 1:
                    current_h1 = cleaned
                    current_h2 = None
                    current_h3 = None
                    if cleaned and cleaned not in ignored_titles and cleaned not in outline_titles:
                        outline_titles.append(cleaned)
                elif level == 2:
                    current_h2 = cleaned
                    current_h3 = None
                else:
                    current_h3 = cleaned
            records.append(
                {
                    "index": idx,
                    "text": text,
                    "level": level,
                    "h1": current_h1,
                    "h2": current_h2,
                    "h3": current_h3,
                }
            )

        placeholder_indices = {
            record["index"]
            for record in records
            if record["level"] is None and record["text"] in placeholder_texts
        }

        for record in records:
            if record["level"] is None and record["text"] in placeholder_texts:
                targets.append(
                    {
                        "index": record["index"],
                        "h1": record["h1"],
                        "h2": record["h2"],
                        "h3": record["h3"],
                        "mode": "replace",
                    }
                )

        if fill_empty_headings:
            for idx, record in enumerate(records):
                level = record["level"]
                if not level:
                    continue
                deeper_found = False
                body_found = False
                placeholder_found = False
                j = idx + 1
                while j < len(records):
                    next_level = records[j]["level"]
                    if next_level and next_level <= level:
                        break
                    if next_level and next_level > level:
                        deeper_found = True
                    if records[j]["index"] in placeholder_indices:
                        placeholder_found = True
                    if records[j]["level"] is None and records[j]["text"] not in placeholder_texts:
                        body_found = True
                    j += 1
                if deeper_found or placeholder_found:
                    continue
                if not body_found:
                    targets.append(
                        {
                            "index": record["index"],
                            "h1": record["h1"],
                            "h2": record["h2"],
                            "h3": record["h3"],
                            "mode": "insert",
                        }
                    )

        if normal_style is None:
            try:
                normal_style = doc.Styles("Normal")
            except Exception:
                normal_style = None

        if not outline_titles:
            outline_titles = [resolved_topic]

        if allow_web_search:
            queries = _build_search_queries(
                resolved_topic,
                [ReportSection(title=title) for title in outline_titles],
            )
            _log(f"开始联网检索：{len(queries)} 条查询")
            try:
                sources, search_errors = _run_search(queries, max_results=max_results_per_query, logger=_log)
            except Exception as exc:  # noqa: BLE001
                _log(f"[search] error: {exc}")
                sources = []
            if sources:
                _log(f"检索完成：共收集 {len(sources)} 条结果")
            if search_errors:
                for err in search_errors:
                    _log(f"[search] error: {err}")
            sources_text = _format_sources(sources, max_items=10, max_chars=2400) if sources else ""

        _log(f"检测到待补全段落：{len(targets)} 处")

        for target in sorted(targets, key=lambda x: int(x.get("index", 0)), reverse=True):
            h1 = (target.get("h1") or "未命名章节").strip()
            h2 = (target.get("h2") or h1).strip()
            h3 = (target.get("h3") or "").strip() or None
            leaf_title = h3 or h2
            target_chars = 260 if h3 else 320
            min_chars = max(120, int(target_chars * 0.6))
            node = OutlineNode(title=h1, level=1)
            prompt = _build_missing_leaf_prompt(
                resolved_topic,
                node,
                h2,
                h3,
                outline_titles,
                sources_text,
                target_chars,
            )
            payload = {"messages": [{"role": "user", "content": prompt}]}
            thread_id = f"report_complete_{dt.datetime.now().timestamp()}"
            result = _invoke_agent(agent, payload, thread_id, section_timeout)
            if result is None:
                text = _fallback_leaf_text(resolved_topic, node, h2, h3)
                text = _pad_text(text, min_chars, resolved_topic, h1, leaf_title)
                _log(f"[fill] 段落超时，使用兜底文本：{leaf_title}")
            else:
                content = result["messages"][-1].content if result.get("messages") else ""
                text = (content or "").strip()
                if not text:
                    text = _fallback_leaf_text(resolved_topic, node, h2, h3)
                text = _pad_text(text, min_chars, resolved_topic, h1, leaf_title)
                _log(f"[fill] 段落补全：{leaf_title}（约{len(text)}字）")

            paragraphs = _split_paragraphs(text)
            try:
                para = doc.Paragraphs(int(target["index"]))
                if target.get("mode") == "insert":
                    _insert_paragraphs_after(para, paragraphs, normal_style)
                else:
                    _replace_paragraph_with_paragraphs(para, paragraphs, normal_style)
            except Exception:
                continue

        try:
            doc.Fields.Update()
        except Exception:
            pass
        try:
            if doc.TablesOfContents.Count > 0:
                for idx in range(1, doc.TablesOfContents.Count + 1):
                    doc.TablesOfContents(idx).Update()
        except Exception:
            pass
        com_retry(lambda: doc.Save())
    finally:
        if doc is not None:
            try:
                com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
            except Exception:
                pass
        if word is not None:
            try:
                com_retry(lambda: word.Quit(), timeout_s=5.0)
            except Exception:
                pass
        if message_filter is not None:
            restore_ole_message_filter(old_filter)
        pythoncom.CoUninitialize()

    profile = (format_profile or os.getenv("REPORT_FORMAT_PROFILE", "")).strip()
    if not profile:
        profile = "thesis_standard"
    profile = resolve_profile(profile)
    if profile != "none":
        try:
            apply_format_profile(docx_path=str(output_real), root_dir=settings.root_dir, profile=profile)
            _log(f"[format] profile={profile}")
        except Exception as exc:  # noqa: BLE001
            message = f"[format] failed: {exc}"
            if is_com_call_rejected(exc):
                message += "；Word 正忙或存在弹窗，建议关闭所有 Word 窗口/对话框后重试"
            elif is_com_not_initialized(exc):
                message += "；COM 未初始化（多线程调用需在该线程先 CoInitialize）"
            _log(message)

    return {
        "output_path": to_virtual_path(output_real),
        "logs": logs,
        "filled": len(targets),
    }
