from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime
import difflib
import json
from pathlib import Path
import re
from typing import Any

from docx import Document

from app.services.preset_service import ReviewPreset, get_review_preset
from app.tools.doc_map import IndexedSection, build_indexed_sections
from app.tools.path_utils import ensure_parent, resolve_path


_NUMERIC_CITATION_PATTERN = re.compile(r"\[\s*\d+(?:\s*[-,，]\s*\d+)*\s*\]")
_AUTHOR_YEAR_PATTERN = re.compile(
    r"(?:（|\()(?P<author>[^（）()]{1,24}?)(?:，|,)\s*(?P<year>(?:19|20)\d{2}[a-z]?)\s*(?:）|\))"
)
_YEAR_PATTERN = re.compile(r"(?:19|20)\d{2}")
_NUMBER_PATTERN = re.compile(r"(?<![\w.])\d+(?:\.\d+)?%?")
_ACRONYM_PATTERN = re.compile(r"\b[A-Z]{2,}[A-Z0-9-]*\b")
_LATIN_TERM_PATTERN = re.compile(r"\b[A-Za-z][A-Za-z0-9\- ]{2,}\b")
_FIGURE_CAPTION_PATTERN = re.compile(r"^\s*图\s*(\d+(?:[.-]\d+)*)")
_TABLE_CAPTION_PATTERN = re.compile(r"^\s*表\s*(\d+(?:[.-]\d+)*)")
_FIGURE_REF_PATTERN = re.compile(r"图\s*(\d+(?:[.-]\d+)*)")
_TABLE_REF_PATTERN = re.compile(r"表\s*(\d+(?:[.-]\d+)*)")
_REFERENCE_HEADING_KEYWORDS = ("参考文献", "references")


def _normalized_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "").strip()


def _safe_first(values: list[dict], default: dict) -> dict:
    return values[0] if values else default


def _severity_from_counts(*, critical: int = 0, warning: int = 0, success_when_zero: bool = True) -> str:
    if critical > 0:
        return "critical"
    if warning > 0:
        return "warning"
    return "success" if success_when_zero else "info"


def _read_doc_profile(path: Path) -> dict[str, Any]:
    sections = build_indexed_sections(str(path))
    doc = Document(str(path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            paragraphs.append(text)
    return {
        "path": str(path),
        "sections": sections,
        "paragraphs": paragraphs,
        "section_titles": [section.title for section in sections if section.title and section.title != "Document"],
        "text": "\n".join(paragraphs),
    }


def _reference_section(section: IndexedSection) -> bool:
    title = (section.title or "").strip().lower()
    return any(keyword in title for keyword in _REFERENCE_HEADING_KEYWORDS)


def _extract_references(profile: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    for section in profile["sections"]:
        if _reference_section(section):
            lines.extend([item.text.strip() for item in section.paragraphs if item.text.strip()])
    if lines:
        return lines
    fallback: list[str] = []
    for paragraph in profile["paragraphs"]:
        text = paragraph.strip()
        if re.search(r"(?:19|20)\d{2}", text) and len(text) >= 12:
            if text.count(".") >= 1 or text.count("：") >= 1 or text.count(":") >= 1:
                fallback.append(text)
    return fallback[-20:]


def _parse_reference_signature(line: str) -> tuple[str, str]:
    text = (line or "").strip()
    year_match = _YEAR_PATTERN.search(text)
    year = year_match.group(0) if year_match else ""
    author = text
    if "，" in text:
        author = text.split("，", 1)[0]
    elif "," in text:
        author = text.split(",", 1)[0]
    author = re.sub(r"[\[\(\（].*$", "", author).strip()
    author = re.sub(r"\s+", " ", author)
    return author[:30], year


def _analyze_citations(profile: dict[str, Any]) -> dict[str, Any]:
    reference_lines = _extract_references(profile)
    text_outside_refs = []
    for section in profile["sections"]:
        if _reference_section(section):
            continue
        text_outside_refs.extend([item.text for item in section.paragraphs if item.text.strip()])
    body_text = "\n".join(text_outside_refs)
    numeric_citations = sorted(set(_NUMERIC_CITATION_PATTERN.findall(body_text)))
    author_year = [
        {"author": match.group("author").strip(), "year": match.group("year").strip()}
        for match in _AUTHOR_YEAR_PATTERN.finditer(body_text)
    ]
    reference_signatures = [_parse_reference_signature(line) for line in reference_lines]
    unmatched_citations: list[dict] = []
    used_reference_indices: set[int] = set()
    for item in author_year:
        matched = False
        author_norm = _normalized_text(item["author"])
        year = item["year"][:4]
        for idx, (ref_author, ref_year) in enumerate(reference_signatures):
            if year and ref_year and year != ref_year[:4]:
                continue
            if author_norm and _normalized_text(ref_author) and author_norm[:4] in _normalized_text(ref_author):
                matched = True
                used_reference_indices.add(idx)
                break
        if not matched:
            unmatched_citations.append(item)
    unused_references = [
        reference_lines[idx]
        for idx in range(len(reference_lines))
        if idx not in used_reference_indices and reference_lines[idx].strip()
    ]
    mixed_styles = bool(numeric_citations and author_year)
    issues = len(unmatched_citations) + len(unused_references) + (1 if mixed_styles else 0)
    return {
        "severity": _severity_from_counts(critical=len(unmatched_citations), warning=len(unused_references) + (1 if mixed_styles else 0)),
        "headline": f"检测到 {len(author_year)} 处作者-年份引文、{len(reference_lines)} 条参考文献。",
        "counts": {
            "numeric_citation_count": len(numeric_citations),
            "author_year_citation_count": len(author_year),
            "reference_count": len(reference_lines),
            "unmatched_citation_count": len(unmatched_citations),
            "unused_reference_count": len(unused_references),
        },
        "issues": {
            "mixed_citation_styles": mixed_styles,
            "unmatched_citations": unmatched_citations[:20],
            "unused_references": unused_references[:20],
        },
        "score": max(0, 100 - issues * 12),
    }


def _match_section_expectation(section_titles: list[str], expectation_keywords: tuple[str, ...]) -> str | None:
    normalized_titles = [(title, _normalized_text(title)) for title in section_titles]
    for keyword in expectation_keywords:
        keyword_norm = _normalized_text(keyword)
        for title, norm in normalized_titles:
            if keyword_norm in norm:
                return title
    return None


def _analyze_structure(profile: dict[str, Any], preset: ReviewPreset) -> dict[str, Any]:
    section_titles = [section.title for section in profile["sections"] if section.title and section.title != "Document"]
    matched_sections: list[dict[str, Any]] = []
    missing: list[str] = []
    weak: list[str] = []
    title_to_para_count = {
        section.title: len([item for item in section.paragraphs if item.text.strip()])
        for section in profile["sections"]
        if section.title
    }
    for expectation in preset.section_expectations:
        matched = _match_section_expectation(section_titles, expectation.keywords)
        if not matched:
            missing.append(expectation.label)
            matched_sections.append({"key": expectation.key, "label": expectation.label, "matched_title": "", "paragraphs": 0})
            continue
        para_count = int(title_to_para_count.get(matched, 0))
        matched_sections.append(
            {
                "key": expectation.key,
                "label": expectation.label,
                "matched_title": matched,
                "paragraphs": para_count,
            }
        )
        if para_count < expectation.min_paragraphs:
            weak.append(expectation.label)
    score = max(0, 100 - len(missing) * 15 - len(weak) * 8)
    return {
        "severity": _severity_from_counts(critical=len(missing), warning=len(weak)),
        "headline": f"命中 {len(preset.section_expectations) - len(missing)}/{len(preset.section_expectations)} 个核心章节。",
        "score": score,
        "matched_sections": matched_sections,
        "missing_sections": missing,
        "weak_sections": weak,
    }


def _find_acronym_definitions(text: str) -> set[str]:
    definitions: set[str] = set()
    patterns = [
        re.compile(r"[A-Za-z\u4e00-\u9fff]{2,20}\s*[（(]\s*([A-Z]{2,}[A-Z0-9-]*)\s*[)）]"),
        re.compile(r"([A-Z]{2,}[A-Z0-9-]*)\s*[（(][A-Za-z\u4e00-\u9fff]{2,20}\s*[)）]"),
    ]
    for pattern in patterns:
        for match in pattern.finditer(text or ""):
            definitions.add(match.group(1).strip())
    return definitions


def _analyze_terms(profile: dict[str, Any]) -> dict[str, Any]:
    text = profile["text"]
    acronyms = Counter(_ACRONYM_PATTERN.findall(text))
    definitions = _find_acronym_definitions(text)
    undefined = sorted([token for token, count in acronyms.items() if count >= 1 and token not in definitions])

    latin_terms: dict[str, set[str]] = defaultdict(set)
    for match in _LATIN_TERM_PATTERN.findall(text):
        term = match.strip()
        if len(term) < 4:
            continue
        normalized = re.sub(r"[\s-]+", "", term).lower()
        latin_terms[normalized].add(term)
    variants = [
        {"normalized": key, "variants": sorted(values)}
        for key, values in latin_terms.items()
        if len(values) > 1
    ]
    variants.sort(key=lambda item: len(item["variants"]), reverse=True)
    warning_count = len(undefined) + len(variants)
    return {
        "severity": _severity_from_counts(warning=warning_count),
        "headline": f"发现 {len(undefined)} 个未显式定义缩略语，{len(variants)} 组术语写法变体。",
        "undefined_acronyms": undefined[:20],
        "term_variants": variants[:20],
        "score": max(0, 100 - len(undefined) * 6 - len(variants) * 4),
    }


def _caption_numbers(paragraphs: list[str], pattern: re.Pattern[str]) -> list[str]:
    values: list[str] = []
    for text in paragraphs:
        match = pattern.match(text.strip())
        if match:
            values.append(match.group(1))
    return values


def _reference_numbers(text: str, pattern: re.Pattern[str]) -> list[str]:
    return [match.group(1) for match in pattern.finditer(text or "")]


def _sequence_gaps(values: list[str]) -> list[str]:
    numeric_values = []
    for raw in values:
        try:
            numeric_values.append(int(raw.split(".", 1)[0].split("-", 1)[0]))
        except Exception:
            continue
    if not numeric_values:
        return []
    gaps: list[str] = []
    ordered = sorted(set(numeric_values))
    for prev, curr in zip(ordered, ordered[1:]):
        if curr - prev > 1:
            gaps.append(f"{prev}->{curr}")
    return gaps


def _analyze_figures_tables(profile: dict[str, Any]) -> dict[str, Any]:
    paragraphs = profile["paragraphs"]
    body_text = profile["text"]
    figure_captions = _caption_numbers(paragraphs, _FIGURE_CAPTION_PATTERN)
    table_captions = _caption_numbers(paragraphs, _TABLE_CAPTION_PATTERN)
    figure_refs = _reference_numbers(body_text, _FIGURE_REF_PATTERN)
    table_refs = _reference_numbers(body_text, _TABLE_REF_PATTERN)
    unreferenced_figures = [item for item in figure_captions if item not in figure_refs]
    unreferenced_tables = [item for item in table_captions if item not in table_refs]
    missing_figure_captions = [item for item in figure_refs if item not in figure_captions]
    missing_table_captions = [item for item in table_refs if item not in table_captions]
    figure_gaps = _sequence_gaps(figure_captions)
    table_gaps = _sequence_gaps(table_captions)
    warning_count = len(unreferenced_figures) + len(unreferenced_tables) + len(figure_gaps) + len(table_gaps)
    critical_count = len(missing_figure_captions) + len(missing_table_captions)
    return {
        "severity": _severity_from_counts(critical=critical_count, warning=warning_count),
        "headline": f"图题 {len(figure_captions)} 个、表题 {len(table_captions)} 个。",
        "counts": {
            "figure_caption_count": len(figure_captions),
            "table_caption_count": len(table_captions),
            "figure_reference_count": len(figure_refs),
            "table_reference_count": len(table_refs),
        },
        "issues": {
            "unreferenced_figures": unreferenced_figures[:20],
            "unreferenced_tables": unreferenced_tables[:20],
            "missing_figure_captions": missing_figure_captions[:20],
            "missing_table_captions": missing_table_captions[:20],
            "figure_number_gaps": figure_gaps,
            "table_number_gaps": table_gaps,
        },
        "score": max(0, 100 - critical_count * 10 - warning_count * 5),
    }


def _analyze_logic(profile: dict[str, Any]) -> dict[str, Any]:
    sections = profile["sections"]
    shallow_sections = [
        section.title
        for section in sections
        if section.title
        and section.title != "Document"
        and len([item for item in section.paragraphs if item.text.strip()]) <= 1
    ]
    repeated_paragraphs: list[dict[str, Any]] = []
    paragraphs = [item for item in profile["paragraphs"] if len(item.strip()) >= 20]
    for idx in range(len(paragraphs) - 1):
        current = paragraphs[idx]
        nxt = paragraphs[idx + 1]
        ratio = difflib.SequenceMatcher(a=_normalized_text(current), b=_normalized_text(nxt)).ratio()
        if ratio >= 0.92:
            repeated_paragraphs.append({"index": idx, "similarity": round(ratio, 3), "text": current[:120]})
    weak_transitions = [
        section.title
        for section in sections[1:]
        if section.title
        and section.paragraphs
        and len((section.paragraphs[0].text or "").strip()) < 30
    ]
    warning_count = len(shallow_sections) + len(repeated_paragraphs) + len(weak_transitions)
    return {
        "severity": _severity_from_counts(warning=warning_count),
        "headline": f"浅层章节 {len(shallow_sections)} 个，疑似重复段落 {len(repeated_paragraphs)} 组。",
        "shallow_sections": shallow_sections[:20],
        "repeated_paragraphs": repeated_paragraphs[:10],
        "weak_transition_sections": weak_transitions[:20],
        "score": max(0, 100 - len(shallow_sections) * 7 - len(repeated_paragraphs) * 10 - len(weak_transitions) * 4),
    }


def _numeric_tokens(text: str) -> set[str]:
    return set(_NUMBER_PATTERN.findall(text or ""))


def _year_tokens(text: str) -> set[str]:
    return set(_YEAR_PATTERN.findall(text or ""))


def _acronym_tokens(text: str) -> set[str]:
    return set(_ACRONYM_PATTERN.findall(text or ""))


def _paragraph_risk_samples(texts: list[str]) -> list[str]:
    samples: list[str] = []
    for paragraph in texts:
        stripped = paragraph.strip()
        if not stripped:
            continue
        if len(_NUMBER_PATTERN.findall(stripped)) >= 2 or len(_YEAR_PATTERN.findall(stripped)) >= 1:
            samples.append(stripped[:160])
    return samples[:12]


def _analyze_factual_risk(input_profile: dict[str, Any], output_profile: dict[str, Any] | None) -> dict[str, Any]:
    if output_profile is None:
        sensitive = _paragraph_risk_samples(input_profile["paragraphs"])
        return {
            "severity": "info",
            "headline": f"识别到 {len(sensitive)} 段含数值/年份的敏感段落，建议重点复核。",
            "sensitive_paragraphs": sensitive,
            "score": 100,
        }

    input_text = input_profile["text"]
    output_text = output_profile["text"]
    removed_numbers = sorted(_numeric_tokens(input_text) - _numeric_tokens(output_text))
    added_numbers = sorted(_numeric_tokens(output_text) - _numeric_tokens(input_text))
    removed_years = sorted(_year_tokens(input_text) - _year_tokens(output_text))
    added_years = sorted(_year_tokens(output_text) - _year_tokens(input_text))
    acronym_drift = sorted(_acronym_tokens(input_text) ^ _acronym_tokens(output_text))
    critical_count = len(removed_years) + len(added_years)
    warning_count = len(removed_numbers) + len(added_numbers) + len(acronym_drift)
    return {
        "severity": _severity_from_counts(critical=critical_count, warning=warning_count, success_when_zero=False),
        "headline": f"数值差异 {len(removed_numbers) + len(added_numbers)} 项，年份差异 {len(removed_years) + len(added_years)} 项。",
        "removed_numbers": removed_numbers[:30],
        "added_numbers": added_numbers[:30],
        "removed_years": removed_years[:20],
        "added_years": added_years[:20],
        "acronym_drift": acronym_drift[:20],
        "score": max(0, 100 - critical_count * 12 - warning_count * 4),
    }


def _analyze_single_document(profile: dict[str, Any], preset: ReviewPreset) -> dict[str, Any]:
    return {
        "citation_reference_check": _analyze_citations(profile),
        "section_structure_score": _analyze_structure(profile, preset),
        "terminology_consistency": _analyze_terms(profile),
        "figure_table_caption_check": _analyze_figures_tables(profile),
        "logic_cohesion": _analyze_logic(profile),
    }


def _build_overview(pre_review: dict[str, Any], post_review: dict[str, Any] | None, change_risk: dict[str, Any]) -> dict[str, Any]:
    cards: list[dict[str, Any]] = []
    labels = {
        "citation_reference_check": "引用与参考文献",
        "section_structure_score": "章节结构",
        "terminology_consistency": "术语一致性",
        "figure_table_caption_check": "图表与题注",
        "logic_cohesion": "逻辑与衔接",
    }
    for key, label in labels.items():
        current = post_review.get(key) if isinstance(post_review, dict) else pre_review.get(key)
        if not isinstance(current, dict):
            continue
        cards.append(
            {
                "key": key,
                "label": label,
                "severity": current.get("severity", "info"),
                "headline": current.get("headline", ""),
                "score": current.get("score", 0),
            }
        )
    cards.append(
        {
            "key": "factual_numeric_risk",
            "label": "事实与数字风险",
            "severity": change_risk.get("severity", "info"),
            "headline": change_risk.get("headline", ""),
            "score": change_risk.get("score", 0),
        }
    )
    critical_count = len([item for item in cards if item["severity"] == "critical"])
    warning_count = len([item for item in cards if item["severity"] == "warning"])
    avg_score = round(sum(int(item.get("score", 0) or 0) for item in cards) / max(len(cards), 1), 1)
    return {
        "cards": cards,
        "critical_count": critical_count,
        "warning_count": warning_count,
        "average_score": avg_score,
        "summary": f"平均分 {avg_score}；critical {critical_count} 项；warning {warning_count} 项。",
    }


def build_review_diagnostics_payload(
    *,
    input_path: str,
    preset_key: str | None,
    output_path: str | None = None,
) -> dict[str, Any]:
    input_real = resolve_path(input_path)
    output_real = resolve_path(output_path) if output_path else None
    preset = get_review_preset(preset_key)
    input_profile = _read_doc_profile(input_real)
    output_profile = _read_doc_profile(output_real) if output_real and output_real.exists() else None
    pre_review = _analyze_single_document(input_profile, preset)
    post_review = _analyze_single_document(output_profile, preset) if output_profile else None
    change_risk = _analyze_factual_risk(input_profile, output_profile)
    payload = {
        "generated_at": datetime.now().replace(microsecond=0).isoformat(sep=" "),
        "preset": preset.to_public_dict(),
        "input_path": str(input_real),
        "output_path": str(output_real) if output_real else "",
        "overview": _build_overview(pre_review, post_review, change_risk),
        "pre_review": pre_review,
        "post_review": post_review or {},
        "change_risk": change_risk,
    }
    return payload


def write_review_diagnostics(
    *,
    input_path: str,
    output_path: str,
    preset_key: str | None,
) -> Path:
    output_real = resolve_path(output_path)
    diagnostics_path = output_real.with_suffix(".diagnostics.json")
    payload = build_review_diagnostics_payload(input_path=input_path, output_path=output_path, preset_key=preset_key)
    ensure_parent(diagnostics_path)
    diagnostics_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return diagnostics_path
