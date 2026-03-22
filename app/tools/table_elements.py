from __future__ import annotations

import datetime as dt
import base64
import hashlib
import json
import os
import re
import shutil
import tempfile
import uuid
from pathlib import Path
from typing import Any
import xml.etree.ElementTree as ET

from app.tools.image_understanding import analyze_image_apiyi
from app.tools.path_utils import ensure_parent, ensure_workspace_dir, resolve_path, to_virtual_path
from app.tools.win32_utils import com_retry, dispatch_word_application, get_win32_constants, win32com_context


_PREFIX_PATTERN = re.compile(r"</?([A-Za-z_][\w.-]*):[A-Za-z_]")
_RID_PATTERN = re.compile(r"^rId\d+$", re.IGNORECASE)
_ATTR_PREFIX_PATTERN = re.compile(r"\s([A-Za-z_][\w.-]*):[A-Za-z_][\w.-]*\s*=")
_DOCX_LIKE_SUFFIXES = {".docx", ".docm", ".dotx", ".dotm"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".wmf", ".emf"}


def _sha1_file(path: Path) -> str:
    hasher = hashlib.sha1()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _find_html_assets_dir(html_path: Path, tmp_dir: Path) -> Path | None:
    assets_candidates = [
        html_path.with_suffix(".files"),
        html_path.with_name(f"{html_path.stem}_files"),
    ]
    assets_dir = next((p for p in assets_candidates if p.is_dir()), None)
    if assets_dir is not None and assets_dir.exists():
        return assets_dir
    try:
        dirs = [p for p in tmp_dir.iterdir() if p.is_dir()]
    except Exception:
        dirs = []
    return next((p for p in dirs if p.name.lower().startswith(html_path.stem.lower())), None)


def _export_clipboard_via_html(
    word,
    *,
    images_real: Path,
    stem: str,
    payload: dict[str, Any],
) -> list[Path]:
    tmp_root = ensure_workspace_dir() / "tmp_word_exports"
    tmp_root.mkdir(parents=True, exist_ok=True)
    tmp_dir = Path(tempfile.mkdtemp(prefix="table_clip_", dir=str(tmp_root)))
    keep_exports = os.getenv("KEEP_WORD_EXPORTS", "").strip().lower() in {"1", "true", "yes", "y", "on"}
    html_path = (tmp_dir / f"{stem}_{uuid.uuid4().hex}.html").resolve()

    tmp_doc = None
    try:
        tmp_doc = com_retry(lambda: word.Documents.Add())
        rng = tmp_doc.Range(0, 0)
        try:
            # Try to paste as a picture to avoid carrying external links / OLE bindings.
            com_retry(lambda: rng.PasteSpecial(DataType=9), timeout_s=5.0)  # wdPasteEnhancedMetafile
        except Exception:
            com_retry(lambda: rng.Paste(), timeout_s=5.0)
        # wdFormatFilteredHTML = 10
        com_retry(lambda: tmp_doc.SaveAs2(str(html_path), FileFormat=10))
    except Exception as exc:  # noqa: BLE001
        payload.setdefault("warnings", []).append(f"clipboard html export failed ({stem}): {exc}")
        return []
    finally:
        if tmp_doc is not None:
            try:
                com_retry(lambda: tmp_doc.Close(SaveChanges=False), timeout_s=5.0)
            except Exception:
                pass

    assets_dir = _find_html_assets_dir(html_path, tmp_dir)
    if assets_dir is None or not assets_dir.exists():
        payload.setdefault("warnings", []).append(f"clipboard html export produced no assets dir ({stem})")
        if not keep_exports:
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass
        return []

    try:
        exported_files = sorted([p for p in assets_dir.iterdir() if p.is_file()], key=lambda p: p.name)
    except Exception:
        exported_files = []

    images = [p for p in exported_files if p.suffix.lower() in _IMAGE_EXTS]
    if not images:
        payload.setdefault("warnings", []).append(f"clipboard html export produced no images ({stem})")
        if not keep_exports:
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass
        return []

    copied: list[Path] = []
    multi = len(images) > 1
    for idx, src in enumerate(images, start=1):
        name = f"{stem}_{idx:02d}{src.suffix}" if multi else f"{stem}{src.suffix}"
        dest = images_real / name
        if dest.exists():
            dest = images_real / f"{stem}_{uuid.uuid4().hex[:6]}{src.suffix}"
        try:
            shutil.copy2(src, dest)
        except Exception as exc:  # noqa: BLE001
            payload.setdefault("warnings", []).append(f"copy exported image failed ({src.name}): {exc}")
            continue
        copied.append(dest)

    if not keep_exports:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass
    return copied


def _export_shape_via_html(
    word,
    shape,
    *,
    images_real: Path,
    stem: str,
    payload: dict[str, Any],
) -> list[Path]:
    """
    Best-effort export a Word Shape to one or more image files by copying it into a
    temporary document and saving as filtered HTML (which materializes shapes/charts
    into raster images).
    """

    try:
        # Copy shape to clipboard. Some Word versions fail on Shape.Copy() for floating shapes,
        # but Select()+Selection.Copy() works more reliably.
        try:
            com_retry(lambda: shape.Select(), timeout_s=5.0)
            com_retry(lambda: word.Selection.Copy(), timeout_s=5.0)
        except Exception:
            com_retry(lambda: shape.Copy(), timeout_s=5.0)
    except Exception as exc:  # noqa: BLE001
        payload.setdefault("warnings", []).append(f"shape html export failed ({stem}): {exc}")
        return []

    return _export_clipboard_via_html(word, images_real=images_real, stem=stem, payload=payload)


def _export_inline_shape_via_html(
    word,
    inline_shape,
    *,
    images_real: Path,
    stem: str,
    payload: dict[str, Any],
) -> list[Path]:
    try:
        try:
            com_retry(lambda: inline_shape.Select(), timeout_s=5.0)
            com_retry(lambda: word.Selection.Copy(), timeout_s=5.0)
        except Exception:
            try:
                com_retry(lambda: inline_shape.Range.Select(), timeout_s=5.0)
                com_retry(lambda: word.Selection.Copy(), timeout_s=5.0)
            except Exception:
                com_retry(lambda: inline_shape.Range.Copy(), timeout_s=5.0)
    except Exception as exc:  # noqa: BLE001
        payload.setdefault("warnings", []).append(f"inline shape copy failed ({stem}): {exc}")
        return []

    return _export_clipboard_via_html(word, images_real=images_real, stem=stem, payload=payload)


def _clean_cell_text(text: str) -> str:
    # Word cell text ends with '\r\x07' (end-of-cell marker) and may contain extra CRs.
    value = (text or "").replace("\r", "").replace("\x07", "")
    return value.strip()


def _attr_local(elem: ET.Element, name: str) -> str | None:
    for key, value in elem.attrib.items():
        if key == name or key.endswith("}" + name):
            return value
    return None


def _localname(tag: str) -> str:
    if "}" in tag:
        return tag.split("}", 1)[1]
    if ":" in tag:
        return tag.split(":", 1)[1]
    return tag


def _wrap_xml_with_dummy_namespaces(xml: str) -> str:
    raw = xml or ""
    if raw:
        # WordOpenXML strings may include XML declarations / processing instructions
        # that break when we wrap fragments into a dummy <root> element.
        raw = re.sub(r"<\?xml[^>]*\?>", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"<\?mso-application[^>]*\?>", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"<\?[^>]*\?>", "", raw).strip()
    prefixes = set(_PREFIX_PATTERN.findall(raw))
    prefixes.update(_ATTR_PREFIX_PATTERN.findall(raw))
    prefixes.discard("xml")
    prefixes.discard("xmlns")
    prefixes = sorted(prefixes)
    if not prefixes:
        return f"<root>{raw}</root>"
    decls = " ".join(f'xmlns:{p}=\"urn:{p}\"' for p in prefixes)
    return f"<root {decls}>{raw}</root>"


def _extract_media_parts_from_flat_opc(xml: str) -> list[dict[str, Any]]:
    """
    Extract embedded /word/media/* binary parts from a "Flat OPC" XML string
    returned by Word's Range.WordOpenXML.

    Returns: [{name, content_type, blob}]
    """

    if not xml:
        return []

    try:
        wrapped = _wrap_xml_with_dummy_namespaces(xml)
        root = ET.fromstring(wrapped)
    except Exception:
        return []

    parts: list[dict[str, Any]] = []
    for elem in root.iter():
        if _localname(elem.tag) != "part":
            continue
        name = _attr_local(elem, "name") or ""
        if "word/media/" not in name.replace("\\", "/").lower():
            continue
        content_type = _attr_local(elem, "contentType") or ""
        if content_type and not content_type.lower().startswith("image/"):
            continue
        binary = None
        for child in list(elem):
            if _localname(child.tag) == "binaryData":
                binary = child
                break
        if binary is None:
            continue
        encoded = (binary.text or "").strip()
        if not encoded:
            continue
        try:
            blob = base64.b64decode(re.sub(r"\s+", "", encoded))
        except Exception:
            continue
        if not blob:
            continue
        parts.append({"name": name, "content_type": content_type, "blob": blob})
    return parts


def _table_spans_from_openxml(xml: str) -> tuple[int, int, dict[tuple[int, int], dict[str, int]]]:
    """
    Best-effort parse to compute (row_count, col_count, spans) where:
      spans[(r, c)] = {"rowspan": X, "colspan": Y} for the *top-left* cell.
    """
    if not xml:
        return 0, 0, {}
    try:
        wrapped = _wrap_xml_with_dummy_namespaces(xml)
        root = ET.fromstring(wrapped)
    except Exception:
        return 0, 0, {}

    tbl = None
    for elem in root.iter():
        if _localname(elem.tag) == "tbl":
            tbl = elem
            break
    if tbl is None:
        return 0, 0, {}

    col_count = 0
    for elem in tbl.iter():
        if _localname(elem.tag) == "tblGrid":
            col_count = sum(1 for child in list(elem) if _localname(child.tag) == "gridCol")
            break

    rows = [r for r in list(tbl) if _localname(r.tag) == "tr"]
    row_count = len(rows)
    spans: dict[tuple[int, int], dict[str, int]] = {}
    if row_count <= 0:
        return 0, col_count, spans

    # Fallback if tblGrid is missing.
    if col_count <= 0:
        max_cols = 0
        for tr in rows:
            total = 0
            for tc in [c for c in list(tr) if _localname(c.tag) == "tc"]:
                tcpr = next((x for x in list(tc) if _localname(x.tag) == "tcPr"), None)
                span_val = 1
                if tcpr is not None:
                    grid_span = next((x for x in list(tcpr) if _localname(x.tag) == "gridSpan"), None)
                    if grid_span is not None:
                        raw = _attr_local(grid_span, "val")
                        if raw and raw.isdigit():
                            span_val = max(1, int(raw))
                total += span_val
            max_cols = max(max_cols, total)
        col_count = max_cols

    start_by_col: dict[int, tuple[int, int]] = {}
    for r in range(1, col_count + 1):
        start_by_col[r] = (0, 0)
    start_by_col = {k: (0, 0) for k in range(1, col_count + 1)}

    for r_idx, tr in enumerate(rows, start=1):
        tcs = [c for c in list(tr) if _localname(c.tag) == "tc"]
        filled: set[int] = set()
        col_cursor = 1
        for tc in tcs:
            while col_cursor <= col_count and col_cursor in filled:
                col_cursor += 1
            if col_cursor > col_count:
                break

            tcpr = next((x for x in list(tc) if _localname(x.tag) == "tcPr"), None)
            colspan = 1
            vmerge_val: str | None = None
            if tcpr is not None:
                grid_span = next((x for x in list(tcpr) if _localname(x.tag) == "gridSpan"), None)
                if grid_span is not None:
                    raw = _attr_local(grid_span, "val")
                    if raw and raw.isdigit():
                        colspan = max(1, int(raw))
                vmerge = next((x for x in list(tcpr) if _localname(x.tag) == "vMerge"), None)
                if vmerge is not None:
                    vmerge_val = _attr_local(vmerge, "val") or "continue"

            c_idx = col_cursor
            colspan = max(1, min(colspan, max(1, col_count - c_idx + 1)))

            if vmerge_val and vmerge_val != "restart":
                # Continuation of a vertical merge: extend the restart cell above.
                start = start_by_col.get(c_idx, (0, 0))
                if start == (0, 0):
                    # If the XML is unexpected, treat as a normal cell.
                    start = (r_idx, c_idx)
                    spans[start] = {"rowspan": 1, "colspan": colspan}
                    for cc in range(c_idx, c_idx + colspan):
                        start_by_col[cc] = start
                        filled.add(cc)
                else:
                    if start in spans:
                        spans[start]["rowspan"] += 1
                    for cc in range(c_idx, c_idx + colspan):
                        start_by_col[cc] = start
                        filled.add(cc)
            else:
                start = (r_idx, c_idx)
                spans.setdefault(start, {"rowspan": 1, "colspan": colspan})
                for cc in range(c_idx, c_idx + colspan):
                    filled.add(cc)
                    if vmerge_val == "restart":
                        start_by_col[cc] = start
                    else:
                        start_by_col[cc] = (0, 0)
            col_cursor = c_idx + colspan

    # Normalize: ensure all spans have valid values.
    for key, val in list(spans.items()):
        val["rowspan"] = max(1, int(val.get("rowspan", 1)))
        val["colspan"] = max(1, int(val.get("colspan", 1)))
        spans[key] = val

    return row_count, col_count, spans


def _build_cover_map(spans: dict[tuple[int, int], dict[str, int]]) -> dict[tuple[int, int], tuple[int, int]]:
    cover: dict[tuple[int, int], tuple[int, int]] = {}
    for (r, c), span in spans.items():
        rs = int(span.get("rowspan", 1) or 1)
        cs = int(span.get("colspan", 1) or 1)
        for rr in range(r, r + rs):
            for cc in range(c, c + cs):
                cover[(rr, cc)] = (r, c)
    return cover


def _normalize_engine(engine: str | None) -> str:
    value = (engine or "").strip().lower()
    if not value or value == "auto":
        return "auto"
    if value in {"win32", "win32com", "com", "word"}:
        return "win32com"
    if value in {"python-docx", "python_docx", "pydocx", "docx"}:
        return "python-docx"
    raise ValueError(f"Unknown table extraction engine: {engine}")


def _win32_disabled() -> bool:
    value = os.getenv("DISABLE_WIN32", "").strip().lower()
    if value in {"1", "true", "yes", "y", "on"}:
        return True
    return os.getenv("REVISION_ENGINE", "").strip().lower() == "python-docx"


def _can_import_win32() -> bool:
    if os.name != "nt":
        return False
    try:
        import win32com.client  # type: ignore  # noqa: F401

        return True
    except Exception:
        return False


def _has_win32() -> bool:
    if _win32_disabled():
        return False
    return _can_import_win32()


def _iter_rel_ids_from_cell_xml(xml: str) -> list[str]:
    if not xml:
        return []
    try:
        wrapped = _wrap_xml_with_dummy_namespaces(xml)
        root = ET.fromstring(wrapped)
    except Exception:
        return []

    rel_ids: list[str] = []
    for elem in root.iter():
        name = _localname(elem.tag)
        if name not in {"blip", "imagedata"}:
            continue
        for key, value in elem.attrib.items():
            local = _localname(key)
            if name == "blip" and local not in {"embed", "link"}:
                continue
            if name == "imagedata" and local not in {"id", "relid"}:
                continue
            if not value:
                continue
            value_str = str(value)
            if not _RID_PATTERN.match(value_str):
                continue
            rel_ids.append(value_str)

    # Preserve order, unique.
    seen: set[str] = set()
    unique: list[str] = []
    for rid in rel_ids:
        if rid in seen:
            continue
        seen.add(rid)
        unique.append(rid)
    return unique


def _collect_docx_tables(doc) -> list[Any]:
    """
    Collect tables recursively (Document.tables excludes nested tables).
    Returns a stable list without duplicates.
    """

    collected: list[Any] = []
    seen: set[int] = set()

    def visit_table(tbl) -> None:
        try:
            key = id(tbl._tbl)
        except Exception:
            key = id(tbl)
        if key in seen:
            return
        seen.add(key)
        collected.append(tbl)
        try:
            rows = list(getattr(tbl, "rows", []) or [])
        except Exception:
            rows = []
        for row in rows:
            try:
                cells = list(getattr(row, "cells", []) or [])
            except Exception:
                cells = []
            for cell in cells:
                try:
                    nested_tables = list(getattr(cell, "tables", []) or [])
                except Exception:
                    nested_tables = []
                for nested in nested_tables:
                    visit_table(nested)

    try:
        top_tables = list(getattr(doc, "tables", []) or [])
    except Exception:
        top_tables = []
    for table in top_tables:
        visit_table(table)
    return collected


def _extract_table_elements_python_docx(
    input_real: Path,
    *,
    images_real: Path,
    analyze_images: bool,
    prompt: str,
    max_cell_text_chars: int,
) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for table extraction") from exc

    payload: dict[str, Any] = {
        "document": str(input_real.name),
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "engine": "python-docx",
        "images_dir": str(images_real),
        "virtual_images_dir": to_virtual_path(images_real),
        "tables": [],
        "errors": [],
        "warnings": [],
    }

    try:
        doc = Document(str(input_real))
    except Exception as exc:  # noqa: BLE001
        payload["errors"].append(str(exc))
        return payload

    tables = _collect_docx_tables(doc)
    for ti, table in enumerate(tables, start=1):
        table_item: dict[str, Any] = {
            "table_index": ti,
            "range": {"start": 0, "end": 0},
            "rows": 0,
            "cols": 0,
            "cells": [],
        }

        spans: dict[tuple[int, int], dict[str, int]] = {}
        xml_rows = 0
        xml_cols = 0
        try:
            xml = str(getattr(getattr(table, "_tbl", None), "xml", "") or "")
            xml_rows, xml_cols, spans = _table_spans_from_openxml(xml)
        except Exception:
            xml_rows, xml_cols, spans = 0, 0, {}

        rows = 0
        cols = 0
        if xml_rows and xml_cols:
            rows, cols = xml_rows, xml_cols
        else:
            try:
                rows = int(len(table.rows))
            except Exception:
                rows = 0
            try:
                cols = int(len(table.columns))
            except Exception:
                cols = 0
            if not cols and rows:
                try:
                    cols = max((len(r.cells) for r in table.rows), default=0)
                except Exception:
                    cols = 0

        table_item["rows"] = rows
        table_item["cols"] = cols

        cell_starts = sorted(spans.keys()) if spans else []
        if not cell_starts and rows and cols:
            cell_starts = [(r, c) for r in range(1, rows + 1) for c in range(1, cols + 1)]

        for r, c in cell_starts:
            try:
                cell = table.cell(int(r) - 1, int(c) - 1)
            except Exception:
                continue

            text = ""
            try:
                text = (cell.text or "").strip()
            except Exception:
                text = ""
            if len(text) > max_cell_text_chars:
                text = text[: max_cell_text_chars - 1] + "…"

            span = spans.get((r, c), {"rowspan": 1, "colspan": 1})
            cell_item: dict[str, Any] = {
                "row": int(r),
                "col": int(c),
                "rowspan": int(span.get("rowspan", 1) or 1),
                "colspan": int(span.get("colspan", 1) or 1),
                "text": text,
                "elements": {
                    "paragraphs": 0,
                    "hyperlinks": 0,
                    "omaths": 0,
                    "fields": 0,
                    "content_controls": 0,
                    "nested_tables": 0,
                },
                "images": [],
            }

            try:
                cell_item["elements"]["paragraphs"] = int(len(getattr(cell, "paragraphs", []) or []))
            except Exception:
                pass
            try:
                cell_item["elements"]["nested_tables"] = int(len(getattr(cell, "tables", []) or []))
            except Exception:
                pass

            cell_xml = ""
            try:
                cell_xml = str(getattr(getattr(cell, "_tc", None), "xml", "") or "")
            except Exception:
                cell_xml = ""
            rel_ids = _iter_rel_ids_from_cell_xml(cell_xml)

            for ii, rid in enumerate(rel_ids, start=1):
                try:
                    rel = doc.part.rels.get(rid)
                except Exception:
                    rel = None
                if rel is None:
                    continue
                try:
                    if bool(getattr(rel, "is_external", False)):
                        payload["warnings"].append(f"skip external image rel: {rid}")
                        continue
                except Exception:
                    pass

                try:
                    part = rel.target_part
                    content_type = str(getattr(part, "content_type", "") or "")
                    if not content_type.startswith("image/"):
                        continue
                    partname = str(getattr(part, "partname", "") or "")
                    ext = Path(partname).suffix or ".bin"
                    blob = bytes(getattr(part, "blob", b"") or b"")
                except Exception:
                    continue
                if not blob:
                    continue

                image_path = images_real / f"T{ti}_R{r}C{c}_{ii:03d}{ext}"
                try:
                    image_path.write_bytes(blob)
                except Exception as exc:  # noqa: BLE001
                    payload["warnings"].append(f"failed to write image {image_path.name}: {exc}")
                    continue

                image_item: dict[str, Any] = {"kind": "embedded", "path": str(image_path), "rel_id": rid}
                image_item["virtual_path"] = to_virtual_path(image_path)
                if analyze_images:
                    try:
                        image_item["analysis"] = analyze_image_apiyi(image_path, prompt=prompt)
                    except Exception as exc:  # noqa: BLE001
                        payload["warnings"].append(f"image analysis failed ({image_path.name}): {exc}")
                cell_item["images"].append(image_item)

            table_item["cells"].append(cell_item)

        payload["tables"].append(table_item)

    return payload


def extract_table_elements(
    input_path: str,
    *,
    output_path: str | None = None,
    images_dir: str | None = None,
    analyze_images: bool | None = None,
    image_prompt: str | None = None,
    engine: str | None = None,
    max_cell_text_chars: int | None = None,
) -> str:
    """
    Extract table elements (cells, merged spans, inline/anchored images) from a Word document.

    Returns a virtual path to the JSON output.

    Engine:
      - auto (default): win32com when available; otherwise python-docx
      - win32com: requires Windows + pywin32 + Microsoft Word (supports .docx and legacy .doc)
      - python-docx: best-effort fallback (no Word required; docx-like only)

    Env vars:
      - EXTRACT_TABLE_ELEMENTS: enable in workflows/UI
      - TABLE_IMAGE_UNDERSTANDING: whether to call image analyzer
      - TABLE_IMAGE_PROMPT: prompt for the analyzer
      - TABLE_MAX_CELL_TEXT_CHARS: truncate cell text in JSON
    """
    input_real = resolve_path(input_path)
    requested_engine = _normalize_engine(engine)
    resolved_engine = requested_engine
    if requested_engine == "auto":
        if input_real.suffix.lower() in _DOCX_LIKE_SUFFIXES:
            resolved_engine = "win32com" if _has_win32() else "python-docx"
        else:
            resolved_engine = "win32com" if _has_win32() else "win32com"
    elif requested_engine == "win32com" and not _can_import_win32():
        raise RuntimeError("win32com table extraction requires Windows + pywin32 + Microsoft Word")
    elif requested_engine == "python-docx" and input_real.suffix.lower() not in _DOCX_LIKE_SUFFIXES:
        raise ValueError("python-docx table extraction requires a docx-like input (.docx/.docm/.dotx/.dotm)")
    elif requested_engine in {"auto", "win32com"} and not _has_win32() and input_real.suffix.lower() not in _DOCX_LIKE_SUFFIXES:
        raise RuntimeError("table extraction for non-docx requires Windows + pywin32 + Microsoft Word (win32com)")

    if output_path:
        output_real = resolve_path(output_path)
    else:
        output_real = input_real.with_suffix(".tables.json")
    ensure_parent(output_real)

    workspace = ensure_workspace_dir()
    if images_dir:
        images_real = resolve_path(images_dir)
    else:
        images_real = workspace / "table_images" / input_real.stem / dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    images_real.mkdir(parents=True, exist_ok=True)

    if analyze_images is None:
        analyze_images = os.getenv("TABLE_IMAGE_UNDERSTANDING", "").strip().lower() in {"1", "true", "yes"}
    prompt = (image_prompt or os.getenv("TABLE_IMAGE_PROMPT", "") or "描述分析这张图").strip()
    if max_cell_text_chars is None:
        try:
            max_cell_text_chars = int(os.getenv("TABLE_MAX_CELL_TEXT_CHARS", "2000"))
        except Exception:
            max_cell_text_chars = 2000
    max_cell_text_chars = max(200, int(max_cell_text_chars))

    if resolved_engine == "python-docx":
        payload = _extract_table_elements_python_docx(
            input_real,
            images_real=images_real,
            analyze_images=bool(analyze_images),
            prompt=prompt,
            max_cell_text_chars=max_cell_text_chars,
        )
        output_real.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return to_virtual_path(output_real)

    payload: dict[str, Any] = {
        "document": str(input_real.name),
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "engine": "win32com",
        "images_dir": str(images_real),
        "virtual_images_dir": to_virtual_path(images_real),
        "tables": [],
        "errors": [],
        "warnings": [],
    }

    with win32com_context():
        import win32com.client as win32  # type: ignore

        word = None
        doc = None
        try:
            word = com_retry(dispatch_word_application)
            word.Visible = False
            word.DisplayAlerts = 0
            constants = get_win32_constants(win32)
            try:
                word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
            except Exception:
                pass
            try:
                word.Options.UpdateLinksAtOpen = False
            except Exception:
                pass
            doc = com_retry(
                lambda: word.Documents.Open(
                    str(input_real),
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
            )

            table_count = int(getattr(doc.Tables, "Count", 0) or 0)
            tables_by_start: dict[int, dict[str, Any]] = {}
            cell_maps_by_start: dict[int, dict[tuple[int, int], dict[str, Any]]] = {}
            cover_maps_by_start: dict[int, dict[tuple[int, int], tuple[int, int]]] = {}
            spans_by_start: dict[int, dict[tuple[int, int], dict[str, int]]] = {}

            for ti in range(1, table_count + 1):
                tbl = doc.Tables(ti)
                try:
                    start = int(tbl.Range.Start)
                    end = int(tbl.Range.End)
                except Exception:
                    start = 0
                    end = 0

                try:
                    rows = int(tbl.Rows.Count)
                except Exception:
                    rows = 0
                try:
                    cols = int(tbl.Columns.Count)
                except Exception:
                    cols = 0

                openxml = ""
                try:
                    openxml = str(getattr(tbl.Range, "WordOpenXML", "") or "")
                except Exception:
                    openxml = ""
                if not openxml:
                    try:
                        openxml = str(getattr(tbl.Range, "XML", "") or "")
                    except Exception:
                        openxml = ""

                xml_rows, xml_cols, spans = _table_spans_from_openxml(openxml)
                if xml_rows:
                    rows = rows or xml_rows
                if xml_cols:
                    cols = cols or xml_cols
                cover_map = _build_cover_map(spans)

                table_item: dict[str, Any] = {
                    "table_index": ti,
                    "range": {"start": start, "end": end},
                    "rows": rows,
                    "cols": cols,
                    "cells": [],
                }
                payload["tables"].append(table_item)
                if start:
                    tables_by_start[start] = table_item
                    cover_maps_by_start[start] = cover_map
                    spans_by_start[start] = spans
                    cell_maps_by_start[start] = {}

                # Prefer iterating only top-left cells derived from spans (more stable with merged cells).
                cell_starts = sorted(spans.keys()) if spans else []
                if not cell_starts:
                    # Fallback: use Word's cell collection.
                    try:
                        for cell in tbl.Range.Cells:
                            try:
                                cell_starts.append((int(cell.RowIndex), int(cell.ColumnIndex)))
                            except Exception:
                                continue
                    except Exception:
                        cell_starts = []
                    cell_starts = sorted(set(cell_starts))

                for r, c in cell_starts:
                    try:
                        cell = tbl.Cell(r, c)
                    except Exception:
                        continue
                    raw_text = ""
                    try:
                        raw_text = str(cell.Range.Text or "")
                    except Exception:
                        raw_text = ""
                    text = _clean_cell_text(raw_text)
                    if len(text) > max_cell_text_chars:
                        text = text[: max_cell_text_chars - 1] + "…"

                    span = spans.get((r, c), {"rowspan": 1, "colspan": 1})
                    cell_item: dict[str, Any] = {
                        "row": int(r),
                        "col": int(c),
                        "rowspan": int(span.get("rowspan", 1) or 1),
                        "colspan": int(span.get("colspan", 1) or 1),
                        "text": text,
                        "elements": {},
                        "images": [],
                    }
                    table_item["cells"].append(cell_item)
                    if start:
                        cell_maps_by_start[start][(r, c)] = cell_item

                    # Basic element stats (best-effort; can be slow on huge docs).
                    try:
                        cell_item["elements"]["paragraphs"] = int(cell.Range.Paragraphs.Count)
                    except Exception:
                        pass
                    try:
                        cell_item["elements"]["hyperlinks"] = int(cell.Range.Hyperlinks.Count)
                    except Exception:
                        pass
                    try:
                        cell_item["elements"]["omaths"] = int(cell.Range.OMaths.Count)
                    except Exception:
                        pass
                    try:
                        cell_item["elements"]["fields"] = int(cell.Range.Fields.Count)
                    except Exception:
                        pass
                    try:
                        cell_item["elements"]["content_controls"] = int(cell.Range.ContentControls.Count)
                    except Exception:
                        pass
                    try:
                        nested_tables = int(cell.Range.Tables.Count)
                        # In Word, the current table is counted as 1; nested tables >1.
                        cell_item["elements"]["nested_tables"] = max(0, nested_tables - 1)
                    except Exception:
                        pass

                    # Embedded images in this cell (Flat OPC package via WordOpenXML).
                    inline_count = 0
                    shape_range_count = 0
                    try:
                        inline_count = int(cell.Range.InlineShapes.Count)
                    except Exception:
                        inline_count = 0
                    try:
                        shape_range_count = int(cell.Range.ShapeRange.Count)
                    except Exception:
                        shape_range_count = 0
                    if inline_count > 0 or shape_range_count > 0:
                        pkg_xml = ""
                        try:
                            pkg_xml = str(getattr(cell.Range, "WordOpenXML", "") or "")
                        except Exception:
                            pkg_xml = ""
                        pkg_extracted = False
                        if pkg_xml and "word/media/" in pkg_xml.replace("\\", "/").lower():
                            parts = _extract_media_parts_from_flat_opc(pkg_xml)
                            for ii, part in enumerate(parts, start=1):
                                partname = str(part.get("name", "") or "")
                                content_type = str(part.get("content_type", "") or "")
                                blob = part.get("blob") or b""
                                if not isinstance(blob, (bytes, bytearray)) or not blob:
                                    continue
                                ext = Path(partname).suffix or ".bin"
                                image_path = images_real / f"T{ti}_R{r}C{c}_{ii:03d}{ext}"
                                try:
                                    image_path.write_bytes(bytes(blob))
                                except Exception as exc:  # noqa: BLE001
                                    payload["warnings"].append(f"failed to write image {image_path.name}: {exc}")
                                    continue
                                image_item: dict[str, Any] = {
                                    "kind": "embedded",
                                    "path": str(image_path),
                                    "virtual_path": to_virtual_path(image_path),
                                    "bytes": int(len(blob)),
                                    "source_part": partname,
                                    "content_type": content_type,
                                }
                                if analyze_images:
                                    try:
                                        image_item["analysis"] = analyze_image_apiyi(image_path, prompt=prompt)
                                    except Exception as exc:  # noqa: BLE001
                                        payload["warnings"].append(f"image analysis failed ({image_path.name}): {exc}")
                                cell_item["images"].append(image_item)
                                pkg_extracted = True

                        if inline_count > 0 and not pkg_extracted:
                            seen_hashes: set[str] = set()
                            for existing in cell_item.get("images") or []:
                                sha1 = existing.get("sha1")
                                if isinstance(sha1, str) and sha1:
                                    seen_hashes.add(sha1)
                            for ii in range(1, inline_count + 1):
                                try:
                                    inline = cell.Range.InlineShapes(ii)
                                except Exception:
                                    continue
                                source = ""
                                try:
                                    link_format = getattr(inline, "LinkFormat", None)
                                    source = str(getattr(link_format, "SourceFullName", "") or "")
                                except Exception:
                                    source = ""
                                if source:
                                    cell_item["images"].append({"kind": "linked", "source": source})
                                    continue
                                stem = f"T{ti}_R{r}C{c}_inline_{ii:03d}"
                                exported_paths = _export_inline_shape_via_html(
                                    word,
                                    inline,
                                    images_real=images_real,
                                    stem=stem,
                                    payload=payload,
                                )
                                for image_path in exported_paths:
                                    try:
                                        sha1 = _sha1_file(image_path)
                                    except Exception:
                                        sha1 = ""
                                    if sha1 and sha1 in seen_hashes:
                                        try:
                                            image_path.unlink()
                                        except Exception:
                                            pass
                                        continue
                                    if sha1:
                                        seen_hashes.add(sha1)
                                    image_item = {
                                        "kind": "inline_shape",
                                        "path": str(image_path),
                                        "virtual_path": to_virtual_path(image_path),
                                    }
                                    try:
                                        image_item["bytes"] = int(image_path.stat().st_size)
                                    except Exception:
                                        image_item["bytes"] = 0
                                    if sha1:
                                        image_item["sha1"] = sha1
                                    if analyze_images:
                                        image_item["analysis"] = analyze_image_apiyi(image_path, prompt=prompt)
                                    cell_item["images"].append(image_item)

            # Floating shapes anchored in tables (pictures/charts/smartart, etc.)
            try:
                shape_count = int(doc.Shapes.Count)
            except Exception:
                shape_count = 0
            for si in range(1, shape_count + 1):
                try:
                    shape = doc.Shapes(si)
                except Exception:
                    continue
                try:
                    in_table = bool(shape.Anchor.Information(constants.wdWithInTable))
                except Exception:
                    in_table = False
                if not in_table:
                    continue
                try:
                    anchor_cell = shape.Anchor.Cells(1)
                except Exception:
                    continue
                try:
                    table_start = int(anchor_cell.Table.Range.Start)
                except Exception:
                    try:
                        table_start = int(anchor_cell.Range.Tables(1).Range.Start)
                    except Exception:
                        table_start = 0
                table_item = tables_by_start.get(table_start)
                if not table_item:
                    continue
                try:
                    ti = int(table_item.get("table_index", 0) or 0)
                except Exception:
                    ti = 0
                try:
                    ar = int(anchor_cell.RowIndex)
                    ac = int(anchor_cell.ColumnIndex)
                except Exception:
                    continue
                cover_map = cover_maps_by_start.get(table_start, {})
                start_coord = cover_map.get((ar, ac), (ar, ac))
                cell_map = cell_maps_by_start.get(table_start, {})
                cell_item = cell_map.get(start_coord)
                if cell_item is None:
                    # If the cell wasn't captured (XML parse failed), attach to a lightweight placeholder.
                    cell_item = {
                        "row": int(start_coord[0]),
                        "col": int(start_coord[1]),
                        "rowspan": 1,
                        "colspan": 1,
                        "text": "",
                        "elements": {},
                        "images": [],
                    }
                    table_item["cells"].append(cell_item)
                    cell_map[start_coord] = cell_item

                stem = f"T{ti}_R{start_coord[0]}C{start_coord[1]}_shape_{si:03d}"
                exported_paths: list[Path] = []
                kind = "shape"
                try:
                    has_chart = bool(getattr(shape, "HasChart", False))
                except Exception:
                    has_chart = False
                try:
                    shape_type = int(getattr(shape, "Type", 0) or 0)
                except Exception:
                    shape_type = 0
                # MsoShapeType.msoPicture == 13. If we've already extracted embedded images for this cell,
                # skip exporting picture-shapes again to reduce duplicates.
                is_picture = shape_type == 13
                if has_chart:
                    try:
                        image_path = images_real / f"{stem}.png"
                        shape.Chart.Export(str(image_path))
                        kind = "chart"
                        exported_paths = [image_path]
                    except Exception as exc:  # noqa: BLE001
                        payload["warnings"].append(f"chart export failed (shape #{si}): {exc}")
                        exported_paths = []

                if not exported_paths:
                    if is_picture and any((img.get("kind") == "embedded") for img in (cell_item.get("images") or [])):
                        continue
                    exported_paths = _export_shape_via_html(
                        word,
                        shape,
                        images_real=images_real,
                        stem=stem,
                        payload=payload,
                    )

                if not exported_paths:
                    continue

                for image_path in exported_paths:
                    image_item: dict[str, Any] = {"kind": kind, "path": str(image_path)}
                    image_item["virtual_path"] = to_virtual_path(image_path)
                    try:
                        image_item["bytes"] = int(image_path.stat().st_size)
                    except Exception:
                        image_item["bytes"] = 0
                    try:
                        image_item["sha1"] = _sha1_file(image_path)
                    except Exception:
                        pass
                    if analyze_images:
                        image_item["analysis"] = analyze_image_apiyi(image_path, prompt=prompt)
                    cell_item["images"].append(image_item)

        except Exception as exc:  # noqa: BLE001
            payload["errors"].append(str(exc))
        finally:
            if doc is not None:
                try:
                    com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
                except Exception:
                    pass
            if word is not None:
                try:
                    com_retry(lambda: word.Quit(), timeout_s=5.0)
                except Exception:
                    pass

    if requested_engine == "auto" and not payload.get("tables"):
        try:
            fallback = _extract_table_elements_python_docx(
                input_real,
                images_real=images_real,
                analyze_images=bool(analyze_images),
                prompt=prompt,
                max_cell_text_chars=max_cell_text_chars,
            )
        except Exception:
            fallback = {}
        if isinstance(fallback, dict) and fallback.get("tables"):
            fallback.setdefault("warnings", [])
            fallback_errors = payload.get("errors") or []
            note = "win32com returned 0 tables; fallback to python-docx"
            if fallback_errors:
                note += f" (errors={fallback_errors})"
            fallback["warnings"].append(note)
            payload = fallback

    output_real.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return to_virtual_path(output_real)
