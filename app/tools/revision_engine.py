from __future__ import annotations

from dataclasses import dataclass
import datetime as dt
import difflib
import hashlib
import importlib.util
import json
import os
from pathlib import Path
import random
import re
import shutil
import sys
import time
from typing import Callable, Iterable

from app.tools.path_utils import ensure_parent, ensure_workspace_dir, resolve_path, to_virtual_path
from app.tools.comment_cleaner import clean_comment_text
from app.tools.win32_utils import (
    com_error_hresult,
    com_retry,
    get_win32_constants,
    install_ole_message_filter,
    restore_ole_message_filter,
)

_CHINESE_NUMERALS = (
    "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d"
    "\u5341\u767e\u5343\u4e07\u4e24"
)
_TITLE_MARKER_PATTERN = re.compile(
    "^(?P<prefix>\\s*)(?P<marker>(?:"
    "(?:\\(?\\d+\\)?[.)\u3001])"
    "|(?:\\d+(?:\\.\\d+)+[.)\u3001])"
    "|(?:\\d+(?:\\.\\d+)+)(?=\\s)"
    "|(?:[A-Za-z][.)\u3001])"
    "|(?:[\u2460-\u2469])"
    "|(?:[\u56fe\u8868]\\s*(?:\\d+(?:[.-]\\d+)*|[一二三四五六七八九十]+)\\s*[)）]?\\s*[:：.\\-–]?)"
    f"|(?:[\uFF08(]?[{_CHINESE_NUMERALS}]+[\uFF09)]?\u3001)"
    f"|(?:\u7b2c[{_CHINESE_NUMERALS}0-9]+[\u7ae0\u8282\u6761\u90e8\u5206])"
    "))(?P<space>\\s*)"
)


_SENTENCE_PATTERN = re.compile(r"[^\u3002\uFF01\uFF1F!?\uFF1B;]+[\u3002\uFF01\uFF1F!?\uFF1B;]?")


def _sentences_with_spans(text: str) -> list[tuple[int, int, str]]:
    spans: list[tuple[int, int, str]] = []
    for match in _SENTENCE_PATTERN.finditer(text or ""):
        if match.group(0).strip():
            spans.append((match.start(), match.end(), match.group(0)))
    return spans


def _parse_sentence_comments(comment: str) -> list[tuple[int, str]]:
    if not comment:
        return []
    matches = list(
        re.finditer(
            r"(?:[Pp]\s*(\d+)\s*-\s*[Ss]\s*(\d+)\s*[:\uff1a]?\s*)"
            r"|(?:(?:正文\s*)?第\s*(\d+)\s*段\s*第\s*(\d+)\s*句\s*[:\uff1a]?\s*)"
            r"|(?:第\s*(\d+)\s*句\s*[:\uff1a]?\s*)",
            comment,
        )
    )
    if not matches:
        return []
    results: list[tuple[int, str]] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(comment)
        text = comment[start:end].strip().strip("\uff1b;\n " )
        if text:
            sentence_idx = None
            if match.group(2):
                sentence_idx = int(match.group(2))
            elif match.group(4):
                sentence_idx = int(match.group(4))
            elif match.group(5):
                sentence_idx = int(match.group(5))
            if sentence_idx is not None:
                results.append((sentence_idx, text))
    return results


def _extract_quoted_text(comment: str) -> list[str]:
    if not comment:
        return []
    pairs = [
        ("\u201c", "\u201d"),
        ("\"", "\""),
        ("\u300c", "\u300d"),
        ("\u300e", "\u300f"),
        ("\u300a", "\u300b"),
        ("\u2018", "\u2019"),
    ]
    results: list[str] = []
    for left, right in pairs:
        pattern = re.escape(left) + r"(.+?)" + re.escape(right)
        for match in re.findall(pattern, comment):
            text = match.strip()
            if text:
                results.append(text)
    return results


def _preserve_title_marker(original: str, revised: str) -> str:
    match = _TITLE_MARKER_PATTERN.match(original or "")
    if not match:
        return revised
    prefix = match.group(0)
    revised_trim = (revised or "").lstrip()
    revised_match = _TITLE_MARKER_PATTERN.match(revised_trim)
    if revised_match:
        revised_trim = revised_trim[revised_match.end() :]
    revised_trim = revised_trim.lstrip()
    if not revised_trim:
        return original
    return f"{prefix}{revised_trim}"


@dataclass
class RevisionInstruction:
    action: str
    paragraph_index: int
    content: str | None = None
    comment: str | None = None


class BaseRevisionEngine:
    name = "base"

    def apply(
        self,
        input_path: str,
        output_path: str,
        instructions: Iterable[RevisionInstruction],
        *,
        progress_callback: Callable[[int, RevisionInstruction, bool], None] | None = None,
        start_index: int = 0,
    ) -> bool:
        raise NotImplementedError


class Win32WordEngine(BaseRevisionEngine):
    name = "win32com"

    def __init__(self) -> None:
        if os.name != "nt":
            raise RuntimeError("win32com engine is only supported on Windows")
        import win32com.client as win32  # type: ignore
        self._win32 = win32
        self._constants = get_win32_constants(win32)

    def _clean_gen_py_modules(self) -> None:
        for key in list(sys.modules.keys()):
            if key.startswith("win32com.gen_py"):
                sys.modules.pop(key, None)

    def _try_fix_gen_py_cache(self, exc: Exception, *, aggressive: bool = False) -> None:
        message = str(exc)
        if "CLSIDToClassMap" not in message:
            raise
        try:
            import win32com.client.gencache as gencache  # type: ignore
        except Exception:
            raise
        self._clean_gen_py_modules()
        try:
            gencache.is_readonly = False
        except Exception:
            pass
        try:
            gencache.Rebuild()
            return
        except Exception:
            pass
        try:
            gen_path = gencache.GetGeneratePath()
        except Exception:
            raise
        try:
            import shutil

            shutil.rmtree(gen_path, ignore_errors=True)
        except Exception:
            raise
        try:
            gencache.Rebuild()
        except Exception:
            raise
        if aggressive:
            try:
                from win32com.client import makepy  # type: ignore
            except Exception:
                return
            try:
                makepy.GenerateFromTypeLibSpec("Microsoft Word 16.0 Object Library")
            except Exception:
                pass

    def _dispatch_word(self):
        if os.getenv("WORD_COM_ISOLATED", "1").strip().lower() not in {"0", "false", "no"}:
            return self._win32.DispatchEx("Word.Application")
        try:
            import win32com.client.gencache as gencache  # type: ignore
        except Exception:
            return self._win32.DispatchEx("Word.Application")
        try:
            return gencache.EnsureDispatch("Word.Application")
        except Exception as exc:  # noqa: BLE001
            self._try_fix_gen_py_cache(exc, aggressive=True)
            try:
                return gencache.EnsureDispatch("Word.Application")
            except Exception:
                return self._win32.DispatchEx("Word.Application")

    def _wd_no_protection(self) -> int:
        return self._const_int("wdNoProtection", -1)

    def _const_int(self, name: str, default: int) -> int:
        try:
            return int(getattr(self._constants, name))
        except Exception:
            return int(default)

    def _insert_after_comment_fallback_enabled(self) -> bool:
        value = os.getenv("REVIEW_INSERT_AFTER_COMMENT_FALLBACK", "true").strip().lower()
        return value in {"1", "true", "yes", "y", "on"}

    def _insert_after_by_new_paragraph(self, para, new_text: str):
        para.Range.InsertParagraphAfter()
        next_para = para.Next()
        if next_para is None:
            raise RuntimeError("Next paragraph unavailable after InsertParagraphAfter")
        rng = self._safe_text_range(next_para)
        if rng is None:
            rng = next_para.Range.Duplicate
            if rng.End > rng.Start:
                rng.End -= 1
        rng.Text = new_text
        return self._safe_comment_range(next_para) or rng

    def _insert_after_by_setrange(self, para, new_text: str):
        anchor = para.Range.Duplicate
        end = int(getattr(anchor, "End", 0))
        rng = para.Range.Duplicate
        start = max(end - 1, 0)
        rng.SetRange(start, start)
        rng.InsertAfter("\r" + new_text)
        return rng

    def _is_document_protected(self, doc) -> tuple[bool, int]:
        wd_no_protection = self._wd_no_protection()
        try:
            protection_type = int(getattr(doc, "ProtectionType", wd_no_protection))
        except Exception:
            protection_type = wd_no_protection
        return protection_type != wd_no_protection, protection_type

    def _try_unprotect_document(self, doc) -> bool:
        try:
            doc.Unprotect(Password="")
        except Exception as exc:  # noqa: BLE001
            print(f"[win32com] unprotect failed: {exc}")
        is_protected, protection_type = self._is_document_protected(doc)
        if is_protected:
            print(f"[win32com] document remains protected: type={protection_type}")
            return False
        print("[win32com] document unprotected")
        return True

    def _error_details_code(self, exc: BaseException) -> int | None:
        try:
            args = getattr(exc, "args", None)
            if not args or len(args) < 3:
                return None
            details = args[2]
            if not isinstance(details, tuple) or len(details) < 6:
                return None
            code = details[5]
            return int(code)
        except Exception:
            return None

    def _is_protected_edit_error(self, exc: BaseException) -> bool:
        text = str(exc) or ""
        if "不允许您编辑此所选内容，因为它受保护" in text or "受保护" in text:
            return True
        hresult = com_error_hresult(exc)
        if isinstance(hresult, int) and hresult == -2146822164:
            return True
        return self._error_details_code(exc) == -2146822164

    def _is_invalid_command_error(self, exc: BaseException) -> bool:
        text = str(exc) or ""
        if "此命令无效" in text:
            return True
        hresult = com_error_hresult(exc)
        if isinstance(hresult, int) and hresult == -2146823683:
            return True
        return self._error_details_code(exc) == -2146823683

    def _instruction_comment_text(self, inst: RevisionInstruction) -> str:
        if inst.comment:
            cleaned = clean_comment_text(inst.comment)
            if cleaned:
                return cleaned
        content = clean_comment_text((inst.content or "").strip())
        if len(content) > 300:
            content = content[:300].rstrip() + "..."
        if inst.action == "replace":
            if content:
                return f"整段：建议改为：{content}"
            return "整段：建议改写该处内容。"
        if inst.action == "insert_after":
            if content:
                return f"整段：建议新增：{content}"
            return "整段：建议补充相关内容。"
        if inst.action == "delete":
            if content:
                return f"整段：建议删除：{content}"
            return "整段：建议删除该处内容。"
        return "整段：建议优化该处表达。"

    def _try_fallback_comment(
        self,
        doc,
        para,
        inst: RevisionInstruction,
        author: str,
        comment_time: dt.datetime,
    ) -> tuple[dt.datetime, bool]:
        rng = self._safe_comment_range(para)
        if rng is None:
            return comment_time, False
        text = self._instruction_comment_text(inst)
        if not text:
            return comment_time, False
        try:
            comment_time += dt.timedelta(seconds=random.randint(30, 60))
            self._add_comment(doc, rng, text, author, comment_time)
            return comment_time, True
        except Exception:
            return comment_time, False

    def _safe_save_document(self, doc, output_path: str, *, fallback_source: str | None = None) -> None:
        target = os.path.abspath(output_path)
        target_exists = os.path.exists(target)
        current_full_name = ""
        try:
            current_full_name = os.path.abspath(str(getattr(doc, "FullName", "") or ""))
        except Exception:
            current_full_name = ""

        if current_full_name and current_full_name == target:
            try:
                com_retry(lambda: doc.Save())
                return
            except Exception as exc:  # noqa: BLE001
                print(f"[win32com] Save failed on target path: {exc}")
                # Best effort: if target already exists, keep last autosaved snapshot.
                if target_exists and os.path.exists(target):
                    print("[win32com] keep existing target snapshot (save best-effort)")
                    return

        save_attempts = [
            ("SaveAs2", lambda: doc.SaveAs2(FileName=target, AddToRecentFiles=False)),
            ("SaveAs", lambda: doc.SaveAs(FileName=target, AddToRecentFiles=False)),
            ("SaveCopyAs", lambda: doc.SaveCopyAs(target)),
        ]
        last_exc: Exception | None = None
        for name, action in save_attempts:
            try:
                com_retry(action)
                return
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                print(f"[win32com] {name} failed: {exc}")
                if not self._is_invalid_command_error(exc):
                    continue

        if current_full_name and current_full_name == target:
            try:
                com_retry(lambda: doc.Save())
                return
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                print(f"[win32com] Save (same path) failed: {exc}")
                if os.path.exists(target):
                    print("[win32com] keep existing target snapshot (save best-effort)")
                    return

        # Last resort: try to persist current doc and copy from its on-disk file.
        if current_full_name and os.path.exists(current_full_name):
            try:
                com_retry(lambda: doc.Save())
            except Exception as exc:  # noqa: BLE001
                print(f"[win32com] Save before filesystem copy failed: {exc}")
            try:
                if current_full_name != target:
                    ensure_parent(Path(target))
                    shutil.copy2(current_full_name, target)
                if os.path.exists(target):
                    print("[win32com] fallback filesystem copy succeeded")
                    return
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                print(f"[win32com] fallback filesystem copy failed: {exc}")

        if fallback_source:
            try:
                src = os.path.abspath(fallback_source)
                if os.path.exists(src):
                    ensure_parent(Path(target))
                    shutil.copy2(src, target)
                    if os.path.exists(target):
                        print("[win32com] fallback source copy succeeded")
                        return
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                print(f"[win32com] fallback source copy failed: {exc}")

        if last_exc is not None:
            raise last_exc
        raise RuntimeError("Failed to save Word document")

    def _safe_text_range(self, para) -> object | None:
        rng = para.Range.Duplicate
        if rng.End - rng.Start <= 1:
            return None
        rng.End -= 1
        if rng.End <= rng.Start:
            return None
        return rng

    def _safe_comment_range(self, para) -> object | None:
        rng = self._safe_text_range(para)
        if rng is None:
            return None
        return rng

    def _comment_author(self) -> str:
        author = os.getenv("COMMENT_AUTHOR", "\u5446\u5854\u5927\u5e08\u5144").strip()
        return author or "\u5446\u5854\u5927\u5e08\u5144"
    def _sentences_with_spans(self, text: str) -> list[tuple[int, int, str]]:
        return _sentences_with_spans(text)
    def _parse_sentence_comments(self, comment: str) -> list[tuple[int, str]]:
        return _parse_sentence_comments(comment)
    def _extract_quoted_text(self, comment: str) -> list[str]:
        return _extract_quoted_text(comment)
    def _find_comment_anchor(self, paragraph_text: str, comment: str) -> tuple[int, int] | None:
        if not paragraph_text or not comment:
            return None
        if re.search(r"(?:\u6574\u6bb5|[Pp]\s*\d+\s*-\s*\u6574\u6bb5)\s*[:\uff1a]", comment):
            return (0, len(paragraph_text))
        for quoted in _extract_quoted_text(comment):
            if quoted in paragraph_text:
                start = paragraph_text.find(quoted)
                return (start, start + len(quoted))
        matcher = difflib.SequenceMatcher(a=paragraph_text, b=comment)
        best = max(matcher.get_matching_blocks(), key=lambda block: block.size, default=None)
        if not best or best.size < 3:
            return None
        snippet = paragraph_text[best.a : best.a + best.size]
        if not re.search(r"[A-Za-z0-9\u4e00-\u9fff]", snippet):
            return None
        return (best.a, best.a + best.size)
    def _find_changed_spans(self, original: str, revised: str) -> list[tuple[int, int]]:
        matcher = difflib.SequenceMatcher(a=original, b=revised)
        spans: list[tuple[int, int]] = []
        has_delete_only = True
        for tag, _i1, _i2, j1, j2 in matcher.get_opcodes():
            if tag in ("replace", "insert") and j1 != j2:
                spans.append((j1, j2))
                has_delete_only = False
            elif tag == "replace":
                has_delete_only = False
        if not spans and has_delete_only:
            return []
        spans.sort()
        merged: list[tuple[int, int]] = []
        for start, end in spans:
            if not merged or start > merged[-1][1]:
                merged.append((start, end))
            else:
                merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        return merged

    def _compute_comment_targets(
        self,
        original: str,
        revised: str,
        comment: str,
    ) -> list[tuple[int, int, str]]:
        if not comment:
            return []
        sentences = self._sentences_with_spans(revised)
        sentence_comments = self._parse_sentence_comments(comment)
        targets: list[tuple[int, int, str]] = []
        if sentence_comments and sentences:
            for idx, text in sentence_comments:
                if 1 <= idx <= len(sentences):
                    start, end, _ = sentences[idx - 1]
                    targets.append((start, end, text))
            if targets:
                return targets

        spans = self._find_changed_spans(original, revised)
        if not spans:
            if revised:
                return [(0, len(revised), comment)]
            return []

        affected_sentences: list[tuple[int, int]] = []
        for start, end, _ in sentences:
            for span_start, span_end in spans:
                if span_start < end and span_end > start:
                    affected_sentences.append((start, end))
                    break
        affected_sentences = list(dict.fromkeys(affected_sentences))
        if len(affected_sentences) == 1:
            if len(spans) == 1:
                span_start, span_end = spans[0]
                if span_end - span_start <= 6:
                    return [(span_start, span_end, comment)]
            start, end = affected_sentences[0]
            return [(start, end, comment)]
        return [(0, len(revised), comment)]

    def _range_from_offsets(self, para, start: int, end: int) -> object | None:
        base = self._safe_text_range(para)
        if base is None:
            return None
        base_text = base.Text or ""
        length = len(base_text)
        if length == 0:
            return None
        start = max(0, min(start, length))
        end = max(0, min(end, length))
        if start >= end:
            return None
        rng = para.Range.Duplicate
        rng.SetRange(base.Start + start, base.Start + end)
        return rng

    def _add_comment(self, doc, rng, text: str, author: str, timestamp: dt.datetime) -> None:
        cleaned = clean_comment_text(text)
        if not cleaned:
            return
        comment = doc.Comments.Add(rng, Text=cleaned)
        try:
            comment.Author = author
        except Exception:
            pass
        try:
            comment.Date = timestamp
        except Exception:
            pass

    def _collect_body_paragraphs(self, doc) -> list[object]:
        body: list[object] = []
        total = doc.Paragraphs.Count
        wd_within_table = self._const_int("wdWithInTable", 12)
        for idx in range(1, total + 1):
            para = doc.Paragraphs(idx)
            try:
                in_table = bool(para.Range.Information(wd_within_table))
            except Exception:
                in_table = False
            if not in_table:
                body.append(para)
        return body

    def apply(
        self,
        input_path: str,
        output_path: str,
        instructions: Iterable[RevisionInstruction],
        *,
        progress_callback: Callable[[int, RevisionInstruction, bool], None] | None = None,
        start_index: int = 0,
    ) -> bool:
        import pythoncom  # type: ignore

        instruction_list = list(instructions)
        autosave_every_raw = os.getenv("REVIEW_AUTOSAVE_EVERY", "10").strip()
        try:
            autosave_every = max(1, int(autosave_every_raw))
        except Exception:
            autosave_every = 10
        autosave_interval_raw = os.getenv("REVIEW_AUTOSAVE_INTERVAL_SECONDS", "1.0").strip()
        try:
            autosave_interval_s = max(0.0, float(autosave_interval_raw))
        except Exception:
            autosave_interval_s = 1.0

        pythoncom.CoInitialize()
        old_filter, message_filter = install_ole_message_filter()
        try:
            word = com_retry(lambda: self._dispatch_word())
        except Exception:
            word = com_retry(lambda: self._win32.DispatchEx("Word.Application"))
        word.Visible = False
        word.DisplayAlerts = 0
        doc = None
        comment_time = dt.datetime.now()
        author = self._comment_author()
        applied_count = 0
        skipped_count = 0
        failed_count = 0
        protected_failures = 0
        fallback_comment_count = 0
        try:
            doc = com_retry(lambda: word.Documents.Open(os.path.abspath(input_path)))
            try:
                doc.TrackRevisions = True
            except Exception as exc:  # noqa: BLE001
                print(f"[win32com] TrackRevisions failed: {exc}")
            try:
                doc.ShowRevisions = True
            except Exception as exc:  # noqa: BLE001
                print(f"[win32com] ShowRevisions failed: {exc}")

            # Create output snapshot upfront so long-running reviews persist progress.
            self._safe_save_document(doc, output_path, fallback_source=input_path)

            is_protected, protection_type = self._is_document_protected(doc)
            if is_protected:
                print(f"[win32com] document protected: type={protection_type}")
                self._try_unprotect_document(doc)
                is_protected, protection_type = self._is_document_protected(doc)

            force_comment_only = False
            if is_protected:
                force_comment_only = (
                    os.getenv("REVIEW_PROTECTED_FORCE_COMMENT_ONLY", "false").strip().lower()
                    in {"1", "true", "yes", "y", "on"}
                )
                if force_comment_only:
                    print("[win32com] force comment-only mode for protected document")

            body_paragraphs = self._collect_body_paragraphs(doc)
            processed_count = 0
            last_autosave_ts = time.monotonic()

            def _maybe_autosave(*, force: bool = False) -> None:
                nonlocal last_autosave_ts
                now = time.monotonic()
                due_by_count = processed_count > 0 and processed_count % autosave_every == 0
                due_by_time = autosave_interval_s <= 0.0 or (now - last_autosave_ts) >= autosave_interval_s
                if not (force or (due_by_count and due_by_time)):
                    return
                try:
                    com_retry(lambda: doc.Save())
                    last_autosave_ts = now
                except Exception as exc:  # noqa: BLE001
                    print(f"[win32com] autosave failed: {exc}")

            for local_idx, inst in enumerate(instruction_list):
                global_next_index = start_index + local_idx + 1
                structural_applied = False
                error_stage = f"{inst.action}:start"
                if inst.paragraph_index < 0 or inst.paragraph_index >= len(body_paragraphs):
                    skipped_count += 1
                    processed_count += 1
                    if progress_callback is not None:
                        try:
                            progress_callback(global_next_index, inst, structural_applied)
                        except Exception:
                            pass
                    _maybe_autosave()
                    continue
                para = body_paragraphs[inst.paragraph_index]
                if force_comment_only and inst.action != "comment":
                    comment_time, ok = self._try_fallback_comment(doc, para, inst, author, comment_time)
                    if ok:
                        fallback_comment_count += 1
                    else:
                        skipped_count += 1
                    processed_count += 1
                    if progress_callback is not None:
                        try:
                            progress_callback(global_next_index, inst, structural_applied)
                        except Exception:
                            pass
                    _maybe_autosave()
                    continue
                try:
                    if inst.action == "comment":
                        error_stage = "comment:prepare"
                        text = inst.comment or inst.content or "Review note"
                        rng = self._safe_comment_range(para)
                        if rng is None:
                            skipped_count += 1
                            continue
                        sentence_targets = self._parse_sentence_comments(text)
                        if sentence_targets:
                            paragraph_text = rng.Text or ""
                            sentences = self._sentences_with_spans(paragraph_text)
                            for idx, reason in sentence_targets:
                                if 1 <= idx <= len(sentences):
                                    start, end, _ = sentences[idx - 1]
                                    target_rng = self._range_from_offsets(para, start, end)
                                    if target_rng is None:
                                        continue
                                    comment_time += dt.timedelta(seconds=random.randint(30, 60))
                                    self._add_comment(doc, target_rng, reason, author, comment_time)
                                    applied_count += 1
                        else:
                            paragraph_text = rng.Text or ""
                            anchor = self._find_comment_anchor(paragraph_text, text)
                            if anchor:
                                start, end = anchor
                                target_rng = self._range_from_offsets(para, start, end)
                                if target_rng is not None:
                                    comment_time += dt.timedelta(seconds=random.randint(30, 60))
                                    self._add_comment(doc, target_rng, text, author, comment_time)
                                    applied_count += 1
                                    continue
                            skipped_count += 1
                            continue
                    elif inst.action == "replace":
                        error_stage = "replace:prepare"
                        new_text = inst.content or ""
                        rng = self._safe_text_range(para)
                        if rng is None:
                            skipped_count += 1
                            continue
                        original_text = rng.Text or ""
                        new_text = _preserve_title_marker(original_text, new_text)
                        error_stage = "replace:set_text"
                        rng.Text = new_text
                        applied_count += 1
                        if inst.comment:
                            targets = self._compute_comment_targets(original_text, new_text, inst.comment)
                            if not targets:
                                targets = [(0, len(new_text), inst.comment)] if new_text else []
                            for start, end, text in targets:
                                target_rng = self._range_from_offsets(para, start, end)
                                if target_rng is None:
                                    continue
                                comment_time += dt.timedelta(seconds=random.randint(30, 60))
                                self._add_comment(doc, target_rng, text, author, comment_time)
                                applied_count += 1
                    elif inst.action == "insert_after":
                        error_stage = "insert_after:prepare"
                        new_text = inst.content or ""
                        comment_anchor = None
                        fallback_commented = False
                        if new_text:
                            try:
                                error_stage = "insert_after:new_paragraph"
                                comment_anchor = self._insert_after_by_new_paragraph(para, new_text)
                                applied_count += 1
                                structural_applied = True
                            except Exception:
                                try:
                                    error_stage = "insert_after:setrange_fallback"
                                    comment_anchor = self._insert_after_by_setrange(para, new_text)
                                    applied_count += 1
                                    structural_applied = True
                                except Exception:
                                    if not self._insert_after_comment_fallback_enabled():
                                        raise
                                    error_stage = "insert_after:comment_fallback"
                                    comment_time, ok = self._try_fallback_comment(
                                        doc,
                                        para,
                                        inst,
                                        author,
                                        comment_time,
                                    )
                                    if not ok:
                                        raise
                                    fallback_comment_count += 1
                                    fallback_commented = True
                        else:
                            skipped_count += 1
                        if inst.comment and not fallback_commented:
                            comment_time += dt.timedelta(seconds=random.randint(30, 60))
                            comment_rng = comment_anchor or self._safe_comment_range(para)
                            if comment_rng is None:
                                skipped_count += 1
                            else:
                                self._add_comment(doc, comment_rng, inst.comment, author, comment_time)
                                applied_count += 1
                    elif inst.action == "delete":
                        error_stage = "delete:prepare"
                        rng = self._safe_text_range(para)
                        if rng is None:
                            skipped_count += 1
                            continue
                        if inst.comment:
                            comment_time += dt.timedelta(seconds=random.randint(30, 60))
                            self._add_comment(doc, rng, inst.comment, author, comment_time)
                            applied_count += 1
                        rng.Delete()
                        applied_count += 1
                        structural_applied = True
                except Exception as exc:  # noqa: BLE001
                    if inst.action != "comment" and self._is_protected_edit_error(exc):
                        protected_failures += 1
                        comment_time, ok = self._try_fallback_comment(doc, para, inst, author, comment_time)
                        if ok:
                            fallback_comment_count += 1
                            continue
                    hresult = com_error_hresult(exc)
                    hresult_text = f" hresult={hresult}" if hresult is not None else ""
                    print(f"[win32com] instruction failed: {inst} (stage={error_stage}{hresult_text}; {exc})")
                    failed_count += 1
                finally:
                    processed_count += 1
                    if progress_callback is not None:
                        try:
                            progress_callback(global_next_index, inst, structural_applied)
                        except Exception:
                            pass
                    _maybe_autosave(force=structural_applied)

            try:
                com_retry(lambda: doc.Save())
            except Exception as exc:  # noqa: BLE001
                print(f"[win32com] final save failed, fallback to SaveAs chain: {exc}")
                self._safe_save_document(doc, output_path, fallback_source=input_path)
            print(
                "[win32com] done:",
                f"applied={applied_count}",
                f"fallback_comments={fallback_comment_count}",
                f"protected_failures={protected_failures}",
                f"failed={failed_count}",
                f"skipped={skipped_count}",
            )
            return True
        finally:
            if doc is not None:
                try:
                    com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
                except Exception:
                    pass
            try:
                com_retry(lambda: word.Quit(), timeout_s=5.0)
            except Exception:
                pass
            if message_filter is not None:
                restore_ole_message_filter(old_filter)
            pythoncom.CoUninitialize()


class PythonDocxEngine(BaseRevisionEngine):
    name = "python-docx"

    def __init__(self) -> None:
        from docx import Document  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.shared import RGBColor  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
        self._Document = Document
        self._OxmlElement = OxmlElement
        self._Paragraph = Paragraph
        self._RGBColor = RGBColor

    def _comment_author(self) -> str:
        author = os.getenv("COMMENT_AUTHOR", "\u5446\u5854\u5927\u5e08\u5144").strip()
        return author or "\u5446\u5854\u5927\u5e08\u5144"
    def _is_in_table(self, para) -> bool:
        try:
            return bool(para._p.xpath("./ancestor::w:tbl"))
        except Exception:
            return False

    def _collect_body_paragraphs(self, doc) -> list[object]:
        body: list[object] = []
        for para in doc.paragraphs:
            if self._is_in_table(para):
                continue
            body.append(para)
        return body

    def _insert_paragraph_after(self, para, text: str):
        new_p = self._OxmlElement("w:p")
        para._p.addnext(new_p)
        new_para = self._Paragraph(new_p, para._parent)
        new_para.text = text
        return new_para

    def _clear_paragraph(self, para) -> None:
        element = para._p
        for child in list(element):
            element.remove(child)

    def _append_run(self, para, text: str, rgb: tuple[int, int, int] | None = None):
        run = para.add_run(text)
        if rgb is not None:
            run.font.color.rgb = self._RGBColor(*rgb)
        return run

    def _sentences_with_spans(self, text: str) -> list[tuple[int, int, str]]:
        return _sentences_with_spans(text)
    def _parse_sentence_comments(self, comment: str) -> list[tuple[int, str]]:
        return _parse_sentence_comments(comment)
    def _extract_quoted_text(self, comment: str) -> list[str]:
        return _extract_quoted_text(comment)
    def _find_sentence_index_for_quote(self, sentences: list[tuple[int, int, str]], quote: str) -> int | None:
        for idx, (_start, _end, sent) in enumerate(sentences, start=1):
            if quote in sent:
                return idx
        return None

    def _add_red_note_after_paragraph(self, para, note: str) -> None:
        note = clean_comment_text(note)
        if not note:
            return
        prefix = " " if re.search(r"[A-Za-z0-9]$", para.text or "") else ""
        self._append_run(para, f"{prefix}\uff08\u5efa\u8bae\uff1a{note}\uff09", (0xFF, 0x00, 0x00))
    def _dedupe_notes(self, notes: list[str]) -> list[str]:
        seen: set[str] = set()
        ordered: list[str] = []
        for note in notes:
            cleaned = (note or "").strip()
            if not cleaned or cleaned in seen:
                continue
            seen.add(cleaned)
            ordered.append(cleaned)
        return ordered

    def _merge_notes(self, notes: list[str]) -> str:
        merged = self._dedupe_notes(notes)
        return "\uff1b".join(merged)
    def _apply_notes_to_paragraph(
        self,
        para,
        base_text: str,
        sentence_notes: dict[int, list[str]],
        paragraph_notes: list[str],
    ) -> None:
        if sentence_notes:
            sentences = self._sentences_with_spans(base_text)
            if not sentences:
                sentence_notes.clear()
            else:
                self._clear_paragraph(para)
                for idx, (_start, _end, sent) in enumerate(sentences, start=1):
                    self._append_run(para, sent)
                    merged = self._merge_notes(sentence_notes.get(idx, []))
                    if merged:
                        prefix = " " if re.search(r"[A-Za-z0-9]$", sent or "") else ""
                        self._append_run(
                            para,
                            f"{prefix}\uff08\u5efa\u8bae\uff1a{merged}\uff09",
                            (0xFF, 0x00, 0x00),
                        )
        merged_paragraph = self._merge_notes(paragraph_notes)
        if merged_paragraph:
            self._add_red_note_after_paragraph(para, merged_paragraph)
    def _collect_notes(self, base_text: str, comments: list[str]) -> tuple[dict[int, list[str]], list[str]]:
        sentence_notes: dict[int, list[str]] = {}
        paragraph_notes: list[str] = []
        sentences = self._sentences_with_spans(base_text)
        for comment in comments:
            comment = (comment or "").strip()
            if not comment:
                continue
            sentence_hits = self._parse_sentence_comments(comment)
            if sentence_hits:
                for idx, note in sentence_hits:
                    note = clean_comment_text(note)
                    if not note:
                        continue
                    sentence_notes.setdefault(idx, []).append(note)
                continue
            if re.search(r"(?:\u6574\u6bb5|[Pp]\s*\d+\s*-\s*\u6574\u6bb5)\s*[:\uff1a]", comment):
                note = re.sub(r"^[Pp]\s*\d+\s*-\s*\u6574\u6bb5\s*[:\uff1a\s]*", "", comment).strip()
                if note == comment.strip():
                    note = re.sub(r"^\u6574\u6bb5\s*[:\uff1a\s]*", "", comment).strip()
                paragraph_notes.append(clean_comment_text(note or comment))
                continue
            quotes = self._extract_quoted_text(comment)
            if quotes and sentences:
                matched = False
                for quote in quotes:
                    idx = self._find_sentence_index_for_quote(sentences, quote)
                    if idx is not None:
                        sentence_notes.setdefault(idx, []).append(clean_comment_text(comment))
                        matched = True
                if matched:
                    continue
            paragraph_notes.append(clean_comment_text(comment))
        for idx, notes in list(sentence_notes.items()):
            sentence_notes[idx] = self._dedupe_notes(notes)
        paragraph_notes = self._dedupe_notes(paragraph_notes)
        return sentence_notes, paragraph_notes
    def apply(
        self,
        input_path: str,
        output_path: str,
        instructions: Iterable[RevisionInstruction],
        *,
        progress_callback: Callable[[int, RevisionInstruction, bool], None] | None = None,
        start_index: int = 0,
    ) -> bool:
        instruction_list = list(instructions)
        doc = self._Document(input_path)

        body_paragraphs = self._collect_body_paragraphs(doc)
        instructions_by_para: dict[int, list[RevisionInstruction]] = {}
        for inst in instruction_list:
            instructions_by_para.setdefault(inst.paragraph_index, []).append(inst)

        for idx, para in enumerate(body_paragraphs):
            insts = instructions_by_para.get(idx)
            if not insts:
                continue
            base_text = para.text or ""
            comments: list[str] = []
            for inst in insts:
                if inst.action == "comment":
                    comments.append(inst.comment or inst.content or "Review note")
                    continue
                if inst.comment:
                    comments.append(inst.comment)
                    continue
                content = (inst.content or "").strip()
                if not content:
                    if inst.action == "delete":
                        comments.append("\u6574\u6bb5\uff1a\u5efa\u8bae\u5220\u9664\u8be5\u5904\u5185\u5bb9\u3002")
                    elif inst.action == "insert_after":
                        comments.append("\u6574\u6bb5\uff1a\u5efa\u8bae\u65b0\u589e\u76f8\u5173\u5185\u5bb9\u3002")
                    else:
                        comments.append("\u6574\u6bb5\uff1a\u5efa\u8bae\u6539\u5199\u8be5\u5904\u5185\u5bb9\u3002")
                    continue
                if inst.action == "insert_after":
                    comments.append(f"\u6574\u6bb5\uff1a\u5efa\u8bae\u65b0\u589e\uff1a{content}")
                elif inst.action == "delete":
                    comments.append(f"\u6574\u6bb5\uff1a\u5efa\u8bae\u5220\u9664\uff1a{content}")
                else:
                    comments.append(f"\u6574\u6bb5\uff1a\u5efa\u8bae\u6539\u4e3a\uff1a{content}")

            if comments:
                sentence_notes, paragraph_notes = self._collect_notes(base_text, comments)
                self._apply_notes_to_paragraph(para, base_text, sentence_notes, paragraph_notes)

        doc.save(output_path)
        if progress_callback is not None:
            for idx, inst in enumerate(instruction_list, start=1):
                try:
                    progress_callback(start_index + idx, inst, False)
                except Exception:
                    pass
        return True


def parse_instructions(revisions_json: str) -> list[RevisionInstruction]:
    try:
        raw = json.loads(revisions_json)
    except Exception as exc:  # noqa: BLE001
        raise ValueError("revisions_json must be a JSON list of instructions") from exc
    if not isinstance(raw, list):
        raise ValueError("revisions_json must be a JSON list of instructions")
    instructions: list[RevisionInstruction] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        try:
            paragraph_index = int(item.get("paragraph_index", 0))
        except Exception:
            paragraph_index = 0
        instructions.append(
            RevisionInstruction(
                action=str(item.get("action", "comment") or "comment"),
                paragraph_index=paragraph_index,
                content=item.get("content"),
                comment=item.get("comment"),
            )
        )
    return instructions


def _win32_disabled() -> bool:
    value = os.getenv("DISABLE_WIN32", "")
    if value.lower() in {"1", "true", "yes"}:
        return True
    return os.getenv("REVISION_ENGINE", "").lower() == "python-docx"


def _has_win32() -> bool:
    if _win32_disabled():
        return False
    if os.name != "nt":
        return False
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False
    return True


def _has_python_docx() -> bool:
    if importlib.util.find_spec("docx") is None:
        return False
    try:
        import docx  # type: ignore  # noqa: F401
    except Exception:
        return False
    return True


def _allow_python_docx_fallback() -> bool:
    value = os.getenv("ALLOW_PYTHON_DOCX_FALLBACK", "")
    return value.lower() in {"1", "true", "yes"}


def _available_engines() -> list[str]:
    engines: list[str] = []
    if _has_win32():
        engines.append("win32com")
    if _has_python_docx():
        engines.append("python-docx")
    return engines


def get_engine(name: str) -> BaseRevisionEngine:
    if name == "win32com":
        if not _has_win32():
            raise RuntimeError("win32com engine is unavailable; install pywin32 or use python-docx/auto")
        return Win32WordEngine()
    if name == "python-docx":
        if not _has_python_docx():
            raise RuntimeError("python-docx engine is unavailable; install python-docx or use win32com/auto")
        return PythonDocxEngine()
    if name == "auto":
        if _has_win32():
            return Win32WordEngine()
        if _has_python_docx() and _allow_python_docx_fallback():
            return PythonDocxEngine()
        raise RuntimeError(
            "win32com engine is unavailable; python-docx fallback is disabled "
            "(it appends red suggestions into the正文)."
        )
    raise RuntimeError(f"Unknown revision engine: {name}")


def _resolve_docx_input(input_real: Path) -> Path:
    if input_real.suffix.lower() == ".docx":
        return input_real
    if input_real.is_dir():
        docx_files = sorted(input_real.glob("*.docx"))
        if docx_files:
            return docx_files[0]
    if input_real.suffix.lower() == ".json":
        candidate = input_real.with_suffix(".docx")
        if candidate.exists():
            return candidate
        docx_files = sorted(input_real.parent.glob("*.docx"))
        if docx_files:
            return docx_files[0]
    raise FileNotFoundError(f"No .docx found for {input_real}")


def _env_flag(name: str, default: bool = False) -> bool:
    raw = os.getenv(name, "true" if default else "false").strip().lower()
    return raw in {"1", "true", "yes", "y", "on"}


def _instructions_signature(instructions: list[RevisionInstruction]) -> str:
    hasher = hashlib.sha1()
    for inst in instructions:
        payload = {
            "a": inst.action,
            "p": int(inst.paragraph_index),
            "c": inst.content or "",
            "m": inst.comment or "",
        }
        hasher.update(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8", errors="ignore"))
        hasher.update(b"\n")
    return hasher.hexdigest()


def _file_fingerprint(path: Path) -> str:
    hasher = hashlib.sha1()
    try:
        with open(path, "rb") as handle:
            while True:
                chunk = handle.read(1024 * 1024)
                if not chunk:
                    break
                hasher.update(chunk)
    except Exception:
        return hashlib.sha1(str(path).encode("utf-8", errors="ignore")).hexdigest()
    return hasher.hexdigest()


def _resume_progress_dir() -> Path:
    raw = os.getenv("REVIEW_RESUME_DIR", "").strip()
    if raw:
        path = Path(raw)
    else:
        path = ensure_workspace_dir() / "agent_state" / "revision_resume"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _load_json(path: Path) -> dict | None:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
    return data if isinstance(data, dict) else None


def _save_json(path: Path, payload: dict) -> None:
    ensure_parent(path)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)


def apply_revisions(input_path: str, output_path: str, revisions_json: str, engine: str) -> str:
    env_engine = os.getenv("REVISION_ENGINE", "").strip().lower()
    if env_engine == "python-docx":
        engine_key = "python-docx"
    else:
        engine_key = (engine or "").strip().lower() or "auto"
    allow_fallback = engine_key == "auto"
    instructions = parse_instructions(revisions_json)
    input_real = resolve_path(input_path)
    input_real = _resolve_docx_input(input_real)
    input_fingerprint = _file_fingerprint(input_real)
    output_real = resolve_path(output_path)
    ensure_parent(output_real)
    instruction_signature = _instructions_signature(instructions)

    resume_enabled = _env_flag("REVIEW_RESUME_ENABLED", default=True)
    resume_key = os.getenv("REVIEW_RESUME_KEY", "").strip()
    if not resume_key:
        resume_key = hashlib.sha1(
            f"{str(input_real.resolve())}|{instruction_signature}".encode("utf-8", errors="ignore")
        ).hexdigest()[:20]
    progress_path = _resume_progress_dir() / f"{resume_key}.apply.json"

    source_input_real = input_real
    start_index = 0
    structural_applied = False
    checkpoint_every_raw = os.getenv("REVIEW_APPLY_CHECKPOINT_EVERY", "1").strip()
    try:
        checkpoint_every = max(1, int(checkpoint_every_raw))
    except Exception:
        checkpoint_every = 1

    existing_progress = _load_json(progress_path) if resume_enabled else None
    if existing_progress:
        matches = (
            str(existing_progress.get("instruction_signature", "")) == instruction_signature
            and int(existing_progress.get("total", -1)) == len(instructions)
            and str(existing_progress.get("input_fingerprint", "")) == input_fingerprint
        )
        if matches:
            status = str(existing_progress.get("status", ""))
            cached_output_raw = str(existing_progress.get("partial_output_path", "") or "")
            cached_output = Path(cached_output_raw) if cached_output_raw else None
            if status == "completed" and cached_output and cached_output.exists():
                if cached_output.resolve() != output_real.resolve():
                    ensure_parent(output_real)
                    shutil.copy2(cached_output, output_real)
                print(f"[resume] completed cache hit: {cached_output}")
                return to_virtual_path(output_real)

            processed = int(existing_progress.get("processed_count", 0) or 0)
            structural_applied = bool(existing_progress.get("structural_applied", False))
            if (
                cached_output
                and cached_output.exists()
                and processed > 0
                and processed < len(instructions)
                and not structural_applied
            ):
                source_input_real = cached_output
                start_index = max(0, min(processed, len(instructions)))
                print(f"[resume] apply resume from {start_index}/{len(instructions)}")

    def _write_progress(status: str, *, processed_count: int, engine_name: str) -> None:
        if not resume_enabled:
            return
        payload = {
            "version": 1,
            "resume_key": resume_key,
            "status": status,
            "instruction_signature": instruction_signature,
            "total": len(instructions),
            "processed_count": int(max(0, min(processed_count, len(instructions)))),
            "input_fingerprint": input_fingerprint,
            "source_input_path": str(source_input_real),
            "partial_output_path": str(output_real),
            "engine": engine_name,
            "structural_applied": bool(structural_applied),
            "updated_at": dt.datetime.now().isoformat(timespec="seconds"),
        }
        try:
            _save_json(progress_path, payload)
        except Exception:
            pass

    log_path = os.getenv("REVISION_LOG_PATH", "").strip()
    if log_path:
        try:
            with open(log_path, "a", encoding="utf-8") as handle:
                handle.write(
                    f"[apply_revisions] engine={engine_key} items={len(instructions)} "
                    f"resume_key={resume_key} start_index={start_index}\n"
                )
                for idx, inst in enumerate(instructions, start=1):
                    content = (inst.content or "").replace("\n", " ")
                    comment = (inst.comment or "").replace("\n", " ")
                    handle.write(
                        f"{idx:04d} action={inst.action} "
                        f"paragraph_index={inst.paragraph_index} "
                        f"content={content[:80]} "
                        f"comment={comment[:120]}\n"
                    )
        except Exception:
            pass

    if start_index >= len(instructions):
        if source_input_real.exists() and source_input_real.resolve() != output_real.resolve():
            ensure_parent(output_real)
            shutil.copy2(source_input_real, output_real)
        _write_progress("completed", processed_count=len(instructions), engine_name=engine_key)
        return to_virtual_path(output_real)

    processed_count = start_index
    current_engine_name = engine_key

    def _progress_callback(next_index: int, inst: RevisionInstruction, structural_changed: bool) -> None:
        nonlocal processed_count, structural_applied
        try:
            resolved_next = int(next_index)
        except Exception:
            resolved_next = processed_count
        processed_count = max(processed_count, min(resolved_next, len(instructions)))
        if structural_changed:
            structural_applied = True
        if processed_count == len(instructions) or (processed_count - start_index) % checkpoint_every == 0:
            _write_progress("running", processed_count=processed_count, engine_name=current_engine_name)

    try:
        revision_engine = get_engine(engine_key)
    except Exception as exc:  # noqa: BLE001
        print(f"[engine] {engine_key} init failed: {exc}")
        if not allow_fallback:
            raise
        available = _available_engines()
        fallback_name = available[0] if available else None
        if fallback_name is None:
            raise
        if fallback_name == "python-docx" and not _allow_python_docx_fallback():
            raise
        print(f"[engine] falling back to {fallback_name}")
        revision_engine = get_engine(fallback_name)
    current_engine_name = revision_engine.name
    _write_progress("running", processed_count=processed_count, engine_name=current_engine_name)
    try:
        revision_engine.apply(
            str(source_input_real),
            str(output_real),
            instructions[start_index:],
            progress_callback=_progress_callback,
            start_index=start_index,
        )
        processed_count = len(instructions)
        _write_progress("completed", processed_count=processed_count, engine_name=current_engine_name)
    except Exception as exc:  # noqa: BLE001
        print(f"[engine] {revision_engine.name} failed: {exc}")
        if not allow_fallback:
            _write_progress("failed", processed_count=processed_count, engine_name=current_engine_name)
            raise
        fallback_name = None
        for candidate in _available_engines():
            if candidate != revision_engine.name:
                fallback_name = candidate
                break
        if fallback_name is None:
            _write_progress("failed", processed_count=processed_count, engine_name=current_engine_name)
            raise
        if fallback_name == "python-docx" and not _allow_python_docx_fallback():
            _write_progress("failed", processed_count=processed_count, engine_name=current_engine_name)
            raise
        print(f"[engine] falling back to {fallback_name}")
        fallback = get_engine(fallback_name)
        current_engine_name = fallback.name
        fallback_start_index = processed_count
        fallback_input = source_input_real
        if processed_count > start_index and output_real.exists() and not structural_applied:
            fallback_input = output_real
            fallback_start_index = processed_count
            print(f"[resume] fallback continue from {fallback_start_index}/{len(instructions)}")
        elif processed_count > start_index:
            fallback_start_index = start_index
            print("[resume] fallback restart from last stable point (structural edits detected)")
        _write_progress("running", processed_count=fallback_start_index, engine_name=current_engine_name)
        fallback.apply(
            str(fallback_input),
            str(output_real),
            instructions[fallback_start_index:],
            progress_callback=_progress_callback,
            start_index=fallback_start_index,
        )
        processed_count = len(instructions)
        _write_progress("completed", processed_count=processed_count, engine_name=current_engine_name)
    return to_virtual_path(output_real)
