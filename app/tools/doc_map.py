from __future__ import annotations

from dataclasses import dataclass
import json
import os
from pathlib import Path
import re
from typing import Any

from docx import Document

from app.tools.path_utils import ensure_parent, resolve_path, to_virtual_path
from app.tools.win32_utils import (
    com_retry,
    dispatch_word_application,
    get_win32_constants,
    install_ole_message_filter,
    restore_ole_message_filter,
)


@dataclass
class SectionMap:
    title: str
    level: int
    paragraphs: list[str]


@dataclass
class DocumentMap:
    sections: list[SectionMap]

    def to_dict(self) -> dict[str, Any]:
        return {
            "sections": [
                {"title": s.title, "level": s.level, "paragraphs": s.paragraphs}
                for s in self.sections
            ]
        }


@dataclass
class IndexedParagraph:
    index: int
    text: str


@dataclass
class IndexedSection:
    title: str
    level: int
    paragraphs: list[IndexedParagraph]


def _heading_level(style_name: str) -> int | None:
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name in {"title", "subtitle"}:
        return 1
    if name.startswith("heading"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return int(parts[1])
        match = re.search(r"heading\s*(\d+)", name)
        if match:
            return int(match.group(1))
        return 1
    if style_name.strip().startswith("标题"):
        parts = style_name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return int(parts[1])
        match = re.search(r"标题\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
        return 1
    if style_name.strip().startswith("副标题"):
        return 2
    return None


_HEADING_TEXT_PATTERN = re.compile(
    r"^("
    r"(第[一二三四五六七八九十零〇两0-9]+[章节部分篇条])"
    r"|([一二三四五六七八九十]+[、.．])"
    r"|([（(][一二三四五六七八九十]+[)）])"
    r"|(\d+(?:\.\d+)+)"
    r"|(\d+[、.．)）])"
    r"|((?:附录|结论|摘要|引言|前言|目录|致谢))"
    r")"
)
_SENTENCE_PATTERN = re.compile(r"[^。！？!?；;]+[。！？!?；;]?")


def _split_sentences(text: str) -> list[str]:
    return [m.group(0).strip() for m in _SENTENCE_PATTERN.finditer(text or "") if m.group(0).strip()]


def _heading_level_from_text(text: str) -> int | None:
    text = (text or "").strip()
    if not text:
        return None
    if len(text) > 40:
        return None
    if re.search(r"[。！？!?；;]$", text):
        return None
    if re.search(r"[。！？!?；;]", text):
        return None
    if text.count("，") >= 2:
        return None
    if not _HEADING_TEXT_PATTERN.match(text):
        return None
    match = re.match(r"^第([一二三四五六七八九十零〇两0-9]+)(章|篇)", text)
    if match:
        return 1
    match = re.match(r"^第([一二三四五六七八九十零〇两0-9]+)(节|部分)", text)
    if match:
        return 2
    match = re.match(r"^第([一二三四五六七八九十零〇两0-9]+)条", text)
    if match:
        return 3
    match = re.match(r"^(\d+(?:\.\d+)+)", text)
    if match:
        return min(4, match.group(1).count(".") + 1)
    if re.match(r"^\d+[、.．)）]", text):
        return 2
    if re.match(r"^[一二三四五六七八九十]+[、.．]", text):
        return 2
    if re.match(r"^[（(][一二三四五六七八九十]+[)）]", text):
        return 3
    if re.match(r"^(附录|结论|摘要|引言|前言|目录|致谢)$", text):
        return 1
    return None


def _infer_heading_level_from_context(prev_text: str, text: str, next_text: str) -> int | None:
    level = _heading_level_from_text(text)
    if level is None:
        return None
    next_text = (next_text or "").strip()
    if not next_text:
        return None
    next_level = _heading_level_from_text(next_text)
    if next_level is not None:
        return level
    if len(next_text) < 10:
        return None
    return level


_CAPTION_STYLE_HINTS = ("caption", "图注", "表注", "题注", "图表", "表题", "图题")
_CAPTION_TEXT_PATTERN = re.compile(
    r"^\s*(?:图|表|Figure|Table)\s*"
    r"(?:\d+(?:[.-]\d+)*|[一二三四五六七八九十]+)"
    r"\s*[:：.\-–]?\s*\S+"
)

_DOCX_LIKE_SUFFIXES = {".docx", ".docm", ".dotx", ".dotm"}


def _is_caption_paragraph(text: str, style_name: str) -> bool:
    if style_name:
        style_lower = style_name.strip().lower()
        for hint in _CAPTION_STYLE_HINTS:
            if hint in style_lower or hint in style_name:
                return True
    if text and _CAPTION_TEXT_PATTERN.match(text.strip()):
        return True
    return False


def _find_docx_path(input_real: Path) -> Path:
    if input_real.suffix.lower() in _DOCX_LIKE_SUFFIXES:
        return input_real
    if input_real.suffix.lower() == ".json" and input_real.exists():
        return input_real
    candidate = input_real.with_suffix(".docx")
    if candidate.exists():
        return candidate
    if input_real.is_dir():
        docx_files = sorted(input_real.glob("*.docx"))
        if docx_files:
            return docx_files[0]
    if input_real.parent.exists():
        docx_files = sorted(input_real.parent.glob("*.docx"))
        if docx_files:
            return docx_files[0]
    raise FileNotFoundError(f"No .docx found for {input_real}")


def _is_in_table(para) -> bool:
    try:
        return bool(para._p.xpath("./ancestor::w:tbl"))
    except Exception:
        return False


def _win32_disabled() -> bool:
    value = os.getenv("DISABLE_WIN32", "")
    if value.lower() in {"1", "true", "yes"}:
        return True
    return os.getenv("REVISION_ENGINE", "").lower() == "python-docx"


def _has_win32() -> bool:
    if _win32_disabled():
        return False
    if os.name != "nt":
        return False
    try:
        import win32com.client  # type: ignore  # noqa: F401
    except Exception:
        return False
    return True


def _win32_heading_level(para, constants, style_name: str) -> int | None:
    outline_level = None
    try:
        outline_level = int(getattr(para, "OutlineLevel", 0))
    except Exception:
        outline_level = None
    if outline_level is not None and outline_level > 0:
        try:
            if constants is None:
                body_level = 10
            else:
                body_level = int(constants.wdOutlineLevelBodyText)
        except Exception:
            body_level = 10
        if outline_level != body_level:
            return outline_level
    return _heading_level(style_name)


def _neighbor_text(records: list[dict], idx: int, direction: int) -> str:
    step = idx + direction
    while 0 <= step < len(records):
        record = records[step]
        if record.get("is_caption"):
            step += direction
            continue
        candidate = (record.get("text") or "").strip()
        if candidate:
            return candidate
        step += direction
    return ""


def _extract_paragraphs_win32(docx_path: Path) -> list[str]:
    if not _has_win32():
        return []
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
    except Exception:
        return []
    pythoncom.CoInitialize()
    old_filter, message_filter = install_ole_message_filter()
    word = None
    doc = None
    paragraphs: list[str] = []
    try:
        word = com_retry(dispatch_word_application)
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            constants = get_win32_constants(win32)
        except Exception:
            constants = None

        doc = com_retry(lambda: word.Documents.Open(str(docx_path)))
        total = doc.Paragraphs.Count
        records: list[dict] = []
        for idx in range(1, total + 1):
            para = doc.Paragraphs(idx)
            try:
                if constants is None:
                    in_table = False
                else:
                    in_table = bool(para.Range.Information(constants.wdWithInTable))
            except Exception:
                in_table = False
            if in_table:
                continue
            text = (para.Range.Text or "").strip()
            style_name = ""
            try:
                style = para.Range.Style
                style_name = str(getattr(style, "NameLocal", "") or getattr(style, "Name", ""))
            except Exception:
                style_name = ""
            is_caption = bool(text) and _is_caption_paragraph(text, style_name)
            records.append(
                {
                    "para": para,
                    "text": text,
                    "style_name": style_name,
                    "is_caption": is_caption,
                }
            )
        for idx, record in enumerate(records):
            text = record["text"]
            style_name = record["style_name"]
            if not text or record["is_caption"]:
                continue
            prev_text = _neighbor_text(records, idx, -1)
            next_text = _neighbor_text(records, idx, 1)
            level = _win32_heading_level(record["para"], constants, style_name)
            if level is None:
                level = _infer_heading_level_from_context(prev_text, text, next_text)
            if level is not None:
                continue
            paragraphs.append(text)
    except Exception:
        paragraphs = []
    finally:
        if doc is not None:
            try:
                com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
            except Exception:
                pass
        if word is not None:
            try:
                com_retry(lambda: word.Quit(), timeout_s=5.0)
            except Exception:
                pass
        if message_filter is not None:
            restore_ole_message_filter(old_filter)
        pythoncom.CoUninitialize()
    return paragraphs


def build_indexed_sections(input_path: str) -> list[IndexedSection]:
    input_real = resolve_path(input_path)
    docx_path = _find_docx_path(input_real)
    if docx_path.suffix.lower() == ".json" and docx_path.exists():
        raise ValueError("Indexed sections require a .docx input")

    if _has_win32():
        win32_sections = _extract_indexed_sections_win32(docx_path)
        if win32_sections and any(section.paragraphs for section in win32_sections):
            return win32_sections

    doc = Document(str(docx_path))
    sections: list[IndexedSection] = []
    current = IndexedSection(title="Document", level=0, paragraphs=[])
    body_index = 0
    records: list[dict] = []
    for para in doc.paragraphs:
        if _is_in_table(para):
            continue
        text = (para.text or "").strip()
        style_name = para.style.name if para.style else ""
        is_caption = bool(text) and _is_caption_paragraph(text, style_name)
        records.append(
            {
                "text": text,
                "style_name": style_name,
                "is_caption": is_caption,
            }
        )
    for idx, record in enumerate(records):
        text = record["text"]
        style_name = record["style_name"]
        if not text:
            body_index += 1
            continue
        if record["is_caption"]:
            body_index += 1
            continue
        prev_text = _neighbor_text(records, idx, -1)
        next_text = _neighbor_text(records, idx, 1)
        level = _heading_level(style_name)
        if level is None:
            level = _infer_heading_level_from_context(prev_text, text, next_text)
        if level is not None:
            if current.paragraphs or current.title != "Document":
                sections.append(current)
            current = IndexedSection(title=text or "Untitled", level=level, paragraphs=[])
            body_index += 1
            continue
        current.paragraphs.append(IndexedParagraph(index=body_index, text=text))
        body_index += 1

    if current.paragraphs or current.title:
        sections.append(current)
    if not sections or all(not section.paragraphs for section in sections):
        sections = _extract_indexed_sections_win32(docx_path)
    return sections


def _extract_indexed_sections_win32(docx_path: Path) -> list[IndexedSection]:
    if not _has_win32():
        return []
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
    except Exception:
        return []
    pythoncom.CoInitialize()
    old_filter, message_filter = install_ole_message_filter()
    word = None
    doc = None
    sections: list[IndexedSection] = []
    current = IndexedSection(title="Document", level=0, paragraphs=[])
    body_index = 0
    try:
        word = com_retry(dispatch_word_application)
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            constants = get_win32_constants(win32)
        except Exception:
            constants = None

        doc = com_retry(lambda: word.Documents.Open(str(docx_path)))
        total = doc.Paragraphs.Count
        records: list[dict] = []
        for idx in range(1, total + 1):
            para = doc.Paragraphs(idx)
            try:
                if constants is None:
                    in_table = False
                else:
                    in_table = bool(para.Range.Information(constants.wdWithInTable))
            except Exception:
                in_table = False
            if in_table:
                continue
            text = (para.Range.Text or "").strip()
            style_name = ""
            try:
                style = para.Range.Style
                style_name = str(getattr(style, "NameLocal", "") or getattr(style, "Name", ""))
            except Exception:
                style_name = ""
            is_caption = bool(text) and _is_caption_paragraph(text, style_name)
            records.append(
                {
                    "para": para,
                    "text": text,
                    "style_name": style_name,
                    "is_caption": is_caption,
                }
            )
        for idx, record in enumerate(records):
            text = record["text"]
            style_name = record["style_name"]
            if not text:
                body_index += 1
                continue
            if record["is_caption"]:
                body_index += 1
                continue
            prev_text = _neighbor_text(records, idx, -1)
            next_text = _neighbor_text(records, idx, 1)
            level = _win32_heading_level(record["para"], constants, style_name)
            if level is None:
                level = _infer_heading_level_from_context(prev_text, text, next_text)
            if level is not None:
                if current.paragraphs or current.title != "Document":
                    sections.append(current)
                current = IndexedSection(title=text or "Untitled", level=level, paragraphs=[])
                body_index += 1
                continue
            current.paragraphs.append(IndexedParagraph(index=body_index, text=text))
            body_index += 1
        if current.paragraphs or current.title:
            sections.append(current)
    except Exception:
        sections = []
    finally:
        if doc is not None:
            try:
                com_retry(lambda: doc.Close(SaveChanges=False), timeout_s=5.0)
            except Exception:
                pass
        if word is not None:
            try:
                com_retry(lambda: word.Quit(), timeout_s=5.0)
            except Exception:
                pass
        if message_filter is not None:
            restore_ole_message_filter(old_filter)
        pythoncom.CoUninitialize()
    return sections


def build_doc_map(input_path: str, output_path: str | None = None) -> str:
    input_real = resolve_path(input_path)
    docx_path = _find_docx_path(input_real)
    if docx_path.suffix.lower() == ".json" and docx_path.exists():
        return to_virtual_path(docx_path)

    if _has_win32():
        win32_sections = _extract_indexed_sections_win32(docx_path)
        if win32_sections and any(section.paragraphs for section in win32_sections):
            sections = [
                SectionMap(
                    title=section.title,
                    level=section.level,
                    paragraphs=[item.text for item in section.paragraphs],
                )
                for section in win32_sections
            ]
            doc_map = DocumentMap(sections=sections)
            if output_path:
                output_real = resolve_path(output_path)
            else:
                output_real = Path(f"{docx_path}.map.json")
            ensure_parent(output_real)
            output_real.write_text(
                json.dumps(doc_map.to_dict(), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            return to_virtual_path(output_real)

    doc = Document(str(docx_path))
    sections: list[SectionMap] = []
    current = SectionMap(title="Document", level=0, paragraphs=[])

    for para in doc.paragraphs:
        if _is_in_table(para):
            continue
        style_name = para.style.name if para.style else ""
        level = _heading_level(style_name)
        if level is not None:
            if current.paragraphs or current.title != "Document":
                sections.append(current)
            current = SectionMap(title=para.text.strip() or "Untitled", level=level, paragraphs=[])
            continue
        text = para.text.strip()
        if text and _is_caption_paragraph(text, style_name):
            continue
        if text:
            current.paragraphs.append(text)

    if current.paragraphs or current.title:
        sections.append(current)

    if not sections or all(not section.paragraphs for section in sections):
        paragraphs = _extract_paragraphs_win32(docx_path)
        if paragraphs:
            sections = [SectionMap(title="Document", level=0, paragraphs=paragraphs)]

    doc_map = DocumentMap(sections=sections)
    if output_path:
        output_real = resolve_path(output_path)
    else:
        output_real = Path(f"{docx_path}.map.json")
    ensure_parent(output_real)
    output_real.write_text(
        json.dumps(doc_map.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return to_virtual_path(output_real)


def load_doc_map(map_path: str) -> DocumentMap:
    map_real = resolve_path(map_path)
    data = json.loads(map_real.read_text(encoding="utf-8"))
    sections = [
        SectionMap(title=s["title"], level=s["level"], paragraphs=s.get("paragraphs", []))
        for s in data.get("sections", [])
    ]
    return DocumentMap(sections=sections)


def extract_section_text(input_path: str, section_index: int) -> str:
    sections = build_indexed_sections(input_path)
    if not sections:
        return ""
    if section_index < 0 or section_index >= len(sections):
        section_index = max(0, min(section_index, len(sections) - 1))
    section = sections[section_index]
    lines = [f"# {section.title}"]
    for item in section.paragraphs:
        lines.append(f"{item.index}: {item.text}")
        sentences = _split_sentences(item.text)
        if sentences:
            lines.append("SENTS:")
            for idx, sentence in enumerate(sentences, start=1):
                lines.append(f"SENT_{idx} (P{item.index}-S{idx}): {sentence}")
    return "\n".join(lines)
