from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.diagnostics_service import build_review_diagnostics_payload, write_review_diagnostics


def _build_sample_docx(path: Path, *, year: str, total_value: str, acronym: str) -> None:
    doc = Document()
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(f"本文围绕乡村治理展开研究，采用 {acronym} 方法进行分析。")
    doc.add_heading("研究背景", level=1)
    doc.add_paragraph(f"研究始于 {year} 年，样本总量为 {total_value}。")
    doc.add_heading("研究内容", level=1)
    doc.add_paragraph("图1 展示治理框架。")
    doc.add_paragraph("图2 结果图。")
    doc.add_paragraph("表1 指标体系。")
    doc.add_heading("创新点", level=1)
    doc.add_paragraph("创新点在于引入跨章节比较。")
    doc.add_heading("结论", level=1)
    doc.add_paragraph("结论认为该路径具有推广价值。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph(f"张三，{year}，《乡村治理研究》。")
    doc.save(path)


def test_build_review_diagnostics_payload_detects_change_risks(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("DISABLE_WIN32", "1")
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _build_sample_docx(input_path, year="2022", total_value="120", acronym="GIS")
    _build_sample_docx(output_path, year="2024", total_value="140", acronym="AHP")

    payload = build_review_diagnostics_payload(
        input_path=str(input_path),
        output_path=str(output_path),
        preset_key="general_academic",
    )

    assert payload["preset"]["key"] == "general_academic"
    assert "overview" in payload
    assert "citation_reference_check" in payload["pre_review"]
    assert "section_structure_score" in payload["pre_review"]
    assert payload["change_risk"]["added_years"]
    assert payload["change_risk"]["removed_years"]
    assert payload["change_risk"]["acronym_drift"]


def test_write_review_diagnostics_writes_json_file(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("DISABLE_WIN32", "1")
    input_path = tmp_path / "paper.docx"
    output_path = tmp_path / "paper_revised.docx"
    _build_sample_docx(input_path, year="2021", total_value="100", acronym="GIS")
    _build_sample_docx(output_path, year="2021", total_value="100", acronym="GIS")

    diagnostics_path = write_review_diagnostics(
        input_path=str(input_path),
        output_path=str(output_path),
        preset_key="literature_review",
    )

    assert diagnostics_path.exists()
    assert diagnostics_path.suffix == ".json"
    content = diagnostics_path.read_text(encoding="utf-8")
    assert "literature_review" in content
